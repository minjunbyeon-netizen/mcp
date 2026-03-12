#!/usr/bin/env python3
"""
공통 유틸리티 모듈
LoadingSpinner, JSON 추출, API 키 로드, 파일 텍스트 추출 등 공통 기능
"""

import os
import base64
import json
import re
import struct
import subprocess
import tempfile
import threading
import time
import logging
from pathlib import Path
from datetime import datetime


def setup_logger(name: str = "auto_blog", log_dir: Path = None) -> logging.Logger:
    """
    [M6] 공통 로거 설정.

    - INFO 레벨: 콘솔 + 파일 출력
    - DEBUG 레벨: 파일 전용 출력
    - ERROR 레벨: 콘솔 + 파일 출력

    Args:
        name: 로거 이름 (기본 "auto_blog")
        log_dir: 로그 저장 폴더 (기본 프로젝트 루트/logs)

    Returns:
        logging.Logger 인스턴스
    """
    logger = logging.getLogger(name)

    # 중복 핸들러 방지
    if logger.handlers:
        return logger

    logger.setLevel(logging.DEBUG)

    # 로그 디렉토리 설정
    if log_dir is None:
        log_dir = Path(__file__).parent / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)

    log_filename = f"auto_blog_{datetime.now().strftime('%Y-%m-%d')}.log"
    log_file = log_dir / log_filename

    # 포맷터
    file_fmt = logging.Formatter(
        "[%(asctime)s] %(levelname)s %(name)s — %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S"
    )
    console_fmt = logging.Formatter("%(levelname)s: %(message)s")

    # 파일 핸들러 (DEBUG 이상 모두 기록)
    fh = logging.FileHandler(log_file, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(file_fmt)

    # 콘솔 핸들러 (INFO 이상만)
    ch = logging.StreamHandler()
    ch.setLevel(logging.INFO)
    ch.setFormatter(console_fmt)

    logger.addHandler(fh)
    logger.addHandler(ch)

    return logger


class LoadingSpinner:
    """로딩 스피너 애니메이션"""
    def __init__(self, message="처리 중"):
        self.message = message
        self.running = False
        self.thread = None

    def start(self):
        self.running = True
        self.thread = threading.Thread(target=self._animate)
        self.thread.start()

    def _animate(self):
        frames = ['|', '/', '-', '\\']
        i = 0
        while self.running:
            print(f"\r  {frames[i % 4]} {self.message}...", end="", flush=True)
            time.sleep(0.2)
            i += 1

    def stop(self, success_msg="완료"):
        self.running = False
        if self.thread:
            self.thread.join()
        print(f"\r  [OK] {success_msg}" + " " * 20)


def extract_json_from_response(response_text: str) -> str:
    """AI 응답에서 JSON 문자열 추출"""
    if "```json" in response_text:
        response_text = response_text.split("```json")[1].split("```")[0]
    elif "```" in response_text:
        response_text = response_text.split("```")[1].split("```")[0]
    return response_text.strip()


def parse_json_response(response_text: str) -> dict:
    """AI 응답에서 JSON을 추출하고 파싱"""
    cleaned = extract_json_from_response(response_text)

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        # 잘못된 이스케이프 문자 정리 후 재시도
        cleaned = re.sub(r'\\(?!["\\/bfnrtu])', r'\\\\', cleaned)
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            # 최종 시도: 제어 문자 제거
            cleaned = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', cleaned)
            return json.loads(cleaned)


def load_api_key(key_name: str = "GEMINI_API_KEY") -> str | None:
    """환경변수 또는 mcp_config.json에서 API 키 로드"""
    from dotenv import load_dotenv
    load_dotenv()

    api_key = os.getenv(key_name)
    if api_key:
        return api_key

    config_path = Path(__file__).parent / "mcp_config.json"
    if config_path.exists():
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
            for server in config.get("mcpServers", {}).values():
                key = server.get("env", {}).get(key_name)
                if key:
                    os.environ[key_name] = key
                    return key

    return None


_SAFE_TEXT_PUNCTUATION = set(".,:;!?()[]{}<>/%&@#*+-_=~'\"`|\\·•※○●△▲▽▼→←↑↓…")


def _is_hangul_char(ch: str) -> bool:
    cp = ord(ch)
    return (
        0xAC00 <= cp <= 0xD7A3 or
        0x3131 <= cp <= 0x318E
    )


def _is_cjk_char(ch: str) -> bool:
    cp = ord(ch)
    return (
        0x3400 <= cp <= 0x4DBF or
        0x4E00 <= cp <= 0x9FFF or
        0xF900 <= cp <= 0xFAFF
    )


def _is_safe_display_char(ch: str, allow_cjk: bool = False) -> bool:
    if ch in "\n\r\t":
        return True
    if ch in _SAFE_TEXT_PUNCTUATION:
        return True
    if " " <= ch <= "~":
        return True
    if _is_hangul_char(ch):
        return True
    if allow_cjk and _is_cjk_char(ch):
        return True
    return False


def sanitize_text_for_display(text: str, allow_cjk: bool = False) -> str:
    """
    화면 표시용 텍스트 정제.

    - 제어 문자와 비정상 유니코드 노이즈 제거
    - 줄바꿈은 보존
    - 공백은 과하게 누적되지 않도록 정리
    """
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    cleaned = "".join(
        ch if _is_safe_display_char(ch, allow_cjk=allow_cjk) else " "
        for ch in text
    )
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r" *\n *", "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def is_meaningful_text_line(text: str) -> bool:
    """잡음 라인을 제외하기 위한 휴리스틱."""
    stripped = (text or "").strip()
    if len(stripped) < 2:
        return False
    if re.fullmatch(r"[-=_.~#* ]{3,}", stripped):
        return False

    meaningful = len(re.findall(r"[A-Za-z0-9가-힣]", stripped))
    hangul = len(re.findall(r"[가-힣]", stripped))
    digits = len(re.findall(r"\d", stripped))
    letters = len(re.findall(r"[A-Za-z]", stripped))
    tokens = stripped.split()
    single_char_tokens = sum(1 for token in tokens if len(token) == 1)

    if meaningful < 2:
        return False
    if meaningful / max(len(stripped), 1) < 0.35:
        return False
    if hangul == 0 and digits == 0 and letters < 3:
        return False
    if letters >= 10 and hangul == 0 and digits == 0 and "http" not in stripped.lower():
        return False
    if len(tokens) >= 4 and single_char_tokens > len(tokens) // 2:
        return False
    return True


def _dedupe_preserve_order(items: list[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        key = re.sub(r"\s+", "", item).lower()
        if key and key not in seen:
            seen.add(key)
            result.append(item)
    return result


# ============================================================
# [M-5] 공통 파일 텍스트 추출 함수 (web/app.py에서 통합)
# 지원 형식: .txt, .pdf (PyMuPDF 우선 → pdfplumber fallback),
#            .docx, .hwp, .jpg/.jpeg/.png
# ============================================================

# HWP 지원 여부
try:
    import olefile as _olefile
    import zlib as _zlib
    _HWP_SUPPORTED = True
except ImportError:
    _HWP_SUPPORTED = False


def _run_windows_ocr(image_path: Path) -> str:
    """
    Windows 기본 OCR(WinRT)을 이용해 이미지에서 텍스트 추출.

    Tesseract 설치 없이도 Windows 환경에서 동작하도록 설계한다.
    실패 시 빈 문자열 반환.
    """
    if os.name != "nt":
        return ""

    image_path = Path(image_path).resolve()
    script = f"""
$ErrorActionPreference='Stop'
[Console]::OutputEncoding = [System.Text.Encoding]::UTF8
Add-Type -AssemblyName System.Runtime.WindowsRuntime
$null = [Windows.Storage.StorageFile, Windows.Storage, ContentType = WindowsRuntime]
$null = [Windows.Storage.FileAccessMode, Windows.Storage, ContentType = WindowsRuntime]
$null = [Windows.Graphics.Imaging.BitmapDecoder, Windows.Graphics.Imaging, ContentType = WindowsRuntime]
$null = [Windows.Graphics.Imaging.SoftwareBitmap, Windows.Graphics.Imaging, ContentType = WindowsRuntime]
$null = [Windows.Storage.Streams.IRandomAccessStream, Windows.Storage.Streams, ContentType = WindowsRuntime]
$null = [Windows.Media.Ocr.OcrEngine, Windows.Foundation, ContentType = WindowsRuntime]
$null = [Windows.Media.Ocr.OcrResult, Windows.Foundation, ContentType = WindowsRuntime]
function Await([object]$Operation, [type]$ResultType) {{
  $asTaskGeneric = [System.WindowsRuntimeSystemExtensions].GetMethods() |
    Where-Object {{ $_.Name -eq 'AsTask' -and $_.IsGenericMethod -and $_.GetParameters().Count -eq 1 }} |
    Select-Object -First 1
  $asTask = $asTaskGeneric.MakeGenericMethod($ResultType)
  $task = $asTask.Invoke($null, @($Operation))
  $task.Wait()
  return $task.Result
}}
$path = {json.dumps(str(image_path), ensure_ascii=False)}
$file = Await ([Windows.Storage.StorageFile]::GetFileFromPathAsync($path)) ([Windows.Storage.StorageFile])
$stream = Await ($file.OpenAsync([Windows.Storage.FileAccessMode]::Read)) ([Windows.Storage.Streams.IRandomAccessStream])
$decoder = Await ([Windows.Graphics.Imaging.BitmapDecoder]::CreateAsync($stream)) ([Windows.Graphics.Imaging.BitmapDecoder])
$bitmap = Await ($decoder.GetSoftwareBitmapAsync()) ([Windows.Graphics.Imaging.SoftwareBitmap])
$engine = [Windows.Media.Ocr.OcrEngine]::TryCreateFromUserProfileLanguages()
$result = Await ($engine.RecognizeAsync($bitmap)) ([Windows.Media.Ocr.OcrResult])
Write-Output $result.Text
"""

    encoded_script = base64.b64encode(script.encode("utf-16le")).decode("ascii")

    try:
        result = subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-EncodedCommand", encoded_script],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            timeout=45,
        )
    except Exception:
        return ""

    if result.returncode != 0:
        return ""
    return result.stdout.strip()


def _extract_text_from_image(file_path: Path) -> str:
    """이미지 파일 OCR."""
    try:
        from rapidocr_onnxruntime import RapidOCR

        engine = RapidOCR()
        result, _ = engine(str(file_path))
        if result:
            lines = [item[1] for item in result if len(item) >= 2 and item[1]]
            text = "\n".join(lines).strip()
            if text:
                return text
    except Exception:
        pass

    text = _run_windows_ocr(file_path)
    if text:
        return text

    try:
        import pytesseract
        from PIL import Image

        return pytesseract.image_to_string(Image.open(file_path), lang="kor+eng").strip()
    except Exception:
        return ""


def _ocr_pdf_with_windows(file_path: Path, max_pages: int = 8) -> str:
    """텍스트가 거의 없는 스캔 PDF를 이미지 OCR로 보완."""
    temp_dir = Path(tempfile.mkdtemp(prefix="pdf_ocr_"))
    ocr_texts = []

    try:
        import fitz

        doc = fitz.open(str(file_path))
        page_count = min(len(doc), max_pages)
        for page_index in range(page_count):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_path = temp_dir / f"page_{page_index + 1}.png"
            pix.save(str(image_path))
            text = _extract_text_from_image(image_path)
            if text.strip():
                ocr_texts.append(text.strip())
        doc.close()
    except Exception:
        return ""
    finally:
        for image_file in temp_dir.glob("*"):
            try:
                image_file.unlink()
            except Exception:
                pass
        try:
            temp_dir.rmdir()
        except Exception:
            pass

    return "\n\n".join(ocr_texts).strip()


def _iter_hwp_records(raw: bytes):
    """압축 해제된 HWP section에서 레코드를 순회."""
    offset = 0
    total = len(raw)
    while offset + 4 <= total:
        header = struct.unpack_from("<I", raw, offset)[0]
        tag_id = header & 0x3FF
        level = (header >> 10) & 0x3FF
        size = (header >> 20) & 0xFFF
        offset += 4
        if size == 0xFFF:
            if offset + 4 > total:
                break
            size = struct.unpack_from("<I", raw, offset)[0]
            offset += 4
        if offset + size > total:
            break
        yield tag_id, level, raw[offset:offset + size]
        offset += size


def _clean_hwp_text_chunk(text: str) -> str:
    """HWP PARA_TEXT 레코드에서 사람이 읽을 수 있는 텍스트만 남긴다."""
    cleaned = sanitize_text_for_display(text, allow_cjk=False)
    cleaned = re.sub(r"^[A-Za-z]{4,}(?=[가-힣])", "", cleaned)
    cleaned = re.sub(r"(?<=[가-힣])[A-Za-z]{4,}$", "", cleaned)

    if re.search(r"[가-힣]", cleaned):
        tokens = []
        for token in cleaned.split():
            if re.fullmatch(r"[a-z]{4,}", token):
                continue
            tokens.append(token)
        cleaned = " ".join(tokens)

    cleaned = re.sub(r"\s+", " ", cleaned).strip(" -|\t")
    return cleaned


def _extract_text_from_hwp(file_path: Path) -> str:
    """HWP 본문에서 PARA_TEXT 레코드만 읽어 텍스트를 추출."""
    if not _HWP_SUPPORTED:
        raise ValueError("HWP 지원을 위해 'pip install olefile'를 실행하세요.")

    text_parts = []
    try:
        ole = _olefile.OleFileIO(str(file_path))
        for stream in ole.listdir():
            if "BodyText" not in stream or not any(part.startswith("Section") for part in stream):
                continue
            try:
                data = ole.openstream(stream).read()
                raw = _zlib.decompress(data, -15)
            except Exception:
                continue

            for tag_id, _level, payload in _iter_hwp_records(raw):
                if tag_id != 67:
                    continue
                try:
                    chunk = payload.decode("utf-16-le", errors="ignore")
                except Exception:
                    continue
                chunk = _clean_hwp_text_chunk(chunk)
                if is_meaningful_text_line(chunk):
                    text_parts.append(chunk)
        ole.close()
    except Exception as e:
        raise ValueError(f"HWP 파일 읽기 실패: {e}")

    return "\n".join(_dedupe_preserve_order(text_parts))


def extract_text_from_file(file_path: Path) -> str:
    """
    다양한 파일 형식에서 텍스트 추출 (공통 유틸).

    지원 형식:
    - .txt  : UTF-8 텍스트
    - .pdf  : PyMuPDF(fitz) 우선, 실패 시 pdfplumber fallback
    - .docx : python-docx (단락 + 표)
    - .hwp  : olefile + zlib 압축 해제
    - .jpg/.jpeg/.png : 이미지 경로 마커 반환 (Gemini Vision 처리용)

    Returns:
        추출된 텍스트 문자열
    Raises:
        ValueError: 지원하지 않는 파일 형식 또는 라이브러리 미설치
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()

    if ext == '.txt':
        for enc in ('utf-8', 'cp949', 'euc-kr'):
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
        raise ValueError(f"TXT 파일 인코딩을 인식할 수 없습니다: {file_path}")

    elif ext == '.pdf':
        text = ""
        # 1차 시도: PyMuPDF (fitz) — 레이아웃 보존, 빠름
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(file_path))
            for page in doc:
                text += page.get_text("text") + "\n"
            doc.close()
        except Exception as e:
            print("  PDF 읽는 중입니다... (레이아웃 방식 전환)")
            # 2차 시도: pdfplumber (백업)
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        text = text.strip()
        if len(re.sub(r"\s+", "", text)) < 80:
            ocr_text = _ocr_pdf_with_windows(file_path)
            if ocr_text:
                text = ocr_text
        return text

    elif ext == '.docx':
        text = ""
        try:
            import docx
            doc = docx.Document(file_path)
            # 단락 텍스트
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            # 표(Table) 데이터
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        text += " | ".join(row_data) + "\n"
        except Exception as e:
            raise ValueError(f"DOCX 파일 읽기 실패: {e}")
        return text

    elif ext == '.hwp':
        return _extract_text_from_hwp(file_path)

    elif ext in ['.jpg', '.jpeg', '.png']:
        return _extract_text_from_image(file_path)

    else:
        raise ValueError(f"지원되지 않는 파일 형식: {ext}")
