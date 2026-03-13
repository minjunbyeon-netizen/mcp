#!/usr/bin/env python3
"""
블로그 자동화 도구 웹 대시보드 - Flask 백엔드
사용법: python app.py
"""

import sys
import os
import json
import io
import tempfile
try:
    import truststore
    truststore.inject_into_ssl()
except ImportError:
    pass  # Linux/Mac — 시스템 인증서 사용
from pathlib import Path
from datetime import datetime
from flask import Flask, request, jsonify, send_from_directory, session, redirect, url_for
from flask_cors import CORS
from authlib.integrations.flask_client import OAuth
from functools import wraps
import pdfplumber
import requests as http_requests
import docx
import fitz  # PyMuPDF
import re

# Windows 터미널 UTF-8 출력 설정
if sys.platform == 'win32' and not isinstance(sys.stdout, io.TextIOWrapper):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

# 프로젝트 루트 설정
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

# 환경 변수 로드
from dotenv import load_dotenv
load_dotenv()

# mcp_config.json에서 API 키 로드
config_path = PROJECT_ROOT / "mcp_config.json"
if config_path.exists():
    with open(config_path, 'r', encoding='utf-8') as f:
        config = json.load(f)
        for _section in config.get("mcpServers", {}).values():
            env_config = _section.get("env", {})
            if not os.getenv("GEMINI_API_KEY") and env_config.get("GEMINI_API_KEY"):
                os.environ["GEMINI_API_KEY"] = env_config["GEMINI_API_KEY"]
            if not os.getenv("UNSPLASH_ACCESS_KEY") and env_config.get("UNSPLASH_ACCESS_KEY"):
                os.environ["UNSPLASH_ACCESS_KEY"] = env_config["UNSPLASH_ACCESS_KEY"]

# HWP 지원
try:
    import olefile
    import zlib
    HWP_SUPPORTED = True
except ImportError:
    HWP_SUPPORTED = False

# Flask 앱 설정
app = Flask(__name__, static_folder='.', static_url_path='')
CORS(app)

# 리버스 프록시 서브패스 지원 (nginx /mcp/ → 앱 내부 경로 보정)
from werkzeug.middleware.proxy_fix import ProxyFix
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)

# APPLICATION_ROOT 환경변수로 서브패스 설정 가능 (예: /mcp)
_app_root = os.getenv('APPLICATION_ROOT', '').rstrip('/')
if _app_root:
    app.config['APPLICATION_ROOT'] = _app_root
    app.config['SESSION_COOKIE_PATH'] = _app_root + '/'

# 세션 시크릿 키 (SSO용)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'dev-secret-key-change-in-production')

# OAuth 설정
oauth = OAuth(app)

# Google OAuth 설정 (환경 변수에서 로드)
GOOGLE_CLIENT_ID = os.getenv('GOOGLE_CLIENT_ID')
GOOGLE_CLIENT_SECRET = os.getenv('GOOGLE_CLIENT_SECRET')

if GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET:
    google = oauth.register(
        name='google',
        client_id=GOOGLE_CLIENT_ID,
        client_secret=GOOGLE_CLIENT_SECRET,
        authorize_url='https://accounts.google.com/o/oauth2/v2/auth',
        access_token_url='https://oauth2.googleapis.com/token',
        jwks_uri='https://www.googleapis.com/oauth2/v3/certs',
        userinfo_endpoint='https://openidconnect.googleapis.com/v1/userinfo',
        client_kwargs={
            'scope': 'openid email profile https://www.googleapis.com/auth/documents',
        },
    )
    SSO_ENABLED = True
    print("[OK] Google OAuth SSO 활성화됨")
else:
    google = None
    SSO_ENABLED = False
    print("[INFO] Google OAuth 미설정 (GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET 환경 변수 필요)")

# 허용된 이메일 도메인 (쉼표로 구분)
ALLOWED_DOMAINS = os.getenv('ALLOWED_DOMAINS', '').split(',')
ALLOWED_DOMAINS = [d.strip() for d in ALLOWED_DOMAINS if d.strip()]

# 허용된 이메일 목록 — 환경변수 + allowed_emails.txt 파일 병합
_ALLOWED_EMAILS_FILE = PROJECT_ROOT / "allowed_emails.txt"

def _load_allowed_emails() -> set:
    """환경변수 ALLOWED_EMAILS + allowed_emails.txt 를 합쳐 반환 (소문자 정규화)"""
    emails = set()
    # 1) 환경변수
    for e in os.getenv('ALLOWED_EMAILS', '').split(','):
        e = e.strip().lower()
        if e:
            emails.add(e)
    # 2) 파일
    if _ALLOWED_EMAILS_FILE.exists():
        for line in _ALLOWED_EMAILS_FILE.read_text(encoding='utf-8').splitlines():
            e = line.strip().lower()
            if e and not e.startswith('#'):
                emails.add(e)
    return emails

def _save_allowed_emails(emails: set):
    """allowed_emails.txt 에 현재 목록 저장"""
    lines = sorted(emails)
    _ALLOWED_EMAILS_FILE.write_text('\n'.join(lines) + '\n', encoding='utf-8')

ALLOWED_EMAILS: set = _load_allowed_emails()  # 시작 시 1회 로드

# 경로 설정
OUTPUT_DIR = Path.home() / "mcp-data" / "outputs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
BLOG_COLLECTIONS_DIR = PROJECT_ROOT / "blog_pull" / "output"
BLOG_COLLECTIONS_DIR.mkdir(parents=True, exist_ok=True)
DNA_DIR = OUTPUT_DIR / "dna"
DNA_DIR.mkdir(parents=True, exist_ok=True)
BUSINESS_DIR = OUTPUT_DIR / "business"
BUSINESS_DIR.mkdir(parents=True, exist_ok=True)
CALIBRATIONS_DIR = OUTPUT_DIR / "calibrations"
CALIBRATIONS_DIR.mkdir(parents=True, exist_ok=True)

# blog_pull 모듈 경로 추가
sys.path.insert(0, str(PROJECT_ROOT / "blog_pull"))
# web/ 폴더도 sys.path에 추가 (run_crawler.py 등 web/ 내 모듈 import용)
sys.path.insert(0, str(Path(__file__).parent))

# 스타일 템플릿 로드
_STYLE_TEMPLATES_PATH = Path(__file__).parent / "style_templates.json"
with open(_STYLE_TEMPLATES_PATH, 'r', encoding='utf-8') as _f:
    STYLE_TEMPLATES: list[dict] = json.load(_f)
_STYLE_TEMPLATES_MAP: dict[str, dict] = {t["id"]: t for t in STYLE_TEMPLATES}

# Google Gemini API 클라이언트
from google import genai

def _normalize_naver_blog_url(url: str) -> str:
    """네이버 블로그 URL을 PostView URL로 변환 (JS 렌더링 우회)"""
    # blog.naver.com/USER/POST_NO 또는 m.blog.naver.com/USER/POST_NO 패턴
    m = re.match(
        r'https?://(?:m\.)?blog\.naver\.com/([^/?#]+)/(\d+)',
        url
    )
    if m:
        blog_id, log_no = m.group(1), m.group(2)
        return f'https://blog.naver.com/PostView.naver?blogId={blog_id}&logNo={log_no}'
    return url


def fetch_url_text(url: str) -> str:
    """URL에서 텍스트 추출 (HTML 페이지 or PDF URL)"""
    # 네이버 블로그는 JS 렌더링 — PostView URL로 변환
    fetch_url = _normalize_naver_blog_url(url)

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept-Language': 'ko-KR,ko;q=0.9,en;q=0.8',
        'Referer': 'https://blog.naver.com/',
    }
    try:
        resp = http_requests.get(fetch_url, headers=headers, timeout=20, stream=True)
        resp.raise_for_status()
        content_type = resp.headers.get('Content-Type', '')

        if 'application/pdf' in content_type or url.lower().endswith('.pdf'):
            # PDF URL → 임시 파일로 저장 후 pdfplumber 추출
            import pdfplumber
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                for chunk in resp.iter_content(chunk_size=8192):
                    tmp.write(chunk)
                tmp_path = Path(tmp.name)
            try:
                with pdfplumber.open(tmp_path) as pdf:
                    pages = [p.extract_text() or '' for p in pdf.pages[:15]]
                    return '\n'.join(pages)[:10000]
            finally:
                tmp_path.unlink(missing_ok=True)
        else:
            # HTML 페이지 → 태그 제거 후 텍스트 추출
            html = resp.text
            # script/style 제거
            html = re.sub(r'<(script|style)[^>]*>.*?</(script|style)>', '', html, flags=re.DOTALL | re.IGNORECASE)
            # HTML 태그 제거
            text = re.sub(r'<[^>]+>', ' ', html)
            # 연속 공백/줄바꿈 정리
            text = re.sub(r'[ \t]+', ' ', text)
            text = re.sub(r'\n{3,}', '\n\n', text)
            text = text.strip()
            if len(text) < 100:
                raise ValueError(f"추출된 텍스트가 너무 짧습니다 ({len(text)}자). JS 렌더링 페이지일 수 있습니다.")
            return text[:10000]
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"URL 크롤링 실패 ({url}): {e}")


def _strip_markdown(text: str) -> str:
    """블로그 본문에서 마크다운 기호를 모두 제거"""
    import re as _re
    # ** bold **, __bold__
    text = _re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
    text = _re.sub(r'_{1,2}(.*?)_{1,2}', r'\1', text)
    # ## 헤딩
    text = _re.sub(r'^#{1,6}\s+', '', text, flags=_re.MULTILINE)
    # ``` 코드블록
    text = _re.sub(r'```.*?```', '', text, flags=_re.DOTALL)
    text = _re.sub(r'`([^`]*)`', r'\1', text)
    # <태그>
    text = _re.sub(r'<[^>]+>', '', text)
    # --- === 구분선
    text = _re.sub(r'^[-=]{3,}\s*$', '', text, flags=_re.MULTILINE)
    # > 인용
    text = _re.sub(r'^>\s+', '', text, flags=_re.MULTILINE)
    return text.strip()


def parse_ai_json(text):
    """AI 응답에서 JSON을 안전하게 추출 및 파싱"""
    if not text:
        return {}
    
    # 1. 마크다운 코드 블록 제거
    if "```json" in text:
        text = text.split("```json")[-1].split("```")[0]
    elif "```" in text:
        text = text.split("```")[-1].split("```")[0]
    
    text = text.strip()
    
    try:
        # 2. 일반적인 파싱 시도
        return json.loads(text)
    except json.JSONDecodeError:
        # 3. 이스케이프 문자 등 정제 후 재시도
        # 제어 문자 및 잘못된 백슬래시 처리
        # (단일 백슬래시가 줄을 깨는 경우가 많음)
        # JSON 문자열 내의 단일 \를 \\로 변경하려고 시도 (매우 조심스럽게)
        # 단, 이미 정당한 이스케이프(\", \\, \/, \b, \f, \n, \r, \t, \uXXXX)는 보존해야 함
        
        # 간단한 정제: 불필요한 제어 문자 제거 (strict=False로도 어느 정도 해결됨)
        try:
            return json.loads(text, strict=False)
        except json.JSONDecodeError as e:
            # 마지막 수단: 텍스트를 더 공격적으로 정제
            # 예: "content": "blah \n blah" -> "content": "blah \\n blah"
            # 하지만 이미 json.loads(strict=False)가 \n 등은 어느 정도 허용함.
            # 문제는 문법에 어긋나는 \ 하나임.
            # 정규식으로 유효하지 않은 \를 찾아 \\로 변환 (lookahead 사용)
            # 유효한 이스케이프 시퀀스가 아닌 \ 를 찾아냄
            cleaned = re.sub(r'\\(?![/"\\bfnrtu])', r'\\\\', text)
            try:
                return json.loads(cleaned, strict=False)
            except:
                print(f"[CRITICAL] JSON 파싱 최종 실패: {e}")
                raise


# ============================================================
# File Text Extraction — [M-5] utils.py로 통합, 여기서는 re-export
# ============================================================
from utils import extract_text_from_file  # noqa: F401
from blog_storage import (
    build_blog_package,
    ensure_blog_package_shape,
    load_blog_package,
    save_blog_package,
    update_blog_package_version,
)
from material_pipeline import build_material_bundle, build_material_bundle_from_paths
from offline_engines import generate_blog_versions_offline


UPLOADS_DIR = Path(__file__).parent / "uploads"
UPLOADS_DIR.mkdir(exist_ok=True)


def save_uploaded_file(file) -> Path:
    """업로드된 파일을 임시 저장"""
    ext = Path(file.filename).suffix.lower()
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    file.save(temp_file.name)
    return Path(temp_file.name)


def extract_images_from_file(path: Path) -> list[str]:
    """파일에서 이미지를 추출해 uploads/에 저장, URL 경로 목록 반환"""
    import uuid
    ext = path.suffix.lower()
    saved = []

    try:
        # ── 이미지 파일 자체 ──────────────────────────
        if ext in {'.jpg', '.jpeg', '.png', '.gif', '.webp'}:
            uid = uuid.uuid4().hex[:12]
            dst = UPLOADS_DIR / f"{uid}{ext}"
            import shutil; shutil.copy2(path, dst)
            saved.append(f"/uploads/{dst.name}")

        # ── PDF → PyMuPDF ────────────────────────────
        elif ext == '.pdf':
            import fitz
            doc = fitz.open(str(path))
            for page_num in range(len(doc)):
                page = doc[page_num]
                for img_index, img in enumerate(page.get_images(full=True)):
                    xref = img[0]
                    base_img = doc.extract_image(xref)
                    img_bytes = base_img["image"]
                    img_ext   = base_img.get("ext", "png")
                    # 너무 작은 이미지(아이콘 등) 제외
                    if len(img_bytes) < 10_000:
                        continue
                    uid = uuid.uuid4().hex[:12]
                    dst = UPLOADS_DIR / f"{uid}.{img_ext}"
                    dst.write_bytes(img_bytes)
                    saved.append(f"/uploads/{dst.name}")
            doc.close()

        # ── DOCX → python-docx ───────────────────────
        elif ext in {'.docx'}:
            import zipfile as _zip
            with _zip.ZipFile(str(path)) as zf:
                for name in zf.namelist():
                    if name.startswith("word/media/"):
                        img_ext = Path(name).suffix.lower()
                        if img_ext not in {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.emf', '.wmf'}:
                            continue
                        if img_ext in {'.emf', '.wmf'}:
                            continue
                        img_bytes = zf.read(name)
                        if len(img_bytes) < 10_000:
                            continue
                        uid = uuid.uuid4().hex[:12]
                        save_ext = img_ext if img_ext != '.jpeg' else '.jpg'
                        dst = UPLOADS_DIR / f"{uid}{save_ext}"
                        dst.write_bytes(img_bytes)
                        saved.append(f"/uploads/{dst.name}")

    except Exception as e:
        print(f"[WARN] 이미지 추출 실패 ({path.name}): {e}")

    return saved


def upload_to_gemini(path: Path, client: genai.Client):
    """Gemini File API에 파일 업로드 (PDF / 이미지 네이티브 지원)"""
    MIME_MAP = {
        '.pdf':  'application/pdf',
        '.jpg':  'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png':  'image/png',
        '.gif':  'image/gif',
        '.webp': 'image/webp',
    }
    ext = path.suffix.lower()
    mime_type = MIME_MAP.get(ext)
    if not mime_type:
        print(f"[INFO] {ext} → 텍스트 추출 방식 사용")
        return None
    try:
        from google.genai import types as _gtypes
        gfile = client.files.upload(
            file=str(path),
            config=_gtypes.UploadFileConfig(mime_type=mime_type),
        )
        print(f"[OK] Gemini 파일 업로드: {path.name} → {gfile.name}")
        return gfile
    except Exception as e:
        print(f"[WARN] Gemini 파일 업로드 실패 ({path.name}): {e}")
        return None


# ============================================================
# SSO Authentication (Google OAuth)
# ============================================================

def login_required(f):
    """로그인 필수 데코레이터"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not SSO_ENABLED:
            # SSO 미설정 시 바로 통과
            return f(*args, **kwargs)
        if 'user' not in session:
            return jsonify({"error": "로그인이 필요합니다.", "login_required": True}), 401
        return f(*args, **kwargs)
    return decorated_function


@app.route('/login')
def login():
    """Google 로그인 시작"""
    if not SSO_ENABLED:
        # Google OAuth 미설정 — 개발 모드: 자동 세션 생성 후 앱 진입
        session['user'] = {
            'email': 'dev@localhost',
            'name': '개발자',
            'picture': ''
        }
        return redirect(url_for('index'))
    redirect_uri = url_for('callback', _external=True)
    return google.authorize_redirect(redirect_uri, access_type='offline', prompt='consent')


@app.route('/callback')
def callback():
    """Google 로그인 콜백"""
    if not SSO_ENABLED:
        return redirect(url_for('index'))
    
    try:
        token = google.authorize_access_token()
        user_info = token.get('userinfo', {})
        
        email = user_info.get('email', '')
        email_lower = email.lower()
        domain = email.split('@')[-1] if '@' in email else ''

        # ── 접근 제어 ──────────────────────────────────────────
        # 최신 허용 목록 다시 로드 (파일이 변경될 수 있음)
        current_allowed = _load_allowed_emails()

        # 허용 조건: 이메일 목록에 있거나 / 도메인 목록에 있거나 / 둘 다 미설정(오픈)
        email_ok   = bool(current_allowed) and email_lower in current_allowed
        domain_ok  = bool(ALLOWED_DOMAINS) and domain in ALLOWED_DOMAINS
        lists_empty = not current_allowed and not ALLOWED_DOMAINS  # 제한 없음

        if not lists_empty and not email_ok and not domain_ok:
            print(f"[DENY] 접근 거부: {email}")
            return f'''<!doctype html>
<html lang="ko">
<head><meta charset="utf-8"><title>접근 거부</title>
<style>
  body{{font-family:-apple-system,Helvetica Neue,sans-serif;
       display:flex;align-items:center;justify-content:center;
       min-height:100vh;margin:0;background:#f5f5f7;}}
  .card{{background:#fff;border-radius:12px;padding:48px 40px;
         text-align:center;max-width:400px;box-shadow:0 1px 4px rgba(0,0,0,.08);}}
  h2{{font-size:20px;font-weight:700;color:#1d1d1f;margin:0 0 12px}}
  p{{font-size:14px;color:#6e6e73;margin:0 0 8px;line-height:1.6}}
  .email{{font-weight:700;color:#1d1d1f}}
  a{{display:inline-block;margin-top:24px;padding:10px 24px;
     background:#000;color:#fff;text-decoration:none;
     border-radius:980px;font-size:14px;}}
</style>
</head>
<body>
<div class="card">
  <h2>접근 권한 없음</h2>
  <p>이 서비스에 접근 권한이 없는 계정입니다.</p>
  <p class="email">{email}</p>
  <p>서비스 관리자에게 접근 권한을 요청하세요.</p>
  <a href="/">돌아가기</a>
</div>
</body>
</html>''', 403
        
        # 세션에 사용자 정보 저장
        session['user'] = {
            'email': email,
            'name': user_info.get('name', ''),
            'picture': user_info.get('picture', '')
        }
        # Google API 호출을 위한 액세스 토큰 + 갱신 토큰 저장
        import time as _time
        session['google_token'] = token.get('access_token', '')
        session['google_refresh_token'] = token.get('refresh_token', '')
        session['google_token_expires_at'] = _time.time() + token.get('expires_in', 3600)
        
        print(f"[OK] 로그인 성공: {email}")
        return redirect(url_for('index'))
        
    except Exception as e:
        print(f"[ERROR] 로그인 실패: {e}")
        return f"로그인 실패: {str(e)}", 500


@app.route('/logout')
def logout():
    """로그아웃"""
    user_email = session.get('user', {}).get('email', 'unknown')
    session.pop('user', None)
    print(f"[INFO] 로그아웃: {user_email}")
    return redirect(url_for('index'))


@app.route('/api/auth/status')
def auth_status():
    """현재 인증 상태 반환"""
    if 'user' in session:
        return jsonify({
            "logged_in": True,
            "sso_enabled": SSO_ENABLED,
            "user": session['user']
        })
    else:
        return jsonify({
            "logged_in": False,
            "sso_enabled": SSO_ENABLED,
            "allowed_domains": ALLOWED_DOMAINS
        })


# ── 허용 이메일 관리 API ─────────────────────────────────────

@app.route('/api/admin/allowed-emails', methods=['GET'])
@login_required
def get_allowed_emails():
    """허용 이메일 목록 조회 (로그인한 사용자 누구나 조회 가능)"""
    emails = _load_allowed_emails()
    return jsonify({
        "emails": sorted(emails),
        "count": len(emails),
        "domains": ALLOWED_DOMAINS,
    })


@app.route('/api/admin/allowed-emails', methods=['POST'])
@login_required
def add_allowed_email():
    """허용 이메일 추가"""
    data = request.json or {}
    new_email = data.get('email', '').strip().lower()
    if not new_email or '@' not in new_email:
        return jsonify({"error": "유효한 이메일 주소를 입력하세요."}), 400
    emails = _load_allowed_emails()
    if new_email in emails:
        return jsonify({"message": "이미 등록된 이메일입니다.", "emails": sorted(emails)})
    emails.add(new_email)
    _save_allowed_emails(emails)
    print(f"[ADMIN] 허용 이메일 추가: {new_email} (by {session['user']['email']})")
    return jsonify({"message": f"{new_email} 추가됨", "emails": sorted(emails)})


@app.route('/api/admin/allowed-emails/<path:email>', methods=['DELETE'])
@login_required
def delete_allowed_email(email):
    """허용 이메일 제거"""
    email = email.strip().lower()
    emails = _load_allowed_emails()
    if email not in emails:
        return jsonify({"error": "등록되지 않은 이메일입니다."}), 404
    # 자기 자신은 제거 불가
    if email == session.get('user', {}).get('email', '').lower():
        return jsonify({"error": "자신의 이메일은 제거할 수 없습니다."}), 400
    emails.discard(email)
    _save_allowed_emails(emails)
    print(f"[ADMIN] 허용 이메일 제거: {email} (by {session['user']['email']})")
    return jsonify({"message": f"{email} 제거됨", "emails": sorted(emails)})


# ============================================================
# Static Files (Frontend)
# ============================================================

NO_CACHE_HEADERS = {
    'Cache-Control': 'no-store, no-cache, must-revalidate, max-age=0',
    'Pragma': 'no-cache',
    'Expires': '0',
}

@app.route('/')
def index():
    resp = send_from_directory('.', 'index.html')
    for k, v in NO_CACHE_HEADERS.items():
        resp.headers[k] = v
    return resp

@app.route('/uploads/<path:filename>')
def serve_upload(filename):
    return send_from_directory(str(UPLOADS_DIR), filename)

@app.route('/<path:filename>')
def static_files(filename):
    resp = send_from_directory('.', filename)
    if filename in ('app.js', 'style.css', 'index.html'):
        for k, v in NO_CACHE_HEADERS.items():
            resp.headers[k] = v
    return resp


# ============================================================
# API: Style Templates
# ============================================================

@app.route('/api/style-templates', methods=['GET'])
def get_style_templates():
    """10가지 블로그 스타일 템플릿 목록 반환"""
    return jsonify({"templates": STYLE_TEMPLATES})


# 페르소나 기능은 C:\work\email-persona 프로젝트로 이전됨

@app.route('/api/persona/list', methods=['GET'])
@login_required
def list_personas():
    """[이전됨] 페르소나 목록 — 항상 빈 목록 반환 (email-persona 프로젝트로 이전)"""
    return jsonify({"personas": []})


@app.route('/api/persona/get', methods=['GET'])
@login_required
def get_persona():
    return jsonify({"error": "페르소나 기능은 이메일 도구로 이전되었습니다."}), 410


@app.route('/api/persona/extract', methods=['POST'])
@login_required
def extract_persona():
    return jsonify({"error": "페르소나 기능은 이메일 도구로 이전되었습니다."}), 410




# ============================================================
# API: Blog Generator
# ============================================================

@app.route('/api/blog/dna-preview', methods=['GET'])
@login_required
def blog_dna_preview():
    """DNA 분석 결과 미리보기 — 모방 스타일 가이드 + 샘플 글 반환"""
    blog_id = request.args.get("blog_id", "")
    if not blog_id:
        return jsonify({"error": "blog_id 필요"}), 400

    dna_analysis = None
    if DNA_DIR.exists():
        candidates = sorted(
            [f for f in DNA_DIR.glob(f"DNA_{blog_id}_*.json")],
            key=lambda f: f.stat().st_mtime, reverse=True
        )
        if candidates:
            with open(candidates[0], 'r', encoding='utf-8') as f:
                dna_analysis = json.load(f)

    # 원본 글 샘플 1개 로드
    sample_post = None
    if BLOG_COLLECTIONS_DIR.exists():
        all_posts = []
        for item in BLOG_COLLECTIONS_DIR.iterdir():
            if item.is_dir() and not item.name.startswith('.'):
                data_file = item / "_data.json"
                if data_file.exists():
                    try:
                        with open(data_file, 'r', encoding='utf-8') as f:
                            col = json.load(f)
                        if col.get("blog_id") == blog_id:
                            all_posts.extend(col.get("posts", []))
                    except Exception:
                        pass
        all_posts.sort(key=lambda x: x.get('addDate', ''), reverse=True)
        seen = set()
        for p in all_posts:
            if p.get("url") not in seen:
                seen.add(p.get("url"))
                if p.get("content", "").strip():
                    sample_post = p
                    break

    if not dna_analysis and not sample_post:
        return jsonify({"error": "DNA 분석 결과 없음. 먼저 글쓰기 DNA 분석을 실행하세요."}), 404

    # 미리보기 텍스트 구성
    lines = []
    if dna_analysis:
        c1 = dna_analysis.get("c1_template_structure", {})
        c2 = dna_analysis.get("c2_tone_mood", {})
        c3 = dna_analysis.get("c3_speech_style", {})
        c5 = dna_analysis.get("c5_frequent_expressions", {})
        c6 = dna_analysis.get("c6_sentence_patterns", {})
        c9 = dna_analysis.get("c9_opening_closing", {})
        c10 = dna_analysis.get("c10_visual_formatting", {})

        lines.append("▌ 글 구조 패턴")
        lines.append(f"  {c1.get('overall_pattern', '-')}")
        if c1.get('section_flow'):
            lines.append("  흐름: " + " → ".join(c1['section_flow'])[:120])

        lines.append("")
        lines.append("▌ 톤 & 어투")
        lines.append(f"  톤: {c2.get('primary_tone', '-')}  /  격식도: {c2.get('formality_level', '-')}/10")
        lines.append(f"  종결어미: {', '.join(c3.get('ending_patterns', []))}")
        lines.append(f"  독자 호칭: {c3.get('reader_address', '-')}")

        lines.append("")
        lines.append("▌ 이모지 & 특수기호")
        emoji_list = ', '.join(c10.get('emoji_types', []))
        lines.append(f"  이모지({c10.get('emoji_usage', '-')}/10): {emoji_list}")
        special = c10.get('special_symbols', c10.get('special_formatting', []))
        if special:
            lines.append(f"  특수기호: {', '.join(special)}")

        c11 = dna_analysis.get("c11_length_stats", {})
        c12 = dna_analysis.get("c12_typography", {})
        c13 = dna_analysis.get("c13_brackets_quotes", {})
        c14 = dna_analysis.get("c14_title_patterns", {})

        lines.append("")
        lines.append("▌ 시각적 포맷팅 (HTML 분석)")
        lines.append(f"  중앙정렬: {c10.get('center_align', '-')}")
        lines.append(f"  볼드 패턴: {c10.get('bold_pattern', '-')}")
        lines.append(f"  이탤릭: {c10.get('italic_usage', '-')}")
        lines.append(f"  줄바꿈: {c10.get('line_break_style', '-')}")
        if c10.get('text_colors'):
            lines.append(f"  강조 색상: {', '.join(c10['text_colors'][:4])}")
        if c10.get('highlight_colors'):
            lines.append(f"  하이라이트: {', '.join(c10['highlight_colors'][:3])}")
        if c10.get('writing_guide'):
            lines.append(f"\n  작성 가이드: {c10['writing_guide']}")

        if c11:
            lines.append("")
            lines.append("▌ 글자수/분량")
            lines.append(f"  글당 평균 글자수: {c11.get('avg_chars_per_post', '-')}")
            lines.append(f"  평균 문장 수: {c11.get('avg_sentences_per_post', '-')} / 문장당: {c11.get('avg_chars_per_sentence', '-')}")
            lines.append(f"  서론:본론:결론 = {c11.get('content_ratio', '-')}")
            if c11.get('writing_density_guide'):
                lines.append(f"  분량 지침: {c11['writing_density_guide']}")

        if c12:
            lines.append("")
            lines.append("▌ 폰트/글꼴 스타일")
            if c12.get('font_families'):
                lines.append(f"  폰트: {', '.join(c12['font_families'])}")
            lines.append(f"  본문 크기: {c12.get('base_font_size', '-')} / 소제목: {c12.get('heading_font_size', '-')}")
            lines.append(f"  볼드 빈도: {c12.get('bold_frequency', '-')}/10  목적: {c12.get('bold_purpose', '-')}")
            lines.append(f"  기울임: {c12.get('italic_usage', '-')} / 밑줄: {c12.get('underline_usage', '-')}")
            if c12.get('font_guide'):
                lines.append(f"  글꼴 지침: {c12['font_guide']}")

        if c13:
            lines.append("")
            lines.append("▌ 꺽쇠/괄호/인용부호")
            if c13.get('angle_bracket_types'):
                lines.append(f"  꺽쇠 종류: {', '.join(c13['angle_bracket_types'])}  빈도: {c13.get('angle_bracket_frequency', '-')}")
                lines.append(f"  꺽쇠 목적: {c13.get('angle_bracket_purpose', '-')}")
            if c13.get('square_bracket_types'):
                lines.append(f"  대괄호: {', '.join(c13['square_bracket_types'])}  목적: {c13.get('square_bracket_purpose', '-')}")
            lines.append(f"  소괄호 패턴: {c13.get('round_bracket_usage', '-')}")
            lines.append(f"  따옴표 방식: {c13.get('quotation_mark_style', '-')}")
            if c13.get('examples'):
                for ex in c13['examples'][:2]:
                    lines.append(f"    예: {ex}")

        if c14:
            lines.append("")
            lines.append("▌ 제목 패턴")
            lines.append(f"  평균 길이: {c14.get('avg_title_length', '-')}  구조: {c14.get('title_structure', '-')}")
            lines.append(f"  숫자 활용: {c14.get('number_usage', '-')} / 감정 후크: {c14.get('emotion_hook', '-')}")
            if c14.get('examples'):
                for ex in c14['examples'][:3]:
                    lines.append(f"    {ex}")

        lines.append("")
        lines.append("▌ 시그니처 표현")
        for phrase in c5.get('signature_phrases', [])[:6]:
            lines.append(f"  • {phrase}")

        lines.append("")
        lines.append("▌ 도입부 예시")
        for ex in c9.get('opening_examples', [])[:2]:
            lines.append(f"  {ex}")

        lines.append("")
        lines.append("▌ 마무리 예시")
        for ex in c9.get('closing_examples', [])[:2]:
            lines.append(f"  {ex}")

    if sample_post:
        lines.append("")
        lines.append("━" * 36)
        lines.append(f"▌ 실제 글 샘플 (AI가 모방하는 글)")
        lines.append(f"  제목: {sample_post.get('title', '')}")
        lines.append(f"  날짜: {sample_post.get('addDate', '')}")
        lines.append("")
        lines.append(sample_post.get('content', '')[:1800])

    # c10 시각 스타일 → JS가 HTML 변환에 쓸 구조화 데이터
    dna_styles = {}
    if dna_analysis:
        c10 = dna_analysis.get("c10_visual_formatting", {})
        # 폰트 패밀리 (Naver CSS font-family 값)
        raw_font = c10.get("font_family", "")
        font_map = {
            "나눔마루부리": "NanumMyeongjo, '나눔마루부리', serif",
            "나눔고딕": "'나눔고딕', NanumGothic, sans-serif",
            "나눔명조": "NanumMyeongjo, '나눔명조', serif",
            "맑은 고딕": "'맑은 고딕', MalgunGothic, sans-serif",
            "돋움": "Dotum, '돋움', sans-serif",
            "굴림": "Gulim, '굴림', sans-serif",
        }
        css_font = font_map.get(raw_font, font_map.get(raw_font.split()[0], "'맑은 고딕', sans-serif"))

        # 중앙정렬 비율 파싱 (숫자 or "57%" 형태 모두 처리)
        center_raw = c10.get("center_align", "0")
        try:
            center_val = float(str(center_raw).replace("%", "").strip())
            if center_val > 1:
                center_val = center_val / 100
        except Exception:
            center_val = 0.0

        # 강조 색상 (첫번째 유효색)
        text_colors = c10.get("text_colors", [])
        accent_color = text_colors[0] if text_colors else ""
        highlight_colors = c10.get("highlight_colors", [])
        highlight_color = highlight_colors[0] if highlight_colors else ""

        dna_styles = {
            "font_family": css_font,
            "center_align_ratio": center_val,
            "accent_color": accent_color,
            "highlight_color": highlight_color,
            "has_italic": bool(c10.get("italic_usage", "")),
            "has_quote": bool(c10.get("quote_style", "")),
            "font_size_body": "11pt",
            "font_size_heading": "14pt",
        }

    return jsonify({
        "blog_id": blog_id,
        "preview_text": "\n".join(lines),
        "has_dna_analysis": dna_analysis is not None,
        "sample_title": sample_post.get("title", "") if sample_post else "",
        "dna_styles": dna_styles,
    })


@app.route('/api/blog/generate', methods=['POST'])
@login_required
def generate_blog():
    """스타일 템플릿 기반 블로그 글 생성 (통합 DNA 데이터 지원)"""
    style_template_id = request.form.get("style_template_id", "informational")
    keywords_str = request.form.get("keywords", "")
    keywords = [k.strip() for k in keywords_str.split(",") if k.strip()]
    blog_dna_id = request.form.get("blog_dna_id", "")
    target_audience = request.form.get("target_audience", "일반 독자")
    content_angle = request.form.get("content_angle", "정보전달형")
    direct_text = request.form.get("press_release", "")
    press_url = request.form.get("press_url", "").strip()
    reference_blog_url = request.form.get("reference_blog_url", "").strip()
    try:
        active_tags = json.loads(request.form.get("active_tags", "[]"))
    except Exception:
        active_tags = []

    # URL로 보도자료 크롤링
    if press_url:
        try:
            url_text = fetch_url_text(press_url)
            if url_text:
                direct_text = (direct_text + "\n\n" + url_text).strip() if direct_text else url_text
                print(f"[OK] 보도자료 URL 크롤링 완료: {len(url_text)}자")
        except Exception as e:
            print(f"[WARN] 보도자료 URL 크롤링 실패: {e}")

    # 참고 블로그 URL 크롤링 (스타일 참고용)
    reference_blog_text = ""
    if reference_blog_url:
        try:
            reference_blog_text = fetch_url_text(reference_blog_url)
            print(f"[OK] 참고 블로그 URL 크롤링 완료: {len(reference_blog_text)}자")
        except Exception as e:
            print(f"[WARN] 참고 블로그 URL 크롤링 실패: {e}")

    uploaded_files = [file for file in request.files.getlist("files") if file and file.filename]
    if not uploaded_files and 'file' in request.files and request.files['file'].filename:
        uploaded_files = [request.files['file']]

    # Gemini가 네이티브로 읽을 수 있는 형식
    GEMINI_NATIVE_EXTS = {'.pdf', '.jpg', '.jpeg', '.png', '.gif', '.webp'}

    temp_paths = []
    gemini_uploaded_files = []   # Gemini File API 업로드 객체
    text_sources = []            # 로컬 텍스트 추출 결과
    extracted_image_urls = []    # 업로드 자료에서 추출한 이미지 URL 목록

    api_key = os.getenv("GEMINI_API_KEY")
    _gemini_client = genai.Client(api_key=api_key) if api_key else None

    try:
        for file in uploaded_files:
            temp_path = save_uploaded_file(file)
            temp_paths.append(temp_path)
            ext = temp_path.suffix.lower()
            # 이미지 추출 (PDF/DOCX/이미지 파일 모두)
            extracted_image_urls.extend(extract_images_from_file(temp_path))

            if ext in GEMINI_NATIVE_EXTS and _gemini_client:
                # Gemini File API로 직접 업로드
                gfile = upload_to_gemini(temp_path, _gemini_client)
                if gfile:
                    gemini_uploaded_files.append(gfile)
                    print(f"[OK] Gemini 네이티브 처리: {file.filename}")
                else:
                    # 업로드 실패 시 텍스트 추출로 폴백
                    try:
                        text = extract_text_from_file(temp_path)
                        text_sources.append({
                            "name": file.filename,
                            "kind": ext.lstrip("."),
                            "text": text,
                        })
                    except Exception:
                        pass
            else:
                # DOCX / HWP / TXT → 로컬 텍스트 추출
                try:
                    text = extract_text_from_file(temp_path)
                    text_sources.append({
                        "name": file.filename,
                        "kind": ext.lstrip("."),
                        "text": text,
                    })
                except Exception as ex:
                    print(f"[WARN] 텍스트 추출 실패 ({file.filename}): {ex}")

        material_bundle = build_material_bundle(sources=text_sources, direct_text=direct_text)

    except Exception as e:
        return jsonify({"error": f"파일 처리 실패: {str(e)}"}), 500
    finally:
        for temp_path in temp_paths:
            try:
                if temp_path.exists():
                    temp_path.unlink()
            except Exception:
                pass

    if not uploaded_files:
        material_bundle = build_material_bundle(direct_text=direct_text)

    has_content = material_bundle.get("combined_text", "").strip() or bool(gemini_uploaded_files)
    if not has_content:
        return jsonify({"error": "보도자료 내용이 필요합니다."}), 400

    # 스타일 템플릿 로드
    style_template = _STYLE_TEMPLATES_MAP.get(style_template_id, STYLE_TEMPLATES[0])

    blog_dna_text = ""
    dna_analysis = None  # 항상 초기화 (UnboundLocalError 방지)
    if blog_dna_id:
        try:
            # ① DNA 분석 결과 (스타일 가이드) 로드
            if DNA_DIR.exists():
                # blog_dna_id가 전체 파일 ID("DNA_blogid_날짜")인 경우 직접 로드
                _direct = DNA_DIR / f"{blog_dna_id}.json"
                if _direct.exists():
                    with open(_direct, 'r', encoding='utf-8') as f:
                        dna_analysis = json.load(f)
                else:
                    # blog_id만 넘어온 경우 glob으로 가장 최근 파일 검색
                    candidates = sorted(
                        [f for f in DNA_DIR.glob(f"DNA_{blog_dna_id}_*.json")],
                        key=lambda f: f.stat().st_mtime, reverse=True
                    )
                    if candidates:
                        with open(candidates[0], 'r', encoding='utf-8') as f:
                            dna_analysis = json.load(f)

            # ② 원본 글 샘플 로드 (전문 1개 + 제목 목록)
            all_dna_posts = []
            if BLOG_COLLECTIONS_DIR.exists():
                for item in BLOG_COLLECTIONS_DIR.iterdir():
                    if item.is_dir() and not item.name.startswith('.'):
                        data_file = item / "_data.json"
                        if data_file.exists():
                            try:
                                with open(data_file, 'r', encoding='utf-8') as f:
                                    collection = json.load(f)
                                if collection.get("blog_id") == blog_dna_id:
                                    all_dna_posts.extend(collection.get("posts", []))
                            except Exception:
                                pass

            seen_urls = set()
            unique_posts = []
            for post in all_dna_posts:
                url = post.get("url")
                if url and url not in seen_urls:
                    seen_urls.add(url)
                    unique_posts.append(post)
            unique_posts.sort(key=lambda x: x.get('addDate', ''), reverse=True)

            # ③ DNA 스타일 가이드 + 샘플 글 조합
            dna_parts = []

            if dna_analysis:
                c1  = dna_analysis.get("c1_template_structure", {})
                c2  = dna_analysis.get("c2_tone_mood", {})
                # c3: 새 이름 우선, 구버전 하위호환
                c3  = dna_analysis.get("c3_speech_endings", dna_analysis.get("c3_speech_style", {}))
                # c4: 문장 구조
                c4  = dna_analysis.get("c4_sentence_structure", dna_analysis.get("c6_sentence_patterns", {}))
                # c5: 시그니처 표현
                c5  = dna_analysis.get("c6_signature_expressions", dna_analysis.get("c5_frequent_expressions", {}))
                # c6: 단락
                c6p = dna_analysis.get("c5_paragraph_composition", dna_analysis.get("c8_paragraph_composition", {}))
                # c8: 화법
                c8  = dna_analysis.get("c8_rhetoric", dna_analysis.get("c4_rhetoric", {}))
                # c9: 도입
                c9  = dna_analysis.get("c9_opening_patterns", dna_analysis.get("c9_opening_closing", {}))
                # c10: 마무리
                c10c = dna_analysis.get("c10_closing_patterns", dna_analysis.get("c9_opening_closing", {}))
                # c11: 시각
                c11v = dna_analysis.get("c11_visual_symbols", dna_analysis.get("c10_visual_formatting", {}))
                c12 = dna_analysis.get("c12_typography", {})
                c13 = dna_analysis.get("c13_brackets_quotes", {})
                c14 = dna_analysis.get("c14_length_stats", dna_analysis.get("c11_length_stats", {}))
                c15 = dna_analysis.get("c15_title_patterns", dna_analysis.get("c14_title_patterns", {}))
                c7  = dna_analysis.get("c7_vocabulary", {})

                dna_parts.append("【블로그 글쓰기 DNA 스타일 가이드 — 모든 항목 100% 재현할 것】")
                # 구조 & 톤
                dna_parts.append(f"구조 패턴: {c1.get('overall_pattern', '')} / 섹션 수: {c1.get('section_count', '')}")
                dna_parts.append(f"섹션 흐름: {' → '.join(c1.get('section_flow', []))}")
                dna_parts.append(f"소제목 포맷: {c1.get('subheading_format', c1.get('heading_style', ''))}")
                dna_parts.append(f"톤: {c2.get('primary_tone', '')} / 격식도: {c2.get('formality_level', '')}/10 / 활발함: {c2.get('energy_level', '')}/10")
                dna_parts.append(f"톤 변화 패턴: {c2.get('tone_shift_pattern', '')}")
                # 어투/종결어미
                primary_endings = c3.get('primary_endings', c3.get('ending_patterns', []))
                dna_parts.append(f"종결어미 TOP5: {', '.join(primary_endings[:5])}")
                dna_parts.append(f"격식:비격식 비율: {c3.get('formality_mix', '')} / 연속 같은 어미: {c3.get('consecutive_same_ending', '')}")
                dna_parts.append(f"독자 호칭: {c3.get('reader_address', '')} / 의문형 패턴: {c3.get('question_ending_style', '')}")
                # 문장 구조
                dna_parts.append(f"문장 길이: 평균 {c4.get('avg_chars_per_sentence', c4.get('avg_length', ''))}자 / 짧은:{c4.get('short_sentence_ratio', '')} 중간:{c4.get('medium_sentence_ratio', '')} 긴:{c4.get('long_sentence_ratio', '')}")
                dna_parts.append(f"리듬 패턴: {c4.get('rhythm_pattern', c4.get('rhythm', ''))}")
                dna_parts.append(f"문장 시작 패턴: {', '.join(c4.get('leading_phrase_patterns', [])[:5])}")
                # 단락
                dna_parts.append(f"단락당 문장: {c6p.get('avg_sentences_per_paragraph', '')} / 여백: {c6p.get('whitespace_style', c6p.get('whitespace_usage', ''))}")
                dna_parts.append(f"단락 시작 패턴: {c6p.get('paragraph_opening_pattern', '')}")
                # 시그니처 표현
                sig = c5.get('signature_phrases', [])
                trans = c5.get('transition_words', [])
                dna_parts.append(f"시그니처 표현: {', '.join(sig[:7])}")
                dna_parts.append(f"전환 표현: {', '.join(trans[:6])}")
                dna_parts.append(f"강조 표현: {', '.join(c5.get('emphasis_expressions', [])[:5])}")
                dna_parts.append(f"긍정 추임새: {', '.join(c5.get('affirmation_expressions', [])[:5])}")
                # 어휘
                dna_parts.append(f"어휘 수준: {c7.get('level', '')} / 순우리말:{c7.get('korean_ratio', '')} 한자어:{c7.get('sino_korean_ratio', '')} 외래어:{c7.get('foreign_word_ratio', '')}")
                dna_parts.append(f"특징 어휘: {', '.join(c7.get('characteristic_words', [])[:8])}")
                # 도입/마무리
                dna_parts.append(f"도입 방식: {', '.join(c9.get('opening_types', []))}")
                dna_parts.append(f"첫 문장 패턴: {c9.get('first_sentence_pattern', '')}")
                dna_parts.append(f"마무리 방식: {', '.join(c10c.get('closing_types', []))}")
                dna_parts.append(f"CTA 방식: {c10c.get('cta_style', '')} / 표현: {', '.join(c10c.get('cta_keywords', [])[:4])}")
                # 시각 요소
                emoji_list = c11v.get('emoji_list', c11v.get('emoji_types', []))
                dna_parts.append(f"이모지 빈도(1-10): {c11v.get('emoji_frequency', c11v.get('emoji_usage', ''))} / 글당 이모지 수: {c11v.get('emoji_per_post', '')}")
                dna_parts.append(f"자주 쓰는 이모지: {', '.join(emoji_list[:15])}")
                dna_parts.append(f"이모지 위치: {c11v.get('emoji_position', '')}")
                dna_parts.append(f"특수기호: {', '.join(c11v.get('special_symbols', [])[:15])}")
                sep = c11v.get('separator_patterns', [c11v.get('separator_style', '')])
                dna_parts.append(f"구분선: {', '.join(sep[:3]) if isinstance(sep, list) else sep}")
                # 글자수/분량 (c14)
                if c14:
                    dna_parts.append(f"\n【분량 기준】")
                    dna_parts.append(f"글당 평균 글자수: {c14.get('avg_chars_per_post', '')}")
                    dna_parts.append(f"평균 문장 수: {c14.get('avg_sentences_per_post', '')}")
                    dna_parts.append(f"문장당 평균 글자수: {c14.get('avg_chars_per_sentence', '')}")
                    dna_parts.append(f"서론:본론:결론 비율: {c14.get('content_ratio', '')}")
                    if c14.get('writing_density_guide'):
                        dna_parts.append(f"분량 지침: {c14.get('writing_density_guide', '')}")
                # 폰트/글꼴 (c12)
                if c12:
                    dna_parts.append(f"\n【폰트/글꼴 스타일】")
                    if c12.get('font_families'):
                        dna_parts.append(f"사용 폰트: {', '.join(c12.get('font_families', []))}")
                    dna_parts.append(f"본문 크기: {c12.get('base_font_size', '')} / 소제목 크기: {c12.get('heading_font_size', '')}")
                    dna_parts.append(f"볼드 빈도(1-10): {c12.get('bold_frequency', '')} / 목적: {c12.get('bold_purpose', '')}")
                    if c12.get('bold_examples'):
                        dna_parts.append(f"볼드 예시: {', '.join(c12.get('bold_examples', [])[:3])}")
                    dna_parts.append(f"기울임꼴: {c12.get('italic_usage', '')} / 밑줄: {c12.get('underline_usage', '')}")
                    if c12.get('font_guide'):
                        dna_parts.append(f"글꼴 지침: {c12.get('font_guide', '')}")
                # 꺽쇠/괄호 (c13)
                if c13:
                    dna_parts.append(f"\n【꺽쇠/괄호/인용부호】")
                    if c13.get('angle_bracket_types'):
                        dna_parts.append(f"꺽쇠 종류: {', '.join(c13.get('angle_bracket_types', []))} / 목적: {c13.get('angle_bracket_purpose', '')} / 빈도: {c13.get('angle_bracket_frequency', '')}")
                    if c13.get('square_bracket_types'):
                        dna_parts.append(f"대괄호 종류: {', '.join(c13.get('square_bracket_types', []))} / 목적: {c13.get('square_bracket_purpose', '')}")
                    dna_parts.append(f"소괄호 패턴: {c13.get('round_bracket_usage', '')}")
                    dna_parts.append(f"따옴표 방식: {c13.get('quotation_mark_style', '')}")
                    if c13.get('examples'):
                        dna_parts.append(f"꺽쇠/괄호 예시: {' | '.join(c13.get('examples', [])[:3])}")
                # 제목 패턴 (c14)
                if c14:
                    dna_parts.append(f"\n【제목 패턴】")
                    dna_parts.append(f"평균 제목 길이: {c14.get('avg_title_length', '')} / 구조: {c14.get('title_structure', '')}")
                    dna_parts.append(f"숫자 활용: {c14.get('number_usage', '')} / 감정 후크: {c14.get('emotion_hook', '')}")
                    if c14.get('examples'):
                        dna_parts.append(f"제목 예시: {' | '.join(c14.get('examples', [])[:3])}")
                # 이미지 (c16 또는 c15 하위호환)
                c16 = dna_analysis.get("c16_image_media", dna_analysis.get("c15_image_media", {}))
                if c16:
                    dna_parts.append(f"\n【이미지/미디어 패턴】")
                    dna_parts.append(f"글당 이미지 수: {c16.get('avg_images_per_post', c16.get('avg_images', ''))} / 배치: {c16.get('image_placement', c16.get('image_position', ''))}")
                    dna_parts.append(f"캡션 방식: {c16.get('caption_style', c16.get('image_caption_style', ''))} / 밀도: {c16.get('media_density', '')}")

                # c17 — 문장부호
                c17 = dna_analysis.get("c17_punctuation", {})
                if c17:
                    dna_parts.append(f"\n【문장부호/구두점 패턴 — 반드시 재현할 것】")
                    dna_parts.append(f"마침표 방식: {c17.get('period_style', '')} / 생략 비율: {c17.get('period_omission_ratio', '')}")
                    dna_parts.append(f"쉼표 빈도(1-10): {c17.get('comma_frequency', '')} / 방식: {c17.get('comma_style', '')}")
                    dna_parts.append(f"느낌표 빈도(1-10): {c17.get('exclamation_frequency', '')} / 방식: {c17.get('exclamation_style', '')}")
                    dna_parts.append(f"물음표 맥락: {c17.get('question_mark_usage', '')}")
                    dna_parts.append(f"말줄임표: {c17.get('ellipsis_usage', '')} / 물결표: {c17.get('tilde_usage', '')} / 대시: {c17.get('dash_usage', '')}")
                    if c17.get('multiple_punct_usage'):
                        dna_parts.append(f"복수 부호(!! ~~): {c17.get('multiple_punct_usage', '')}")
                    if c17.get('examples'):
                        dna_parts.append(f"구두점 예시:\n" + "\n".join(c17['examples'][:2]))

                # c18 — 숫자/데이터
                c18 = dna_analysis.get("c18_numbers_data", {})
                if c18:
                    dna_parts.append(f"\n【숫자/단위/데이터 표현 — 반드시 재현할 것】")
                    dna_parts.append(f"숫자 선호: {c18.get('numeral_preference', '')}")
                    dna_parts.append(f"가격 형식: {c18.get('price_format', '')} / 날짜 형식: {c18.get('date_format', '')}")
                    dna_parts.append(f"단위 스타일: {c18.get('unit_style', '')} / 순위 형식: {c18.get('ranking_format', '')}")
                    dna_parts.append(f"어림수 표현: {c18.get('approximation_style', '')}")
                    if c18.get('examples'):
                        dna_parts.append(f"수치 예시: {' | '.join(c18['examples'][:3])}")

                # c19 — 독자 상호작용
                c19 = dna_analysis.get("c19_reader_engagement", {})
                if c19:
                    dna_parts.append(f"\n【독자 참여 유도 방식 — 반드시 재현할 것】")
                    dna_parts.append(f"독자 질문 빈도(1-10): {c19.get('direct_question_frequency', '')}")
                    if c19.get('empathy_phrases'):
                        dna_parts.append(f"공감 표현: {', '.join(c19['empathy_phrases'][:4])}")
                    if c19.get('inclusive_expressions'):
                        dna_parts.append(f"포용 표현: {', '.join(c19['inclusive_expressions'][:4])}")
                    dna_parts.append(f"추천 강도: {c19.get('recommendation_strength', '')}")
                    if c19.get('recommendation_expressions'):
                        dna_parts.append(f"추천 표현: {', '.join(c19['recommendation_expressions'][:4])}")
                    if c19.get('urgency_patterns'):
                        dna_parts.append(f"긴박감 표현: {', '.join(c19['urgency_patterns'][:3])}")
                    if c19.get('examples'):
                        dna_parts.append(f"참여 유도 예시:\n" + "\n".join(c19['examples'][:2]))

                # c20 — 감탄사/추임새
                c20 = dna_analysis.get("c20_interjections_fillers", {})
                if c20:
                    dna_parts.append(f"\n【감탄사/추임새/습관어 — 반드시 재현할 것】")
                    if c20.get('interjections'):
                        dna_parts.append(f"감탄사: {', '.join(c20['interjections'][:8])}")
                    if c20.get('filler_starters'):
                        dna_parts.append(f"시작 습관어: {', '.join(c20['filler_starters'][:6])}")
                    if c20.get('affirmations'):
                        dna_parts.append(f"긍정 추임새: {', '.join(c20['affirmations'][:5])}")
                    if c20.get('excitement_expressions'):
                        dna_parts.append(f"흥분 표현: {', '.join(c20['excitement_expressions'][:5])}")
                    dna_parts.append(f"사용 빈도(1-10): {c20.get('frequency', '')} / 위치: {c20.get('position_pattern', '')}")
                    if c20.get('examples'):
                        dna_parts.append(f"추임새 예시:\n" + "\n".join(c20['examples'][:2]))

                # c21 — 인라인 서식 상세 (확장 버전)
                c21 = dna_analysis.get("c21_inline_formatting", {})
                if c21:
                    dna_parts.append(f"\n【인라인 서식 상세 (Naver SE3) — 반드시 재현할 것】")
                    dna_parts.append(f"폰트 전환 빈도: {c21.get('font_switch_frequency', '')} / 트리거: {c21.get('font_switch_trigger', '')}")
                    if c21.get('font_families_used'):
                        dna_parts.append(f"사용 폰트: {', '.join(c21['font_families_used'])}")
                    if c21.get('size_switch_pattern'):
                        dna_parts.append(f"크기 전환 패턴: {c21.get('size_switch_pattern', '')}")
                    if c21.get('size_examples_by_level'):
                        _sz = c21['size_examples_by_level']
                        if isinstance(_sz, dict):
                            dna_parts.append(f"크기별 사용: " + " | ".join(f"{k}={v}" for k, v in _sz.items() if v and v != "어떤 텍스트에 사용?"))
                    if c21.get('color_switch_pattern'):
                        dna_parts.append(f"색상 전환 패턴: {c21.get('color_switch_pattern', '')}")
                    if c21.get('color_switch_examples'):
                        dna_parts.append(f"색상 전환 예시: {' | '.join(c21['color_switch_examples'][:4])}")
                    if c21.get('italic_usage') and c21['italic_usage'] != '없음':
                        dna_parts.append(f"기울임체: {c21.get('italic_usage', '')}")
                    if c21.get('underline_usage') and c21['underline_usage'] != '없음':
                        dna_parts.append(f"밑줄: {c21.get('underline_usage', '')}")
                    if c21.get('strikethrough_usage') and c21['strikethrough_usage'] != '없음':
                        dna_parts.append(f"취소선: {c21.get('strikethrough_usage', '')}")
                    if c21.get('background_color_pattern') and c21['background_color_pattern'] != '없음':
                        dna_parts.append(f"배경색 강조: {c21.get('background_color_pattern', '')}")
                    if c21.get('background_color_examples'):
                        dna_parts.append(f"배경색 예시: {' | '.join(c21['background_color_examples'][:3])}")
                    if c21.get('center_align_pattern'):
                        dna_parts.append(f"가운데 정렬: {c21.get('center_align_pattern', '')}")
                    if c21.get('box_quote_pattern') and c21['box_quote_pattern'] != '없음':
                        dna_parts.append(f"박스/인용구: {c21.get('box_quote_pattern', '')}")
                    if c21.get('combined_format_examples'):
                        dna_parts.append(f"복합 서식 예시: {' | '.join(c21['combined_format_examples'][:3])}")
                    if c21.get('naver_se3_pattern_guide'):
                        dna_parts.append(f"SE3 서식 가이드: {c21.get('naver_se3_pattern_guide', '')}")

                # c22 — 콘텐츠 패턴
                c22 = dna_analysis.get("c22_content_patterns", {})
                if c22:
                    dna_parts.append(f"\n【콘텐츠 주제/정보 패턴 — 반드시 따를 것】")
                    if c22.get('main_topics'):
                        dna_parts.append(f"주요 주제: {', '.join(c22['main_topics'][:5])}")
                    if c22.get('content_angle'):
                        dna_parts.append(f"콘텐츠 각도: {c22.get('content_angle', '')}")
                    if c22.get('must_include_elements'):
                        dna_parts.append(f"반드시 포함: {', '.join(c22['must_include_elements'][:6])}")
                    if c22.get('never_include_elements'):
                        dna_parts.append(f"절대 포함 금지: {', '.join(c22['never_include_elements'][:4])}")
                    if c22.get('info_ordering'):
                        dna_parts.append(f"정보 제시 순서: {c22.get('info_ordering', '')}")
                    if c22.get('local_terminology'):
                        dna_parts.append(f"지역/기관 용어: {', '.join(c22['local_terminology'][:6])}")
                    if c22.get('typical_post_template'):
                        dna_parts.append(f"전형적 글 구성: {c22.get('typical_post_template', '')}")

                # 도입/마무리 실제 예시
                c9_open = dna_analysis.get("c9_opening_patterns", dna_analysis.get("c9_opening_closing", {}))
                c10_close = dna_analysis.get("c10_closing_patterns", dna_analysis.get("c9_opening_closing", {}))
                if c9_open.get("opening_examples"):
                    dna_parts.append(f"\n실제 도입부 예시 (이대로 따라 쓸 것):\n" + "\n".join(c9_open["opening_examples"][:3]))
                if c10_close.get("closing_examples"):
                    dna_parts.append(f"\n실제 마무리 예시 (이대로 따라 쓸 것):\n" + "\n".join(c10_close["closing_examples"][:3]))

            # 원본 글 전문 최대 3개 (스타일 레퍼런스)
            if unique_posts:
                dna_parts.append(f"\n【실제 글 샘플 (이 스타일을 최대한 그대로 따라 쓸 것 — 어투·구조·길이·표현 모두)】")
                for si, sample in enumerate(unique_posts[:3], 1):
                    dna_parts.append(f"\n[샘플 {si}] 제목: {sample.get('title', '')}")
                    dna_parts.append(sample.get('content', '')[:1800])

                # 제목 목록 (참고용)
                title_list = [f"- {p.get('title', '')}" for p in unique_posts[3:12]]
                if title_list:
                    dna_parts.append(f"\n【최근 글 제목 목록 (주제 참고용)】\n" + "\n".join(title_list))

            # 활성 페르소나 태그 추가
            if active_tags:
                tag_labels = [t.get("label", "") for t in active_tags if t.get("label")]
                if tag_labels:
                    dna_parts.append(f"\n[활성화된 스타일 태그 — 이 특징들을 반드시 반영]\n" + "\n".join(f"- {l}" for l in tag_labels))

            blog_dna_text = "\n".join(dna_parts)

        except Exception as e:
            print(f"[WARN] 블로그 DNA 로드 실패: {e}")

    # ── DNA에서 블로그 신원·분량·구조 가이드 추출 ──────────────
    dna_blog_id = (dna_analysis or {}).get('blog_id', blog_dna_id or '')

    # 분량 가이드: c14_length_stats 또는 c11_length_stats
    _c14 = (dna_analysis or {}).get("c14_length_stats",
             (dna_analysis or {}).get("c11_length_stats", {}))
    if _c14:
        _avg    = _c14.get('avg_chars_per_post', '')
        _guide  = _c14.get('density_guide', _c14.get('writing_density_guide', ''))
        _ratio  = _c14.get('content_ratio', '')
        length_guide_text = "▸ 이 블로그의 평균 글 길이: " + str(_avg)
        if _guide: length_guide_text += f"\n▸ 분량 지침: {_guide}"
        if _ratio: length_guide_text += f"\n▸ 서론:본론:결론 = {_ratio}"
    else:
        length_guide_text = "▸ 최소 1,500자 이상, 권장 2,000~2,500자\n▸ 서론(10%) → 본론(75%) → 결론(15%)"

    # 구조 가이드: c1_template_structure
    _c1 = (dna_analysis or {}).get("c1_template_structure", {})
    if _c1:
        _pattern  = _c1.get('overall_pattern', '')
        _flow     = _c1.get('section_flow', [])
        _subfmt   = _c1.get('subheading_format', _c1.get('heading_style', ''))
        struct_guide_text = f"▸ 구조 패턴: {_pattern}"
        if _flow: struct_guide_text += "\n▸ 섹션 흐름: " + " → ".join(_flow)
        if _subfmt: struct_guide_text += f"\n▸ 소제목 포맷: {_subfmt}"
    else:
        struct_guide_text = "▸ 서론 → 본론(3~5단락) → 결론 구조"

    # 꺽쇠/인용/폰트크기 마커 가이드 (c13 기반)
    _c13 = (dna_analysis or {}).get("c13_brackets_quotes", {})
    _bracket_guide = ""
    if _c13:
        _ab  = _c13.get("angle_bracket_examples", [])
        _sq  = _c13.get("square_bracket_purpose", "")
        _qt  = _c13.get("quotation_style", "")
        _cbo = _c13.get("bracket_combo_pattern", "")
        _abf = _c13.get("angle_bracket_frequency", "")
        if _ab or _qt:
            _bracket_guide = "\n▸ 이 블로그의 꺽쇠/인용 패턴 (반드시 재현):"
            if _ab:   _bracket_guide += f"\n  - 꺽쇠 예시 ({_abf}): {' / '.join(_ab[:3])}"
            if _sq:   _bracket_guide += f"\n  - 대괄호 용도: {_sq[:80]}"
            if _qt:   _bracket_guide += f"\n  - 인용부호: {_qt}"
            if _cbo:  _bracket_guide += f"\n  - 조합 패턴: {_cbo}"

    # 폰트 크기 마커 가이드
    _fs_levels = (dna_analysis or {}).get("c12_typography", {}).get("font_size_levels", [])
    _size_guide = ""
    if len(_fs_levels) >= 2:
        _size_guide = "\n▸ 인라인 크기 마커 (필요시 사용):\n  - [작게]텍스트[/작게] → 작은 글씨 (부가설명, 주석)\n  - [크게]텍스트[/크게] → 강조 큰 글씨"

    struct_guide_text += _bracket_guide + _size_guide

    template_name = style_template.get("name", "정보전달형")
    formality_score = style_template.get("formality_score", 5)
    custom_prompt = style_template.get("custom_prompt", "")

    # 보정 기록 로드 (같은 style_template_id + blog_dna_id 기준, 최근 3개)
    calibration_prompt = ""
    try:
        cal_files = sorted(
            CALIBRATIONS_DIR.glob("CAL_*.json"),
            key=lambda f: f.stat().st_mtime,
            reverse=True
        )
        matching = []
        for cf in cal_files:
            try:
                d = json.loads(cf.read_text(encoding='utf-8'))
                if d.get("style_template_id") == style_template_id:
                    matching.append(d)
                    if len(matching) >= 3:
                        break
            except Exception:
                pass
        if matching:
            tips = []
            for d in matching:
                tips.append(d.get("calibration_prompt", "").strip())
            combined = "\n".join(t for t in tips if t)
            if combined:
                calibration_prompt = f"[실제 통과된 글 기반 보정 지침 — 반드시 따를 것]\n{combined}"
            print(f"[OK] 보정 기록 {len(matching)}개 적용")
    except Exception as e:
        print(f"[WARN] 보정 기록 로드 실패: {e}")

    api_key = os.getenv("GEMINI_API_KEY")
    generation_mode = "offline"
    versions = []

    if api_key:
        try:
            client = genai.Client(api_key=api_key)

            # Gemini 네이티브 파일 처리 대기 (PROCESSING → ACTIVE)
            active_gfiles = []
            for gf in gemini_uploaded_files:
                for _ in range(30):
                    if gf.state.name == 'ACTIVE':
                        active_gfiles.append(gf)
                        break
                    _time.sleep(1)
                    gf = client.files.get(name=gf.name)
                else:
                    print(f"[WARN] Gemini 파일 처리 타임아웃: {gf.name}")

            # 파일 첨부 안내 문구
            native_note = ""
            if active_gfiles:
                native_note = f"첨부된 파일 {len(active_gfiles)}개를 직접 읽고 핵심 내용을 파악하세요."

            # ── 블로그 신원 헤더 ──────────────────────────────────
            identity_header = ""
            if dna_blog_id:
                identity_header = f"""당신은 '{dna_blog_id}' 네이버 블로그의 전담 콘텐츠 작성자입니다.
아래 [글쓰기 DNA]는 이 블로그가 지금까지 써온 글을 정밀 분석한 결과입니다.
당신은 이 DNA를 완벽히 내재화했습니다. 이 블로그처럼 생각하고, 이 블로그의 언어로 씁니다.
DNA에서 벗어나는 순간 실패입니다. DNA를 따르는 것이 곧 이 블로그답게 쓰는 것입니다."""
            else:
                identity_header = f"""당신은 실전형 블로그 콘텐츠 에디터입니다.
아래 스타일 가이드를 바탕으로 블로그 글을 작성하세요.
[글쓰기 스타일: {template_name} / 격식도 {formality_score}/10]
{custom_prompt}"""

            # ── DNA 없는 경우 기본 안내 ───────────────────────────
            dna_section = ""
            if blog_dna_text:
                dna_section = f"""
{'━'*52}
[{dna_blog_id} 블로그 글쓰기 DNA — 이것이 당신의 글쓰기 정체성]
{'━'*52}
{blog_dna_text}
{'━'*52}"""
            else:
                dna_section = f"""
[글쓰기 스타일]
- 템플릿: {template_name} / 격식도: {formality_score}/10
- 종결어미: {', '.join(style_template.get('ending_patterns', []))}
{custom_prompt}"""

            # ── 참고 블로그 URL 샘플 ─────────────────────────────
            ref_section = ""
            if reference_blog_text:
                ref_section = f"""
[추가 참고 블로그 스타일 (URL에서 수집)]
{reference_blog_text[:2500]}"""

            # ── 보정 기록 ────────────────────────────────────────
            cal_section = f"\n{calibration_prompt}\n" if calibration_prompt else ""

            prompt = f"""{identity_header}
반드시 JSON만 출력하세요.
{dna_section}
{ref_section}
{cal_section}
{'━'*52}
[이번 작업: 아래 자료를 {dna_blog_id or '이'} 블로그 스타일로 변환]
{'━'*52}

[첨부 파일]
{native_note or '없음'}

[텍스트 자료 / 보도자료]
{material_bundle.get('briefing', '')[:8000] or '없음'}

[타겟 독자]
{target_audience}

[콘텐츠 앵글]
{content_angle}

[핵심 키워드]
{", ".join(keywords) if keywords else "없음"}

{'━'*52}
[분량 기준 — DNA 기반, 반드시 준수]
{'━'*52}
{length_guide_text}

{'━'*52}
[구조 기준 — DNA 기반, 반드시 준수]
{'━'*52}
{struct_guide_text}

{'━'*52}
[작성 지침 — DNA 100% 일치 단일 버전]
{'━'*52}
★ 이 블로그의 DNA를 완벽히 재현한 단 1개의 최고 품질 버전을 작성한다.
★ DNA 스타일(톤·어투·이모지·구조·길이·문장부호)을 100% 따를 것.
★ 볼드 텍스트가 필요한 경우 반드시 **텍스트** 형식으로 표시할 것 (DNA에 볼드가 있으면 사용, 없으면 금지).
★ 소제목이 필요한 경우 반드시 ## 소제목 형식으로 표시할 것 (DNA에 소제목 구조가 있으면 사용).
★ 인라인 서식 마커 — DNA c21 패턴 기반, 해당하는 경우만 사용:

  [크기 변환]
  - [작게]부연설명/주석[/작게]           → 작은 글씨 (fs13, 보조 설명)
  - [크게]핵심 수치나 강조[/크게]         → 큰 글씨 (fs19, 핵심 강조)
  - [매우크게]제목급 텍스트[/매우크게]    → 최대 크기 (fs28)
  - [아주작게]법적 고지 등[/아주작게]     → 극소 (fs11)
  - [fs13]텍스트[/fs13]                  → 정확한 px 지정 (13/16/19/24/28 등)

  [스타일 변환]
  - **볼드텍스트**                         → 굵게 (DNA bold 패턴이 있으면 사용)
  - [진하게]중요 포인트 텍스트[/진하게]   → 굵게 (마커 형식, 핵심 정보·포인트에 사용)
  - *기울임텍스트*                         → 기울임 (DNA에 italic 쓰면 사용)
  - [기울임]텍스트[/기울임]               → 기울임 (마커 형식)
  - [밑줄]텍스트[/밑줄]                   → 밑줄 (DNA에 underline 패턴 있으면 사용)
  - [취소선]텍스트[/취소선]               → 취소선 (~~text~~ 형식도 가능)
  - ~~취소선~~                             → 취소선 (마크다운 형식)

  [색상/배경 변환]
  - [색상:#0078cb]텍스트[/색상]           → 특정 hex 색상 (DNA color_switch_pattern 기반)
  - [배경:#fff9a0]텍스트[/배경]           → 배경 강조색 (형광펜 효과)
  - [형광펜]텍스트[/형광펜]               → DNA highlight_color 형광펜
  - ==형광펜텍스트==                       → 형광펜 (마크다운 형식)
  - {{액센트색텍스트}}                     → DNA accent 색상

  [폰트 변환]
  - [폰트:강조]텍스트[/폰트]              → DNA 두 번째 폰트 (c21 font_switch_trigger 기반)
  - [폰트:나눔고딕]텍스트[/폰트:나눔고딕] → NanumGothic
  - [폰트:나눔명조]텍스트[/폰트:나눔명조] → NanumMyeongjo (명조 계열)
  - [폰트:나눔스퀘어]텍스트[/폰트:나눔스퀘어] → NanumSquare
  - [폰트:나눔바른히피]텍스트[/폰트:나눔바른히피] → 손글씨 느낌 폰트
  - [폰트:맑은고딕]텍스트[/폰트:맑은고딕] → Malgun Gothic

  [복합 서식]
  - [강조]핵심키워드[/강조]               → 굵게+액센트색 (가장 강한 강조)
  - [인라인강조]텍스트[/인라인강조]       → inline accent + 굵게
  - [볼드밑줄]텍스트[/볼드밑줄]           → 굵게+밑줄
  - [크게색상:#0078cb]텍스트[/크게색상]   → 크게+색상 복합
  - [기울임색상:#555]텍스트[/기울임색상]  → 기울임+색상 복합

  [특수 기호 — DNA c13 기반]
  - < 장소명 >  《제목》  【소제목】  「인용」  ≪강조≫
  - "인용구"  '작은인용'  (유니코드 따옴표 사용)
  - [첨자위:주1]  [첨자아래:참고]

  [추가 6가지 마커]
  - [링크:https://...]텍스트[/링크]        → 하이퍼링크 (색상+밑줄)
  - [뱃지]라벨[/뱃지]                      → 인라인 필 배지 (흰글자+액센트배경)
  - [밑줄색상:#e74c3c]텍스트[/밑줄색상]   → 색상 밑줄 (border-bottom)
  - [자간:2]텍스트[/자간]                  → 자간 넓히기 (px 단위)
  - [회색]보조설명[/회색]                  → 흐린 회색 텍스트
  - [흰글자]텍스트[/흰글자]               → 흰색 글자 ([배경:#hex]와 조합)

  [블록 정렬 — 줄 전체를 마커로 감쌀 것]
  - [중앙]텍스트[/중앙]                   → 가운데 정렬 단락
  - [우측]텍스트[/우측]                   → 우측 정렬 단락
  - [들여쓰기]텍스트[/들여쓰기]           → 들여쓰기 단락
  - [박스]인용구 내용[/박스]              → 인용구/박스 스타일
  - [박스배경:#f0f8ff]텍스트[/박스배경]   → 배경색 있는 박스
  - [구분선]                               → 구분선 (단독 줄로)

  ⚠️ 규칙: DNA c21에 해당 서식 패턴이 있을 때만 사용. 없는 서식을 임의로 추가 금지.
★ 꺽쇠·인용부호는 DNA c13 패턴 그대로 재현할 것 (예시: < 장소명 >, 《제목》, ''인용'', "인용")

{'━'*52}
[절대 준수 규칙]
{'━'*52}
- 작성자 실명·개인정보 노출 금지
- ``` (백틱 코드블록) 사용 금지
- --- === (마크다운 구분선) 사용 금지
- DNA에 꺽쇠(《》 【】 「」 등)가 있으면 적극 사용, 없으면 사용 금지
- DNA에 이모지/특수기호가 있으면 동일하게 사용, 없으면 사용 금지
- 핵심 일정·대상·혜택·문의처 누락 금지
- 단락 사이 빈 줄은 DNA의 여백 패턴 그대로 재현할 것
- HTML 태그 절대 금지: <p>, <br>, <div>, <span>, <h2> 등 어떤 HTML 태그도 content 안에 쓰지 말 것
- content는 순수 텍스트 + 마크다운(** ##)만 사용할 것. HTML을 출력하면 즉시 실패로 간주됨

[출력 JSON]
{{
  "versions": [
    {{"version_type":"dna_match","version_label":"DNA 일치형","title":"제목","content":"본문","tags":["태그"],"meta_description":"설명"}}
  ]
}}"""

            # 멀티모달 contents: 텍스트 프롬프트 + 네이티브 파일 객체들
            from google.genai import types as _gtypes
            contents = [_gtypes.Part.from_text(text=prompt)]
            for gf in active_gfiles:
                contents.append(_gtypes.Part.from_uri(
                    file_uri=gf.uri,
                    mime_type=gf.mime_type,
                ))

            from google.genai import types as _cfg_types
            response = client.models.generate_content(
                model='gemini-2.5-pro',
                contents=contents,
                config=_cfg_types.GenerateContentConfig(
                    max_output_tokens=8192,
                    temperature=0.8,
                ),
            )

            # Gemini 업로드 파일 사용 후 삭제 (48시간 자동 만료지만 즉시 정리)
            for gf in active_gfiles:
                try:
                    client.files.delete(name=gf.name)
                except Exception:
                    pass
            blog_result = parse_ai_json(response.text)
            versions = blog_result.get("versions", [])
            for v in versions:
                if "title" in v:
                    v["title"] = _strip_markdown(v["title"])
                # AI가 HTML 태그를 출력한 경우 제거 (순수 텍스트+마크다운만 유지)
                if "content" in v:
                    import re as _re_html
                    v["content"] = _re_html.sub(r'<[^>]+>', '', v["content"])
                    v["content"] = v["content"].replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&quot;', '"')
            if versions:
                generation_mode = "ai"
        except Exception as e:
            print(f"[WARN] AI 블로그 생성 실패, 오프라인 모드로 전환: {e}")

    if not versions:
        versions = generate_blog_versions_offline(
            persona_data={},
            material_bundle=material_bundle,
            keywords=keywords,
            target_audience=target_audience,
            content_angle=content_angle,
        )

    output_id = f"BLOG_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    package = build_blog_package(
        output_id=output_id,
        client_id=style_template_id,
        client_name=template_name,
        versions=versions,
        source_bundle=material_bundle,
        extra={"generation_mode": generation_mode, "style_template_id": style_template_id},
    )
    save_blog_package(package, OUTPUT_DIR)

    # ── DNA → HTML 스타일 완전 추출 ────────────────────────────
    try:
      _dna_style_ok = True
      _c12 = (dna_analysis or {}).get("c12_typography", {})
      _c11 = (dna_analysis or {}).get("c11_visual_symbols",
              (dna_analysis or {}).get("c10_visual_formatting", {}))
    except Exception:
      _dna_style_ok = False
      _c12, _c11 = {}, {}

    # 1) 폰트 패밀리: Naver SE 클래스명 → CSS font-family
    _font_map = {
        # Naver SE3 클래스명
        "system":             "'Malgun Gothic', '맑은 고딕', 'Apple SD Gothic Neo', sans-serif",
        "nanumgothic":        "'NanumGothic', '나눔고딕', 'Malgun Gothic', sans-serif",
        "nanumsquare":        "'NanumSquare', '나눔스퀘어', 'NanumGothic', sans-serif",
        "nanumbareunhipi":    "'NanumBarunhiPi', '나눔바른히피', 'NanumGothic', sans-serif",
        "nanumbarungothic":   "'NanumBarunGothic', 'NanumGothic', sans-serif",
        "nanumsquareround":   "'NanumSquareRound', 'NanumSquare', sans-serif",
        "nanummyeongjo":      "'NanumMyeongjo', '나눔명조', serif",
        "dotum":              "Dotum, '돋움', sans-serif",
        "gulim":              "Gulim, '굴림', sans-serif",
        "malgun":             "'Malgun Gothic', '맑은 고딕', sans-serif",
        # 한글 표기
        "나눔고딕":           "'NanumGothic', '나눔고딕', 'Malgun Gothic', sans-serif",
        "나눔명조":           "'NanumMyeongjo', '나눔명조', serif",
        "나눔스퀘어":         "'NanumSquare', 'NanumGothic', sans-serif",
        "나눔바른히피":       "'NanumBarunhiPi', 'NanumGothic', sans-serif",
        "맑은 고딕":          "'Malgun Gothic', '맑은 고딕', sans-serif",
        "맑은고딕":           "'Malgun Gothic', '맑은 고딕', sans-serif",
        "기본 폰트":          "'NanumGothic', 'Malgun Gothic', '맑은 고딕', sans-serif",
        "기본폰트":           "'NanumGothic', 'Malgun Gothic', '맑은 고딕', sans-serif",
    }
    _font_names = _c12.get("font_families", [])
    _raw_font = _font_names[0] if _font_names else "기본 폰트"
    html_font = _font_map.get(_raw_font.lower().replace(" ", ""),
                _font_map.get(_raw_font,
                f"'{_raw_font}', 'NanumGothic', 'Malgun Gothic', sans-serif"))

    # 2) 폰트 크기: Naver SE 클래스명(fs9~fs30) 및 텍스트 → px
    def _to_px(val, default):
        if not val: return default
        val = str(val).strip()
        if val.endswith("px"): return val
        # fsNN 형식 (Naver SE3)
        import re as _r
        m = _r.match(r'fs(\d+)', val.lower().split()[0])
        if m: return f"{m.group(1)}px"
        _tmap = {"기본": "15px", "작게": "13px", "보통": "15px",
                 "크게": "17px", "매우크게": "19px"}
        return _tmap.get(val, default)

    _base_raw = _c12.get("base_font_size", "기본")
    _head_raw = _c12.get("heading_font_size", "기본")
    html_base_size = _to_px(_base_raw, "15px")
    # heading_font_size는 종종 "fs16 (폰트 종류나 굵기로 구분)" 같이 설명 포함
    html_head_size = _to_px(_head_raw.split()[0] if _head_raw else "기본", "17px")
    # 소제목은 base보다 2px 크게 (명시 없을 때)
    if html_head_size == html_base_size:
        html_head_size = f"{int(html_base_size.replace('px','')) + 2}px"

    # 3) 색상: c11_visual_symbols.text_colors 에서 실제 hex 추출
    #    (c12.color_examples 는 hex가 아닌 텍스트 설명이므로 사용 금지)
    _raw_text_colors = _c11.get("text_colors", [])
    _hex_colors = [c for c in _raw_text_colors
                   if c and c.startswith("#") and len(c) in (4, 7, 9)]
    # 기본 글자색(#141414 #333333 등) 및 무채색(gray) 제외하고 강조색 우선
    _body_defaults = {"#141414", "#333333", "#1a1a1a", "#000000", "#222222"}
    def _is_gray(hex_c):
        """R≈G≈B 인 무채색 여부 (threshold 15)"""
        try:
            h = hex_c.lstrip('#')
            if len(h) == 3: h = h[0]*2 + h[1]*2 + h[2]*2
            r,g,b = int(h[0:2],16), int(h[2:4],16), int(h[4:6],16)
            return max(r,g,b) - min(r,g,b) < 15
        except Exception: return False
    _accent_hex = [c for c in _hex_colors
                   if c.lower() not in _body_defaults and not _is_gray(c)]
    html_accent_color  = _accent_hex[0] if _accent_hex else ""
    html_text_colors   = _hex_colors[:5]  # 전체 색상 목록

    # 4) 하이라이트 색상
    _raw_highlights = _c11.get("highlight_colors", [])
    _hl_hex = [c for c in _raw_highlights if c and c.startswith("#")]
    html_highlight_color = _hl_hex[0] if _hl_hex else ""

    # 5) 중앙정렬
    _center_txt = str(_c11.get("center_align_usage", "")).lower()
    html_center_headings = any(w in _center_txt for w in
                               ["매우", "빈번", "자주", "항상", "많이", "중앙 정렬 사용"])

    # 6) 볼드 허용 (빈도 3 이상, 또는 텍스트값 "빈번/자주/많이" 등)
    _bold_raw = _c12.get("bold_frequency", 1)
    try:
        html_bold_allowed = int(_bold_raw or 1) >= 3
    except (ValueError, TypeError):
        _bold_txt = str(_bold_raw).lower()
        html_bold_allowed = any(w in _bold_txt for w in ["빈번", "자주", "많이", "high", "freq"])

    # 7) 줄간격: Naver 기본 line-height=2.2, 빈번한 줄바꿈 스타일이면 1.8
    _lbreak = str(_c11.get("line_break_style", "")).lower()
    html_line_height = "1.8" if any(w in _lbreak for w in ["과감", "잦은", "짧은"]) else "2.2"

    # 8) 소제목용 두 번째 폰트 (font_families[1])
    _heading_raw_font = _font_names[1] if len(_font_names) > 1 else _raw_font
    html_heading_font = _font_map.get(_heading_raw_font.lower().replace(" ", ""),
                        _font_map.get(_heading_raw_font, html_font))

    # 9) 작은 폰트 크기 (font_size_levels 에서 가장 작은 값)
    _fs_levels = _c12.get("font_size_levels", [])
    _fs_px = sorted([_to_px(s, "13px") for s in _fs_levels],
                    key=lambda x: int(x.replace("px", "")))
    html_small_font_size = _fs_px[0] if _fs_px else "13px"

    # 10) c21 인라인 서식 — 색상 전환 패턴 추출
    _c21 = (dna_analysis or {}).get("c21_inline_formatting", {})
    _c21_color_ex = _c21.get("color_switch_examples", [])
    # c21 색상 예시에서 hex 추출 (없으면 c11 accent 사용)
    _c21_hex = [c for c in _c21_color_ex if isinstance(c, str) and c.startswith("#")]
    html_inline_accent = _c21_hex[0] if _c21_hex else html_accent_color
    # 인라인 폰트 전환 빈도
    _font_sw = str(_c21.get("font_switch_frequency", "")).lower()
    html_inline_font_switch = any(w in _font_sw for w in ["자주", "매우", "빈번", "항상"])

    # 11) c22 콘텐츠 패턴 — 반드시 포함 요소 (프롬프트용)
    _c22 = (dna_analysis or {}).get("c22_content_patterns", {})
    html_must_include = _c22.get("must_include_elements", [])

    try:
        html_style = {
            "font_family":          html_font,
            "heading_font_family":  html_heading_font,
            "base_font_size":       html_base_size,
            "heading_font_size":    html_head_size,
            "small_font_size":      html_small_font_size,
            "accent_color":         html_accent_color,
            "inline_accent_color":  html_inline_accent,
            "text_colors":          html_text_colors,
            "highlight_color":      html_highlight_color,
            "bold_allowed":         html_bold_allowed,
            "line_height":          html_line_height,
            "center_headings":      html_center_headings,
            "inline_font_switch":   html_inline_font_switch,
            "must_include":         html_must_include,
        }
    except Exception:
        html_style = {}

    return jsonify({
        "success": True,
        "versions": package.get("versions", []),
        "output_id": output_id,
        "output_dir": str(OUTPUT_DIR),
        "generation_mode": generation_mode,
        "images": extracted_image_urls,
        "html_style": html_style,
        "source_bundle": {
            "sources": material_bundle.get("sources", []),
            "warnings": material_bundle.get("warnings", []),
        },
    })


# ============================================================
# API: Blog Image Generation (On-Demand)
# ============================================================

@app.route('/api/blog/suggest-prompts', methods=['POST'])
@login_required
def suggest_blog_prompts():
    """블로그 본문 기반 이미지 생성 프롬프트 제안"""
    data = request.json
    blog_content = data.get("content", "")
    target_audience = data.get("target_audience", "일반 시민")
    content_angle = data.get("content_angle", "정보전달형")
    
    if not blog_content or not blog_content.strip():
        return jsonify({"error": "프롬프트 제안을 위한 블로그 본문이 필요합니다."}), 400
    
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return jsonify({"error": "GEMINI_API_KEY가 설정되지 않았습니다."}), 500
    
    try:
        client = genai.Client(api_key=api_key)
        from image_service import extract_image_prompts
        prompts = extract_image_prompts(blog_content, client, target_audience, content_angle)
        
        return jsonify({
            "success": True,
            "prompts": prompts
        })
    except Exception as e:
        return jsonify({"error": f"프롬프트 추출 실패: {str(e)}"}), 500


# ============================================================
# 보정 루프: AI글 vs 실제 통과된 글 비교 분석
# ============================================================

@app.route('/api/blog/calibrate', methods=['POST'])
@login_required
def calibrate_blog():
    """AI 생성 글과 실제 통과된 글을 비교 분석해 보정 기록 저장"""
    data = request.get_json()
    ai_title    = data.get("ai_title", "")
    ai_content  = data.get("ai_content", "")
    approved_title   = data.get("approved_title", "")
    approved_content = data.get("approved_content", "")
    style_template_id = data.get("style_template_id", "")
    blog_dna_id       = data.get("blog_dna_id", "")

    if not ai_content or not approved_content:
        return jsonify({"error": "AI 생성 글과 실제 통과된 글 모두 필요합니다."}), 400

    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return jsonify({"error": "Gemini API 키가 설정되지 않았습니다."}), 500

    client = genai.Client(api_key=api_key)

    prompt = f"""당신은 블로그 글쓰기 스타일 분석 전문가입니다.
아래 두 글을 비교해 실제 통과된 글의 특징을 분석해주세요.
두 글은 같은 배포자료(보도자료)를 기반으로 작성되었지만,
하나는 AI가 초안을 쓴 것이고, 다른 하나는 실제 검수를 통과한 글입니다.

[AI 생성 글]
제목: {ai_title}
{ai_content[:3000]}

[실제 통과된 글]
제목: {approved_title}
{approved_content[:3000]}

다음 JSON 형식으로 분석해주세요:
{{
  "do_more": [
    "통과된 글에서 더 많이 사용된 표현/패턴 (구체적으로, 예시 포함)"
  ],
  "do_less": [
    "AI 글에만 있고 통과된 글에서 없앤 것 (구체적으로)"
  ],
  "tone_shift": "전반적인 톤 변화 설명 (예: '더 친근하고 구어체로 바뀜')",
  "structure_diff": "구조적 차이점 (예: '소제목 없이 자연스럽게 흐르는 방식 선호')",
  "length_diff": "길이/분량 차이 및 방향 (예: '더 짧고 단락 사이 여백 많음')",
  "key_phrases": [
    "통과된 글에서 쓰인 특징적 표현이나 문장 패턴"
  ],
  "calibration_prompt": "다음 글 생성 시 AI에게 줄 한국어 지침 3~5문장 (이 분석을 요약한 핵심 가이드)"
}}"""

    try:
        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=prompt
        )
        analysis = parse_ai_json(response.text)
    except Exception as e:
        return jsonify({"error": f"분석 실패: {e}"}), 500

    cal_id = f"CAL_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    record = {
        "calibration_id": cal_id,
        "style_template_id": style_template_id,
        "blog_dna_id": blog_dna_id,
        "ai_title": ai_title,
        "ai_content": ai_content[:2000],
        "approved_title": approved_title,
        "approved_content": approved_content[:2000],
        "analysis": analysis,
        "calibration_prompt": analysis.get("calibration_prompt", ""),
        "created_at": datetime.now().isoformat()
    }
    save_path = CALIBRATIONS_DIR / f"{cal_id}.json"
    save_path.write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f"[OK] 보정 기록 저장: {cal_id}")

    return jsonify({"ok": True, "calibration_id": cal_id, "analysis": analysis})


@app.route('/api/blog/calibrate-from-url', methods=['POST'])
@login_required
def calibrate_from_url():
    """AI 블로그 + 실제 통과된 글 URL로 보정 분석.
    blog_id로 저장본 로드 OR ai_title/ai_content 직접 전달 모두 지원.
    """
    data = request.get_json()
    blog_id       = data.get("blog_id", "").strip()
    approved_url  = data.get("approved_url", "").strip()
    style_template_id = data.get("style_template_id", "")
    blog_dna_id       = data.get("blog_dna_id", "")

    # AI 글: 직접 전달 or 저장본 로드
    ai_title   = data.get("ai_title", "").strip()
    ai_content = data.get("ai_content", "").strip()

    if not ai_title and not ai_content:
        if not blog_id:
            return jsonify({"error": "blog_id 또는 ai_title/ai_content가 필요합니다."}), 400
        fp = OUTPUT_DIR / f"{blog_id}.json"
        if not fp.exists():
            return jsonify({"error": "저장된 블로그를 찾을 수 없습니다."}), 404
        blog_data = load_blog_package(fp)
        versions  = blog_data.get("versions", [])
        version_type = data.get("version_type", "formal")
        ai_version = next((v for v in versions if v.get("version_type") == version_type), None)
        if not ai_version and versions:
            ai_version = versions[0]
        if not ai_version:
            return jsonify({"error": "블로그 버전 데이터가 없습니다."}), 400
        ai_title   = ai_version.get("title", "")
        ai_content = ai_version.get("content", "")
        if not style_template_id:
            style_template_id = blog_data.get("style_template_id", "")

    if not approved_url:
        return jsonify({"error": "approved_url이 필요합니다."}), 400

    # 통과된 글 URL 크롤링
    try:
        approved_text = fetch_url_text(approved_url)
    except Exception as e:
        return jsonify({"error": f"URL 크롤링 실패: {e}"}), 400

    if not approved_text or len(approved_text) < 100:
        return jsonify({"error": "URL에서 충분한 텍스트를 추출하지 못했습니다."}), 400

    # 첫 줄을 제목으로, 나머지를 본문으로 분리
    lines = [l.strip() for l in approved_text.splitlines() if l.strip()]
    approved_title   = lines[0] if lines else ""
    approved_content = "\n".join(lines[1:]) if len(lines) > 1 else approved_text

    # Gemini 비교 분석
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return jsonify({"error": "Gemini API 키가 설정되지 않았습니다."}), 500

    client = genai.Client(api_key=api_key)

    prompt = f"""당신은 블로그 글쓰기 스타일 분석 전문가입니다.
아래 두 글을 비교해 실제 통과된 글의 특징을 분석해주세요.
같은 배포자료(보도자료)를 기반으로 작성됐지만,
하나는 AI 초안이고 다른 하나는 실제 검수를 통과해 게시된 글입니다.

[AI 초안]
제목: {ai_title}
{ai_content[:3000]}

[실제 게시된 글 (URL: {approved_url})]
제목: {approved_title}
{approved_content[:3000]}

다음 JSON 형식으로 분석해주세요:
{{
  "do_more": ["통과된 글에서 더 많이 사용된 표현/패턴 (구체적으로, 예시 포함)"],
  "do_less": ["AI 글에만 있고 통과된 글에서 없앤 것 (구체적으로)"],
  "tone_shift": "전반적인 톤 변화 설명",
  "structure_diff": "구조적 차이점",
  "length_diff": "길이/분량 차이 및 방향",
  "key_phrases": ["통과된 글에서 쓰인 특징적 표현이나 문장 패턴"],
  "similarity_score": 현재_일치율_0_to_100_정수,
  "calibration_prompt": "다음 글 생성 시 AI에게 줄 한국어 지침 3~5문장 (이 분석을 요약한 핵심 가이드)"
}}"""

    try:
        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=prompt
        )
        analysis = parse_ai_json(response.text)
    except Exception as e:
        return jsonify({"error": f"분석 실패: {e}"}), 500

    # 분석만 반환 — 저장은 /api/blog/calibration/save 에서 수동으로
    return jsonify({
        "ok": True,
        "analysis": analysis,
        "meta": {
            "blog_id": blog_id,
            "approved_url": approved_url,
            "style_template_id": style_template_id,
            "blog_dna_id": blog_dna_id,
            "ai_title": ai_title,
            "approved_title": approved_title,
        }
    })


@app.route('/api/blog/calibration/save', methods=['POST'])
@login_required
def save_calibration():
    """사용자가 선택한 항목으로 보정 기록 저장"""
    data = request.get_json()
    analysis        = data.get("analysis", {})
    selected_items  = data.get("selected_items", [])   # [{category, text}, ...]
    custom_prompt   = data.get("calibration_prompt", "").strip()
    meta            = data.get("meta", {})

    if not selected_items:
        return jsonify({"error": "선택된 항목이 없습니다."}), 400

    # 선택된 항목 기반으로 calibration_prompt 재구성 (없으면 전달받은 것 사용)
    if not custom_prompt:
        do_more   = [i["text"] for i in selected_items if i.get("category") == "더 활용"]
        do_less   = [i["text"] for i in selected_items if i.get("category") == "줄일 것"]
        singles   = [i for i in selected_items if i.get("category") not in ("더 활용", "줄일 것")]
        parts = []
        if do_more:  parts.append(f"더 활용할 것: {', '.join(do_more)}")
        if do_less:  parts.append(f"줄일 것: {', '.join(do_less)}")
        for s in singles:
            parts.append(f"{s.get('category', '')}: {s.get('text', '')}")
        custom_prompt = ". ".join(parts)

    cal_id = f"CAL_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    record = {
        "calibration_id": cal_id,
        "blog_id":            meta.get("blog_id", ""),
        "approved_url":       meta.get("approved_url", ""),
        "style_template_id":  meta.get("style_template_id", ""),
        "blog_dna_id":        meta.get("blog_dna_id", ""),
        "ai_title":           meta.get("ai_title", ""),
        "approved_title":     meta.get("approved_title", ""),
        "analysis":           analysis,
        "selected_items":     selected_items,
        "calibration_prompt": custom_prompt,
        "similarity_score":   analysis.get("similarity_score", 0),
        "created_at":         datetime.now().isoformat()
    }
    save_path = CALIBRATIONS_DIR / f"{cal_id}.json"
    save_path.write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f"[OK] 보정 기록 저장: {cal_id} ({len(selected_items)}개 항목 선택)")

    return jsonify({"ok": True, "calibration_id": cal_id})


@app.route('/api/blog/calibrations', methods=['GET'])
@login_required
def list_calibrations():
    """보정 기록 목록 조회"""
    style_template_id = request.args.get("style_template_id", "")
    records = []
    for f in sorted(CALIBRATIONS_DIR.glob("CAL_*.json"),
                    key=lambda x: x.stat().st_mtime, reverse=True):
        try:
            d = json.loads(f.read_text(encoding='utf-8'))
            if style_template_id and d.get("style_template_id") != style_template_id:
                continue
            records.append({
                "calibration_id": d.get("calibration_id"),
                "style_template_id": d.get("style_template_id"),
                "blog_dna_id": d.get("blog_dna_id"),
                "approved_url": d.get("approved_url", ""),
                "ai_title": d.get("ai_title", "")[:50],
                "approved_title": d.get("approved_title", "")[:50],
                "tone_shift": d.get("analysis", {}).get("tone_shift", ""),
                "similarity_score": d.get("similarity_score", d.get("analysis", {}).get("similarity_score", 0)),
                "calibration_prompt": d.get("calibration_prompt", "")[:120],
                "created_at": d.get("created_at", "")[:10]
            })
        except Exception:
            pass
    return jsonify({"calibrations": records})


@app.route('/api/blog/calibration/<cal_id>', methods=['DELETE'])
@login_required
def delete_calibration(cal_id):
    path = CALIBRATIONS_DIR / f"{cal_id}.json"
    if path.exists():
        path.unlink()
    return jsonify({"ok": True})


@app.route('/api/blog/generate-images', methods=['POST'])
@login_required
def generate_blog_images():
    """블로그 본문 또는 커스텀 프롬프트 기반 AI 이미지 생성 (On-Demand)"""
    data = request.json
    
    blog_content = data.get("content", "")
    target_audience = data.get("target_audience", "일반 시민")
    content_angle = data.get("content_angle", "정보전달형")
    custom_prompts = data.get("prompts", [])
    
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return jsonify({"error": "GEMINI_API_KEY가 설정되지 않았습니다."}), 500
    
    try:
        client = genai.Client(api_key=api_key)
        from image_service import generate_images_for_blog, generate_images
        
        if custom_prompts:
            # 사용자가 수정한 프롬프트로 생성
            images = generate_images(custom_prompts, client, OUTPUT_DIR)
        else:
            # 블로그 본문에서 자동 추출하여 생성 (기존 방식)
            if not blog_content or not blog_content.strip():
                return jsonify({"error": "이미지 생성을 위한 본문 또는 프롬프트가 필요합니다."}), 400
            images = generate_images_for_blog(blog_content, client, target_audience, content_angle, OUTPUT_DIR)
        
        if not images:
            return jsonify({"error": "이미지를 생성하지 못했습니다. 다시 시도해주세요."}), 500
        
        return jsonify({
            "success": True,
            "images": images,
            "count": len(images)
        })
    except Exception as e:
        return jsonify({"error": f"이미지 생성 실패: {str(e)}"}), 500

# ============================================================
# API: Blog Save & Export
# ============================================================

@app.route('/api/blog/save', methods=['POST'])
@login_required
def save_blog():
    """편집된 블로그 글을 JSON 파일로 저장"""
    data = request.get_json()
    if not data:
        return jsonify({"error": "데이터가 없습니다."}), 400
    
    title = data.get("title", "")
    content = data.get("content", "")
    version_type = data.get("version_type", "unknown")
    version_label = data.get("version_label", "")
    tags = data.get("tags", [])
    output_id = data.get("output_id", "")
    output_dir = data.get("output_dir", "")
    
    if not title or not content:
        return jsonify({"error": "제목과 본문을 입력해주세요."}), 400
    
    try:
        if output_dir and Path(output_dir).exists():
            save_dir = Path(output_dir)
        else:
            save_dir = OUTPUT_DIR

        if output_id:
            package, save_path = update_blog_package_version(
                output_dir=save_dir,
                output_id=output_id,
                version_type=version_type,
                updated_fields={
                    "title": title,
                    "content": content,
                    "tags": tags,
                    "meta_description": data.get("meta_description", ""),
                },
            )
        else:
            filename = f"edited_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            save_path = save_dir / filename
            with open(save_path, 'w', encoding='utf-8') as f:
                json.dump({
                    "title": title,
                    "content": content,
                    "version_type": version_type,
                    "version_label": version_label,
                    "tags": tags,
                    "edited_at": datetime.now().isoformat(),
                    "is_edited": True,
                }, f, ensure_ascii=False, indent=2)

        return jsonify({
            "success": True,
            "blog_id": output_id or Path(save_path).stem,
            "saved_path": str(save_path),
            "filename": Path(save_path).name
        })
    except Exception as e:
        return jsonify({"error": f"저장 실패: {str(e)}"}), 500


@app.route('/api/blog/export-docs', methods=['POST'])
@login_required
def export_blog_docs():
    """편집된 블로그 글을 DOCX 파일로 내보내기"""
    data = request.get_json()
    if not data:
        return jsonify({"error": "데이터가 없습니다."}), 400
    
    title = data.get("title", "제목 없음")
    content = data.get("content", "")
    version_label = data.get("version_label", "")
    tags = data.get("tags", [])
    
    if not content:
        return jsonify({"error": "본문 내용이 없습니다."}), 400
    
    try:
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        # 스타일 설정
        style = doc.styles['Normal']
        font = style.font
        font.name = '맑은 고딕'
        font.size = Pt(11)
        
        # 버전 라벨
        if version_label:
            label_para = doc.add_paragraph()
            label_run = label_para.add_run(f"[{version_label}]")
            label_run.font.size = Pt(9)
            label_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 제목
        heading = doc.add_heading(title, level=1)
        
        # 본문
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            stripped = para_text.strip()
            if stripped:
                doc.add_paragraph(stripped)
        
        # 태그
        if tags:
            doc.add_paragraph()
            tag_para = doc.add_paragraph()
            tag_run = tag_para.add_run(' '.join(f'#{t}' for t in tags))
            tag_run.font.size = Pt(9)
            tag_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # 메모리에 저장 후 전송
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        from flask import send_file
        safe_title = re.sub(r'[^\w가-힣\s]', '', title)[:30].strip() or 'blog'
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{safe_title}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({"error": f"DOCX 생성 실패: {str(e)}"}), 500


# ============================================================
# API: Blog Collection (blog_pull 통합)
# ============================================================

# ─── match-test 제거 (persona 의존) ─── 이하 생략됨

# ============================================================
# API: Blog Collection (blog_pull 통합)
# ============================================================

@app.route('/api/blog/search-users', methods=['GET'])
def search_naver_blog_users():
    """네이버 블로그 사용자 검색 (키워드 → 블로그 ID 목록)"""
    query = request.args.get('q', '').strip()
    if not query or len(query) < 2:
        return jsonify({"results": []})
    try:
        from bs4 import BeautifulSoup as _BS
        _headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
            "Accept-Language": "ko-KR,ko;q=0.9",
            "Referer": "https://search.naver.com/",
        }
        import urllib.parse as _up
        url = f"https://search.naver.com/search.naver?where=blog&query={_up.quote(query)}&sm=tab_opt&nso=so%3Add%2Cp%3Aall"
        resp = http_requests.get(url, headers=_headers, timeout=8)
        soup = _BS(resp.text, 'html.parser')

        results = []
        seen = set()
        _SKIP = {'postview', 'blognews', 'blogid', 'search', 'nblog', 'connect'}

        for a in soup.find_all('a', href=True):
            href = a['href']
            m = re.search(r'(?:m\.)?blog\.naver\.com/([A-Za-z0-9_]{3,})', href)
            if not m:
                continue
            bid = m.group(1).lower()
            if bid in _SKIP or bid in seen:
                continue

            # 블로그명 추출 시도 (상위 요소에서)
            name = ''
            try:
                card = a.find_parent(class_=re.compile(r'title|blog|author|writer', re.I))
                if card:
                    name_el = card.find(class_=re.compile(r'name|title|author', re.I))
                    if name_el:
                        name = name_el.get_text(strip=True)[:40]
            except Exception:
                pass

            seen.add(bid)
            results.append({
                "blog_id": m.group(1),
                "name": name,
                "url": f"https://blog.naver.com/{m.group(1)}"
            })
            if len(results) >= 10:
                break

        return jsonify({"results": results})
    except Exception as e:
        return jsonify({"results": [], "error": str(e)})


try:
    from run_crawler import get_blog_id, get_post_list, get_post_content, get_post_content_with_style, save_results
    import run_crawler as _run_crawler
    _run_crawler.OUTPUT_DIR = str(BLOG_COLLECTIONS_DIR)
    BLOG_PULL_AVAILABLE = True
except Exception as e:
    BLOG_PULL_AVAILABLE = False
    get_blog_id = get_post_list = get_post_content = get_post_content_with_style = save_results = None
    print(f"[WARN] blog_pull 크롤러를 불러오지 못했습니다: {e}")
import time as _time

@app.route('/api/blog/collect', methods=['POST'])
@login_required
def collect_blog():
    """네이버 블로그 글 수집 (blog_pull 크롤러 통합)"""
    if not BLOG_PULL_AVAILABLE:
        return jsonify({"error": "blog_pull 크롤러가 설치되지 않아 블로그 수집 기능을 사용할 수 없습니다."}), 503

    data = request.json
    
    blog_input = data.get("blog_id", "").strip()
    count = min(int(data.get("count", 10)), 30)
    
    if not blog_input:
        return jsonify({"error": "블로그 주소 또는 ID를 입력해주세요."}), 400
    
    blog_id = get_blog_id(blog_input)
    if not blog_id:
        return jsonify({"error": f"올바른 블로그 주소가 아닙니다: {blog_input}"}), 400
    
    try:
        # STEP 1: 글 목록 가져오기
        posts = get_post_list(blog_id, count)
        
        if not posts:
            return jsonify({"error": "글 목록을 가져올 수 없습니다. 블로그 주소를 확인해주세요."}), 404
        
        # STEP 2: 본문 + 스타일 메타 수집
        for post in posts:
            result = get_post_content_with_style(blog_id, post['logNo'])
            post['content'] = result.get('text') or "(본문 추출 실패)"
            post['style_meta'] = result.get('style_meta', {})
            _time.sleep(0.3)
        
        # STEP 3: 저장
        folder = save_results(blog_id, posts)
        
        total_chars = sum(len(p.get('content', '')) for p in posts)
        
        return jsonify({
            "blog_id": blog_id,
            "blog_url": f"https://blog.naver.com/{blog_id}",
            "post_count": len(posts),
            "total_chars": total_chars,
            "output_folder": os.path.basename(folder),
            "posts": [
                {
                    "title": p['title'],
                    "date": p['addDate'],
                    "url": p['url'],
                    "content_length": len(p.get('content', '')),
                    "content_preview": p.get('content', '')[:200]
                }
                for p in posts
            ]
        })
        
    except Exception as e:
        return jsonify({"error": f"블로그 수집 실패: {str(e)}"}), 500


@app.route('/api/blog/collections', methods=['GET'])
@login_required
def list_blog_collections():
    """수집된 블로그 컬렉션 목록 (블로그 ID별 통합 버전)"""
    blog_map = {}
    
    if BLOG_COLLECTIONS_DIR.exists():
        # 날짜 내림차순 정렬 (최신 폴더 먼저)
        for item in sorted(BLOG_COLLECTIONS_DIR.iterdir(), reverse=True):
            if item.is_dir() and not item.name.startswith('.'):
                data_file = item / "_data.json"
                if data_file.exists():
                    try:
                        with open(data_file, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                        
                        blog_id = data.get("blog_id", "")
                        if not blog_id: continue
                        
                        posts = data.get("posts", [])
                        post_count = len(posts)
                        total_chars = sum(len(p.get("content", "")) for p in posts)
                        collected_at = data.get("collected_at", item.name)
                        
                        if blog_id not in blog_map:
                            blog_map[blog_id] = {
                                "blog_id": blog_id,
                                "folders": [item.name],
                                "post_count": post_count,
                                "total_chars": total_chars,
                                "last_collected_at": collected_at
                            }
                        else:
                            # 기존 데이터에 합치기 (Merging)
                            blog_map[blog_id]["folders"].append(item.name)
                            blog_map[blog_id]["post_count"] += post_count
                            blog_map[blog_id]["total_chars"] += total_chars
                            # (이미 날짜 역순 정렬이므로 처음 발견된 게 가장 최신일 가능성이 높음)
                    except:
                        pass
    
    # 결과를 리스트로 변환
    collections = list(blog_map.values())
    return jsonify({"collections": collections})


@app.route('/api/blog/analyze-status', methods=['POST'])
@login_required
def analyze_blog_status():
    """수집된 블로그 글을 AI로 분석하여 상태 파악 (모든 수집 데이터 통합 분석)"""
    data = request.json
    blog_id = data.get("blog_id", "")  # 이제 folder 대신 blog_id를 직접 받음
    
    if not blog_id:
        return jsonify({"error": "분석할 블로그 ID를 선택해주세요."}), 400
    
    all_posts = []
    
    # 해당 blog_id를 가진 모든 폴더 찾아서 데이터 합치기
    if BLOG_COLLECTIONS_DIR.exists():
        for item in BLOG_COLLECTIONS_DIR.iterdir():
            if item.is_dir() and not item.name.startswith('.'):
                data_file = item / "_data.json"
                if data_file.exists():
                    try:
                        with open(data_file, 'r', encoding='utf-8') as f:
                            collection = json.load(f)
                        if collection.get("blog_id") == blog_id:
                            all_posts.extend(collection.get("posts", []))
                    except:
                        pass
    
    if not all_posts:
        return jsonify({"error": f"'{blog_id}'에 대한 수집 데이터를 찾을 수 없습니다."}), 404
    
    # 중복 글 제거 (url 또는 logNo 기준)
    seen_urls = set()
    unique_posts = []
    for post in all_posts:
        url = post.get("url")
        if url and url not in seen_urls:
            seen_urls.add(url)
            unique_posts.append(post)
    
    try:
        # 날짜순 정렬 (최신순)
        unique_posts.sort(key=lambda x: x.get('addDate', ''), reverse=True)
        
        # 블로그 글 요약 텍스트 생성 (통합된 데이터 중 최근 15개 분석)
        blog_summary = ""
        for i, post in enumerate(unique_posts[:15], 1):
            content = post.get("content", "")[:1500]
            blog_summary += f"\n\n--- 글 {i}: {post.get('title', '')} (날짜: {post.get('addDate', '')}) ---\n{content}"

        # 시각적 스타일 메타 집계 (HTML에서 추출한 데이터)
        from collections import Counter as _Counter
        style_metas = [p.get("style_meta", {}) for p in unique_posts if p.get("style_meta")]
        visual_summary = ""
        if style_metas:
            avg_center = round(sum(m.get("center_align_ratio", 0) for m in style_metas) / len(style_metas), 2)
            all_accent = [c for m in style_metas for c, _ in m.get("accent_colors", [])]
            all_highlight = [c for m in style_metas for c, _ in m.get("highlight_colors", [])]
            all_fonts = [f for m in style_metas for f in m.get("dominant_fonts", [])]
            all_sizes = [s for m in style_metas for s in m.get("font_sizes", [])]
            bold_total = sum(m.get("bold_count", 0) for m in style_metas)
            italic_total = sum(m.get("italic_count", 0) for m in style_metas)
            has_quote = any(m.get("has_quote_block") for m in style_metas)

            top_accents = _Counter(all_accent).most_common(5)
            top_highlights = _Counter(all_highlight).most_common(3)
            top_fonts = _Counter(all_fonts).most_common(3)
            top_sizes = _Counter(all_sizes).most_common(3)

            visual_summary = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
【HTML 시각적 스타일 데이터 ({len(style_metas)}개 글 분석)】
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
중앙정렬 비율: {avg_center*100:.0f}% ({"매우 자주 씀" if avg_center > 0.5 else "가끔 씀" if avg_center > 0.2 else "거의 안 씀"})
강조 색상(텍스트): {', '.join(f'{c}(x{n})' for c, n in top_accents) or '없음'}
강조 색상(배경/하이라이트): {', '.join(f'{c}(x{n})' for c, n in top_highlights) or '없음'}
사용 폰트: {', '.join(f for f, _ in top_fonts) or '기본 폰트'}
폰트 크기: {', '.join(s for s, _ in top_sizes) or '기본'}
볼드 사용 횟수: {bold_total}회
이탤릭 사용 횟수: {italic_total}회
인용구 블록 사용: {"있음" if has_quote else "없음"}
"""

        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            return jsonify({"error": "GEMINI_API_KEY가 설정되지 않았습니다."}), 500

        client = genai.Client(api_key=api_key)

        analysis_prompt = f"""당신은 블로그 글쓰기 DNA 분석 전문가입니다.
아래는 네이버 블로그 '{blog_id}'에서 수집한 최근 글들입니다.
이 블로거의 글쓰기 스타일을 **100% 재현**할 수 있을 만큼 철저하게 분석하세요.
목표: 이 분석 결과만 보고 AI가 써도 원본 블로거와 구분이 안 될 정도로 완벽한 스타일 복제.

{blog_summary}
{visual_summary}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
【분석 지침 — 반드시 준수】
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- 22가지 카테고리 전부 빠짐없이 채울 것
- 모든 수치(비율%, 1-10점, 글자수)는 반드시 숫자로 기재
- examples 필드에는 반드시 원문 그대로 발췌 (요약 금지)
- 꺽쇠·괄호·이모지·특수기호는 실제 문자 그대로 기재 (설명 금지)
- 이 분석만으로 원본 블로거처럼 글 쓸 수 있을 만큼 구체적으로

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
【출력 JSON 스키마 (22가지 카테고리 — 완전 심층 분석)】
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{{
  "c1_template_structure": {{
    "title": "글 전체 구조/템플릿",
    "overall_pattern": "주된 구조 유형 (서론-본론-결론형 / 리스트나열형 / Q&A형 / 체험기형 / 비교형 / 혼합형 등)",
    "section_count": "평균 섹션 수 (숫자로 기재, 예: 4~6개)",
    "section_flow": ["섹션 순서 패턴 — 구체적으로 (예: '인사+상황설명 → 본론A → 본론B → 정리 → CTA')"],
    "heading_style": "소제목 사용 방식 (있음/없음, 사용 빈도, 포맷)",
    "subheading_format": "소제목 포맷 패턴 (꺽쇠 포함 여부, 이모지 선행 여부, 굵기 등 — 실제 예시 포함)",
    "info_hierarchy": "정보 배치 순서 (결론 먼저형 / 과정 순서형 / 배경→결론형)",
    "template_consistency": "글마다 템플릿 일관성 (높음/보통/낮음)",
    "examples": ["실제 글 하나의 구조를 요약한 예시 — 섹션별로"]
  }},
  "c2_tone_mood": {{
    "title": "전체 톤/분위기",
    "primary_tone": "주된 톤 — 구체적으로 (예: '친근하고 수다스러운', '따뜻한 정보전달형', '캐주얼한 일기체')",
    "secondary_tone": "보조 톤",
    "formality_level": "격식 수준 (1-10, 1=완전 비격식, 10=완전 격식)",
    "warmth_level": "친밀감 수준 (1-10)",
    "energy_level": "활발함/에너지 수준 (1-10, 1=차분, 10=매우 활발)",
    "positivity_level": "긍정성 수준 (1-10)",
    "consistency": "톤 일관성 (높음/보통/글마다 다름)",
    "tone_shift_pattern": "글 내 톤 변화 패턴 (예: '도입부 친근 → 본문 정보전달 → 마무리 감성')",
    "examples": ["톤이 잘 드러나는 실제 문장 3개 이상"]
  }},
  "c3_speech_endings": {{
    "title": "종결어미/어투 패턴",
    "primary_endings": ["가장 자주 쓰는 종결어미 TOP5 — 각 어미와 예문 포함 (예: '~요 (60%)', '~해요 (20%)')"],
    "secondary_endings": ["2순위 종결어미 목록"],
    "formality_mix": "격식체:비격식체 비율 (예: '2:8 — 비격식 압도적')",
    "question_ending_style": "의문형 어미 패턴 (예: '~인가요?', '~지 않을까요?', '~죠?' 등)",
    "exclamation_style": "감탄/강조형 어미 패턴",
    "consecutive_same_ending": "같은 어미 연속 사용 여부 (예: '~요로 3문장 연속 가능')",
    "reader_address": "독자 호칭 방식 (예: 직접 호칭 없음 / '여러분' / '~하시는 분들' 등)",
    "sentence_end_variety": "문장 끝 변화 다양성 (높음/보통/낮음)",
    "examples": ["종결어미 패턴이 잘 드러나는 연속 문장 발췌"]
  }},
  "c4_sentence_structure": {{
    "title": "문장 구조/길이/리듬",
    "avg_chars_per_sentence": "문장당 평균 글자수 (숫자, 예: 28자)",
    "short_sentence_ratio": "짧은 문장(20자 이하) 비율 (예: 35%)",
    "medium_sentence_ratio": "중간 문장(21-50자) 비율 (예: 50%)",
    "long_sentence_ratio": "긴 문장(51자 이상) 비율 (예: 15%)",
    "complexity": "문장 복잡도 (단문 위주/복문 위주/혼합 — 근거 포함)",
    "rhythm_pattern": "리듬 패턴 — 구체적으로 (예: '짧은 문장 2~3개 → 긴 문장 1개 → 짧은 문장')",
    "list_usage": "리스트/나열 구조 사용 빈도 (1-10)",
    "parenthetical_usage": "삽입구 사용 빈도 (1-10)",
    "leading_phrase_patterns": ["문장 시작 패턴 목록 (예: '그런데', '사실', '이렇게 해서')"],
    "examples": ["리듬/길이 특징이 잘 드러나는 연속 문단 발췌"]
  }},
  "c5_paragraph_composition": {{
    "title": "단락/문단 구성",
    "avg_sentences_per_paragraph": "문단당 평균 문장 수 (숫자, 예: 2.5문장)",
    "avg_paragraphs_per_post": "글당 평균 문단 수 (숫자)",
    "paragraph_opening_pattern": "단락 첫 문장 시작 패턴 (예: '접속사로 시작', '주제문 먼저', '질문으로 시작')",
    "paragraph_closing_pattern": "단락 마지막 문장 패턴",
    "transition_style": "단락 간 전환 방식",
    "whitespace_style": "여백/줄바꿈 활용 방식 (예: '단문 1개=단락 1개', '짧은 단락 선호')",
    "single_sentence_paragraph_ratio": "한 문장짜리 단락 비율 (예: 40%)",
    "content_density": "정보 밀도 (높음/중간/낮음)",
    "examples": ["단락 구성이 잘 보이는 실제 단락 2~3개 발췌"]
  }},
  "c6_signature_expressions": {{
    "title": "시그니처/자주 쓰는 표현",
    "signature_phrases": ["이 블로거만의 시그니처 표현 7개 이상 — 실제 표현 그대로"],
    "paragraph_openers": ["단락 시작에 자주 쓰는 표현 (예: '그리고', '근데', '사실 말하면')"],
    "transition_words": ["접속/전환 표현 7개 이상 — 실제 표현"],
    "emphasis_expressions": ["강조할 때 쓰는 표현 목록"],
    "filler_expressions": ["습관적/군더더기 표현 목록"],
    "affirmation_expressions": ["긍정/동의 추임새 목록 (예: '맞아요', '그렇죠', '역시')"],
    "reaction_expressions": ["반응/감탄 표현 목록 (예: '대박', '진짜', '헐')"],
    "examples": ["실제 글에서 위 표현들이 연속으로 쓰인 발췌"]
  }},
  "c7_vocabulary": {{
    "title": "어휘/용어 선택",
    "level": "어휘 수준 (매우 쉬움/쉬움/보통/전문적)",
    "korean_ratio": "순우리말 비율 % (예: 60%)",
    "sino_korean_ratio": "한자어 비율 % (예: 30%)",
    "foreign_word_ratio": "외래어/영어 혼용 비율 % (예: 10%)",
    "jargon_usage": "전문용어 사용 빈도 (없음/가끔/자주)",
    "trendy_words": "유행어/신조어 사용 여부와 실제 예시",
    "characteristic_words": ["이 블로거만의 특징 어휘 10개 이상 — 실제 단어"],
    "avoided_expressions": "의도적으로 안 쓰는 표현 패턴 (발견된 경우)",
    "number_word_preference": "숫자 표현 선호 (아라비아숫자 선호 / 한글 수사 선호 / 혼용)",
    "examples": ["어휘 특성이 잘 드러나는 문장 발췌"]
  }},
  "c8_rhetoric": {{
    "title": "화법/수사법",
    "storytelling_style": "스토리텔링 방식 (경험 공유형/관찰 서술형/정보 나열형 등) + 비중",
    "persuasion_technique": "설득/전달 기법 — 구체적으로",
    "humor_level": "유머 수준 (1-10)",
    "humor_style": "유머 스타일 (자기비하형/관찰형/상황형/언어유희형 등)",
    "metaphor_frequency": "비유/은유 빈도 (1-10)",
    "repetition_usage": "반복법 사용 여부와 패턴",
    "contrast_usage": "대조/대비 활용 여부",
    "rhetorical_question": "수사 의문문 사용 빈도 (1-10)",
    "self_disclosure_level": "개인 경험/감정 노출 수준 (1-10)",
    "examples": ["화법/수사법이 잘 드러나는 실제 문장 3개"]
  }},
  "c9_opening_patterns": {{
    "title": "도입부 패턴",
    "opening_types": ["자주 쓰는 도입 방식 목록 — 구체적으로 (예: '날씨/계절 언급 + 개인 상황 연결', '직접 결론 제시')"],
    "first_sentence_pattern": "첫 문장의 전형적 패턴 — 구체적으로",
    "opening_length": "도입부 길이 (짧음=1~2문장 / 보통=3~4문장 / 김=5문장 이상)",
    "hook_type": "독자 후킹 방식 (감정 공감형/정보 예고형/질문형/상황 묘사형)",
    "personal_intro_style": "개인 상황 도입 방식",
    "keyword_in_opening": "도입부 키워드 배치 여부",
    "opening_examples": ["실제 도입부 원문 발췌 3개 이상 — 첫 3~4줄 그대로"]
  }},
  "c10_closing_patterns": {{
    "title": "마무리 패턴",
    "closing_types": ["자주 쓰는 마무리 방식 목록 — 구체적으로"],
    "last_sentence_pattern": "마지막 문장의 전형적 패턴",
    "cta_style": "CTA 방식 (구독요청/공감클릭요청/댓글유도/재방문요청/없음 등)",
    "cta_keywords": ["CTA에 자주 쓰는 표현 — 실제 표현으로"],
    "farewell_style": "인사 마무리 방식 (없음/짧게/감성적으로 등)",
    "emotion_at_close": "마무리의 감정 톤",
    "summary_habit": "마지막에 요약하는 습관 여부",
    "closing_examples": ["실제 마무리부 원문 발췌 3개 이상 — 마지막 3~4줄 그대로"]
  }},
  "c11_visual_symbols": {{
    "title": "시각 요소/이모지/기호",
    "emoji_frequency": "이모지 사용 빈도 (1-10)",
    "emoji_list": ["자주 쓰는 이모지 — 실제 문자로 15개 이상"],
    "emoji_position": "이모지 위치 패턴 (소제목 앞/문장 끝/단독 줄/문장 중간 등)",
    "emoji_per_post": "글당 평균 이모지 수 (예: 약 12개)",
    "center_align_usage": "중앙정렬 사용 여부와 빈도",
    "text_colors": ["강조 텍스트 색상 — HTML hex 또는 색상명"],
    "highlight_colors": ["하이라이트 색상 — HTML hex 또는 색상명"],
    "separator_patterns": ["구분선/구분 요소 패턴 — 실제 기호로 (예: '─────', '✦ ✦ ✦', '===')"],
    "special_symbols": ["특수기호 목록 — 실제 문자로 20개 이상 (✅ ▶ ➡️ ⭐ ♥ 등)"],
    "symbol_usage_context": "특수기호 사용 맥락 (강조/리스트마커/구분/감정 표현 등)",
    "line_break_style": "줄바꿈 방식 — 구체적으로",
    "formatting_guide": "시각 스타일 재현을 위한 구체적 지침 (3~5문장)",
    "examples": ["시각 요소가 잘 보이는 실제 단락 발췌 — 기호/이모지 그대로"]
  }},
  "c12_typography": {{
    "title": "폰트/글꼴/타이포그래피",
    "font_families": ["사용 폰트명 목록 — HTML font-family 속성 기반"],
    "base_font_size": "본문 기본 크기 (px 또는 '설정 안됨')",
    "heading_font_size": "소제목/강조부 크기",
    "font_size_levels": ["사용하는 모든 크기 단계 목록"],
    "bold_frequency": "볼드 사용 빈도 (1-10)",
    "bold_purpose": "볼드 주 용도",
    "bold_scope": "볼드 범위 (단어 단위/구 단위/문장 전체)",
    "bold_examples": ["볼드 강조 실제 예시 5개 이상 — **볼드부분** 표시"],
    "italic_usage": "기울임꼴 사용 여부/빈도/목적",
    "underline_usage": "밑줄 사용 여부/빈도/목적",
    "strikethrough_usage": "취소선 사용 여부",
    "color_text_frequency": "색상 텍스트 빈도 (1-10)",
    "color_examples": ["색상 텍스트 실제 예시"],
    "highlight_frequency": "형광펜 사용 빈도 (1-10)",
    "highlight_examples": ["형광펜 사용 실제 예시"],
    "typography_guide": "글꼴 스타일 재현 지침 (구체적으로)"
  }},
  "c13_brackets_quotes": {{
    "title": "꺽쇠/괄호/인용부호",
    "angle_brackets": ["사용하는 꺽쇠 종류 — 실제 기호로 (《 》 〈 〉 「 」 『 』 ≪ ≫ < > 중 실제 쓰는 것)"],
    "angle_bracket_purpose": "꺽쇠 주 용도 (소제목마커/강조/인용/카테고리태그 등)",
    "angle_bracket_frequency": "꺽쇠 사용 빈도 (없음/가끔/자주/항상)",
    "angle_bracket_examples": ["꺽쇠 사용 실제 예시 5개 이상 — 기호 그대로"],
    "square_brackets": ["대괄호 종류 — 실제 기호로 (【 】 [ ] ［ ］ 중 실제 쓰는 것)"],
    "square_bracket_purpose": "대괄호 용도와 빈도",
    "round_bracket_pattern": "소괄호 () 사용 패턴 (부연설명/영문병기/수치표시/생략 등) + 빈도",
    "quotation_style": "따옴표 방식 (''/\"\" /「」/없음) + 용도",
    "bracket_combo_pattern": "복합 패턴 (예: 꺽쇠 안에 이모지, 대괄호+소괄호 동시 사용 등)",
    "examples": ["모든 유형 괄호 사용 실제 예시 — 기호 반드시 그대로 포함"]
  }},
  "c14_length_stats": {{
    "title": "글자수/분량 통계",
    "avg_chars_per_post": "평균 글자수 (공백 포함, 숫자로)",
    "min_chars": "최소 글자수",
    "max_chars": "최대 글자수",
    "avg_sentences_per_post": "평균 문장 수 (숫자로)",
    "avg_paragraphs_per_post": "평균 단락 수 (숫자로)",
    "avg_sentences_per_paragraph": "단락당 평균 문장 수 (숫자로)",
    "content_ratio": "서론:본론:결론 분량 비율 (예: '15%:70%:15%')",
    "length_consistency": "분량 일관성 (높음/보통/낮음)",
    "density_guide": "분량 재현 지침 — 구체적으로 (예: '단락 1개 = 2~3문장, 총 10~12단락, 약 1500자 목표')"
  }},
  "c15_title_patterns": {{
    "title": "제목 작성 패턴",
    "avg_length": "평균 제목 글자수 (숫자, 예: 22자)",
    "min_length": "최단 제목 글자수",
    "max_length": "최장 제목 글자수",
    "structure_types": ["제목 구조 유형 목록 — 구체적으로 (예: '지역명+장소명+후기형', '숫자+혜택 나열형')"],
    "keyword_position": "핵심 키워드 주 위치 (앞/중간/끝)",
    "number_usage": "숫자 활용 여부와 형식 (없음/아라비아숫자/한글 수사/TOP형)",
    "bracket_in_title": "제목 내 괄호/꺽쇠 사용 여부와 예시",
    "emotion_words": ["감정 유발/클릭 유도 단어 목록 — 실제 단어"],
    "location_brand_inclusion": "지명/브랜드명 포함 패턴",
    "title_ending_pattern": "제목 끝 패턴 (마침표 없음/어미 형태)",
    "title_keywords": ["자주 등장하는 핵심 키워드 목록"],
    "seo_pattern": "SEO 키워드 배치 패턴",
    "examples": ["실제 제목 7개 이상 — 패턴이 다양하게 드러나는 것"]
  }},
  "c16_image_media": {{
    "title": "이미지/미디어 활용",
    "avg_images_per_post": "글당 평균 이미지 수 (숫자)",
    "image_placement": "이미지 배치 패턴 — 구체적으로 (예: '본문 사이사이 1~2장씩', '글 상단 대표 1장+중간 분산')",
    "caption_style": "캡션 스타일 (없음/짧은 설명/이모지+설명/상세 설명)",
    "image_types": "이미지 종류 (직접 촬영/스크린샷/공식 홍보 이미지/혼합)",
    "layout": "이미지 정렬 (좌정렬/중앙/전체폭/혼합)",
    "use_tables": "표 사용 여부/빈도/목적",
    "use_maps_links": "지도/외부링크 삽입 여부",
    "media_density": "텍스트 대비 이미지 비중 (이미지 중심/균형/텍스트 중심)",
    "thumbnail_pattern": "대표 이미지/썸네일 패턴",
    "examples": ["이미지 활용 특징이 드러나는 캡션/설명 실제 예시"]
  }},
  "c17_punctuation": {{
    "title": "문장부호/구두점 패턴",
    "period_style": "마침표 사용 방식 (매 문장 사용/자주 생략/완전 생략 등)",
    "comma_frequency": "쉼표 사용 빈도 (1-10)",
    "comma_style": "쉼표 사용 방식 (나열할 때만/긴 문장에서만/자유롭게)",
    "exclamation_frequency": "느낌표 빈도 (1-10)",
    "exclamation_style": "느낌표 사용 방식 (단독!/중복!!/감탄사 뒤 등)",
    "question_mark_usage": "물음표 사용 맥락 (의문문에만/수사의문문에도/감탄형으로도)",
    "ellipsis_usage": "말줄임표 (...) 사용 빈도와 맥락",
    "dash_usage": "대시(— 또는 –) 사용 여부와 목적",
    "tilde_usage": "물결표(~) 사용 빈도와 맥락 (1-10)",
    "colon_usage": "콜론(:) 사용 방식",
    "multiple_punct_usage": "복수 부호 (예: !! ?? !? ~~) 사용 여부와 빈도",
    "period_omission_ratio": "마침표 생략 비율 (예: '70% 문장에서 생략')",
    "examples": ["구두점 특징이 잘 드러나는 연속 문장 발췌"]
  }},
  "c18_numbers_data": {{
    "title": "숫자/단위/데이터 표현",
    "numeral_preference": "아라비아숫자 vs 한글 수사 선호도 (예: '아라비아숫자 90% 사용')",
    "price_format": "가격 표기 방식 — 실제 예시 (예: '15,000원' / '1만5천원' / '만오천원')",
    "date_format": "날짜 표기 방식 — 실제 예시 (예: '2024.03.10' / '3월 10일' / '24년 3월')",
    "time_format": "시간 표기 방식 — 실제 예시",
    "unit_style": "단위 표기 스타일 (숫자에 붙여쓰기/띄어쓰기, 한글단위/영문단위)",
    "ranking_format": "순위 표기 방식 (1위/1등/TOP1/첫 번째 중 실제 쓰는 것)",
    "statistics_usage": "통계/수치 인용 빈도 (없음/가끔/자주)",
    "approximation_style": "어림수 표현 방식 (약/대략/~정도/거의 등)",
    "large_number_format": "큰 숫자 표기 (예: 1,000,000 / 100만 / 백만)",
    "examples": ["숫자/데이터 표현이 잘 드러나는 실제 예시 5개 이상"]
  }},
  "c19_reader_engagement": {{
    "title": "독자 상호작용/참여 유도",
    "direct_question_frequency": "독자에게 직접 질문 빈도 (1-10)",
    "empathy_phrases": ["공감 유도 표현 목록 — 실제 표현 (예: '저만 그런 건 아니죠?', '다들 아시죠?')"],
    "inclusive_expressions": ["'우리/함께' 등 포용 표현 목록 — 실제 표현"],
    "experience_sharing_style": "개인 경험 공유 방식 (구체적/추상적/감정 중심/사실 중심)",
    "recommendation_strength": "추천 강도 (약한 제안형/강한 추천형/명령형)",
    "recommendation_expressions": ["추천할 때 쓰는 표현 목록"],
    "social_proof_usage": "사회적 증거 활용 (입소문 언급/후기 인용/전문가 언급 등)",
    "urgency_patterns": ["긴박감/희소성 유도 표현 목록 — 실제 표현"],
    "reader_benefit_emphasis": "독자 이익 강조 방식",
    "community_building": "공동체 감각 형성 방식 (있다면)",
    "examples": ["참여 유도가 잘 드러나는 실제 문장 3개 이상"]
  }},
  "c20_interjections_fillers": {{
    "title": "감탄사/추임새/습관어",
    "interjections": ["감탄사 목록 — 실제 표현 10개 이상 (와, 오, 아, 헉, 대박, 진짜, 하 등)"],
    "filler_starters": ["문장/단락 시작 습관어 목록 (그래서, 사실, 근데, 아무튼, 그리고 등)"],
    "affirmations": ["동의/긍정 추임새 목록 (맞아요, 그렇죠, 역시, 당연히 등)"],
    "hesitation_expressions": ["망설임/생각 표현 목록 (음..., 그러니까..., 어... 등)"],
    "excitement_expressions": ["흥분/놀람 표현 목록 (진짜요?!, 대박이에요!, 말이 돼요? 등)"],
    "self_talk_patterns": ["혼잣말/독백 표현 목록 (있다면)"],
    "frequency": "전반적 감탄사/추임새 사용 빈도 (1-10)",
    "position_pattern": "주로 어디서 씀 (문장 앞/중간/끝, 단독 줄)",
    "examples": ["감탄사/추임새가 잘 드러나는 실제 단락 발췌"]
  }},
  "c21_inline_formatting": {{
    "title": "인라인 서식 상세 (Naver SE3 기반 — 모든 서식 유형)",
    "font_switch_frequency": "본문 내 폰트 전환 빈도 (없음/가끔/자주/매우자주)",
    "font_switch_trigger": "폰트가 바뀌는 트리거 — 구체적으로 (예: '장소명에만 nanumbareunhipi 적용')",
    "font_families_used": ["이 블로그에서 실제 사용된 폰트 목록 — naver SE3 폰트명으로 기재 (예: nanumgothic, nanumbareuneunhipi, system)"],
    "font_switch_examples": ["폰트 전환 실제 텍스트 발췌 — 어느 부분에서 어떤 폰트"],
    "size_switch_pattern": "본문 내 크기 변화 패턴 — 구체적으로 (예: '중요 수치 fs19, 부연 설명 fs13')",
    "size_examples_by_level": {{"fs13": "어떤 텍스트에 사용?", "fs16": "어떤 텍스트에 사용?", "fs19": "어떤 텍스트에 사용?", "fs24": "어떤 텍스트에 사용?"}},
    "color_switch_pattern": "본문 내 색상 변화 패턴 (예: '키워드 강조 #0078cb, 날짜 #ff0000')",
    "color_switch_examples": ["색상 변화 실제 예시 5개 이상 — 텍스트+hex 색상"],
    "italic_usage": "기울임체 사용 여부와 패턴 (없음/가끔/자주, 어떤 텍스트에?)",
    "underline_usage": "밑줄 사용 여부와 패턴 (없음/가끔/자주, 어떤 텍스트에?)",
    "strikethrough_usage": "취소선 사용 여부와 패턴 (없음/가끔/자주, 어떤 맥락에?)",
    "background_color_pattern": "배경색(하이라이트) 사용 패턴 — 색상과 적용 텍스트 유형",
    "background_color_examples": ["배경색 강조 실제 예시 — 텍스트와 배경 hex 색상"],
    "center_align_pattern": "본문 내 가운데 정렬 사용 패턴 (없음/소제목만/특정단락/자주)",
    "box_quote_pattern": "박스/인용구 블록 사용 여부와 패턴 (없음/가끔/자주, 어떤 내용에?)",
    "combined_format_examples": ["볼드+색상, 크게+색상, 기울임+색상 등 복합 서식 실제 예시"],
    "naver_se3_pattern_guide": "이 블로그의 네이버 SE3 서식 재현 가이드 — 폰트/크기/색상/기울임/밑줄/배경 전환 규칙을 마커 사용법과 함께 구체적으로 (5~8문장)"
  }},
  "c22_content_patterns": {{
    "title": "콘텐츠 주제/정보 패턴",
    "main_topics": ["이 블로그가 주로 다루는 주제 카테고리 목록 (예: '지역 행사', '복지 정책', '관광 명소')"],
    "content_angle": "콘텐츠 접근 각도 — 구체적으로 (예: '시민 입장에서 혜택 중심으로 정보 전달', '현장 체험형 후기')",
    "must_include_elements": ["이 블로그 글에 항상 포함되는 정보 요소 목록 (예: '신청 방법', '문의처', '운영 시간', '가격')"],
    "never_include_elements": ["이 블로그가 절대 안 다루는 것 또는 피하는 표현"],
    "info_ordering": "정보 제시 순서 패턴 — 구체적으로 (예: '날짜/장소 먼저 → 대상 → 신청방법 → 혜택 강조 → CTA')",
    "local_terminology": ["이 블로그에서 자주 쓰는 지역명/기관명/특수 용어 목록"],
    "promotion_style": "홍보성 콘텐츠 작성 방식 (직접 홍보형/정보 제공형/체험 공유형 등)",
    "typical_post_template": "이 블로그의 전형적인 1개 포스트 구성을 요약 (섹션명과 각 섹션의 내용/길이/특징 포함)",
    "content_freshness_pattern": "최신성/시의성 강조 방식 (날짜 명시/계절 언급/이번 달 강조 등)",
    "examples": ["콘텐츠 패턴이 잘 드러나는 실제 글 제목+첫 단락 발췌 2~3개"]
  }}
}}

반드시 유효한 JSON으로만 응답하세요. 다른 텍스트는 포함하지 마세요."""

        response = client.models.generate_content(
            model="gemini-2.5-pro",
            contents=analysis_prompt
        )
        
        result_text = response.text.strip()
        result = parse_ai_json(result_text)
        result["blog_id"] = blog_id
        result["post_count"] = len(unique_posts)
        result["created_at"] = datetime.now().isoformat()
        
        # DNA 분석 결과 자동 저장
        dna_id = f"DNA_{blog_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        result["dna_id"] = dna_id
        dna_path = DNA_DIR / f"{dna_id}.json"
        with open(dna_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        
        return jsonify(result)
        
    except json.JSONDecodeError:
        return jsonify({"error": "AI 분석 결과 파싱 실패", "raw": result_text}), 500
    except Exception as e:
        return jsonify({"error": f"블로그 분석 실패: {str(e)}"}), 500


@app.route('/api/persona/business-analysis', methods=['POST'])
@login_required
def business_analysis():
    """[이전됨] 페르소나 기반 업무성격 분석은 email-persona 프로젝트로 이전"""
    return jsonify({"error": "이 기능은 이메일 도구로 이전되었습니다."}), 410




# ============================================================
# API: My Page (데이터 관리)
# ============================================================

@app.route('/api/mypage/personas', methods=['GET'])
@login_required
def mypage_personas():
    """[이전됨] 페르소나 목록 — 항상 빈 목록 반환"""
    return jsonify({"items": []})


@app.route('/api/mypage/personas/<client_id>', methods=['GET'])
@login_required
def mypage_persona_detail(client_id):
    return jsonify({"error": "이 기능은 이메일 도구로 이전되었습니다."}), 410


@app.route('/api/mypage/blogs', methods=['GET'])
@login_required
def mypage_blogs():
    """생성된 블로그 글 목록"""
    items = []
    for fp in OUTPUT_DIR.glob("BLOG_*.json"):
        try:
            data = load_blog_package(fp)
            versions = data.get("versions", [])
            title = versions[0].get("title", "") if versions else ""
            items.append({
                "id": data.get("output_id", fp.stem),
                "title": title,
                "client_id": data.get("client_id", ""),
                "version_count": len(versions),
                "created_at": data.get("created_at", ""),
                "filename": fp.name
            })
        except Exception:
            pass
    items.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return jsonify({"items": items})


@app.route('/api/mypage/blogs/<blog_id>', methods=['GET'])
@login_required
def mypage_blog_detail(blog_id):
    """블로그 상세 (3버전 포함)"""
    fp = OUTPUT_DIR / f"{blog_id}.json"
    if not fp.exists():
        return jsonify({"error": "블로그를 찾을 수 없습니다."}), 404
    data = load_blog_package(fp)
    return jsonify(data)


@app.route('/api/mypage/dna', methods=['GET'])
@login_required
def mypage_dna_list():
    """DNA 분석 결과 목록"""
    items = []
    for fp in DNA_DIR.glob("DNA_*.json"):
        try:
            with open(fp, 'r', encoding='utf-8') as f:
                data = json.load(f)
            items.append({
                "id": data.get("dna_id", fp.stem),
                "blog_id": data.get("blog_id", ""),
                "folder": data.get("folder", ""),
                "post_count": data.get("post_count", 0),
                "created_at": data.get("created_at", ""),
                "filename": fp.name,
                "has_c21": bool(data.get("c21_inline_formatting")),
                "has_c22": bool(data.get("c22_content_patterns")),
            })
        except Exception:
            pass
    items.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return jsonify({"items": items})


@app.route('/api/mypage/dna/<dna_id>', methods=['GET'])
@login_required
def mypage_dna_detail(dna_id):
    """DNA 상세"""
    fp = DNA_DIR / f"{dna_id}.json"
    if not fp.exists():
        return jsonify({"error": "DNA 분석 결과를 찾을 수 없습니다."}), 404
    with open(fp, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return jsonify(data)


@app.route('/api/mypage/business', methods=['GET'])
@login_required
def mypage_business_list():
    """업무적 성격 분석 결과 목록"""
    items = []
    for fp in BUSINESS_DIR.glob("BIZ_*.json"):
        try:
            with open(fp, 'r', encoding='utf-8') as f:
                data = json.load(f)
            bp = data.get("business_personality", {})
            items.append({
                "id": data.get("biz_id", fp.stem),
                "client_id": data.get("client_id", ""),
                "blog_folder": data.get("blog_folder", ""),
                "type": bp.get("type", ""),
                "created_at": data.get("created_at", ""),
                "filename": fp.name
            })
        except Exception:
            pass
    items.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return jsonify({"items": items})


@app.route('/api/mypage/business/<biz_id>', methods=['GET'])
@login_required
def mypage_business_detail(biz_id):
    """업무적 성격 상세"""
    fp = BUSINESS_DIR / f"{biz_id}.json"
    if not fp.exists():
        return jsonify({"error": "업무적 성격 분석 결과를 찾을 수 없습니다."}), 404
    with open(fp, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return jsonify(data)


@app.route('/api/mypage/dna/<dna_id>/tags', methods=['PATCH'])
@login_required
def update_dna_tags(dna_id):
    """DNA 활성 페르소나 태그 업데이트"""
    fp = DNA_DIR / f"{dna_id}.json"
    if not fp.exists():
        return jsonify({"error": "DNA를 찾을 수 없습니다."}), 404
    with open(fp, 'r', encoding='utf-8') as f:
        data = json.load(f)
    body = request.json or {}
    data["active_tags"] = body.get("active_tags", [])
    data["updated_at"] = datetime.now().isoformat()
    with open(fp, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return jsonify({"success": True})


@app.route('/api/mypage/<data_type>/<item_id>', methods=['DELETE'])
@login_required
def mypage_delete(data_type, item_id):
    """마이페이지 항목 삭제"""
    dir_map = {
        "blogs": OUTPUT_DIR,
        "dna": DNA_DIR,
        "business": BUSINESS_DIR
    }
    target_dir = dir_map.get(data_type)
    if not target_dir:
        return jsonify({"error": "잘못된 데이터 타입입니다."}), 400
    
    fp = target_dir / f"{item_id}.json"
    if not fp.exists():
        return jsonify({"error": "항목을 찾을 수 없습니다."}), 404
    
    fp.unlink()
    
    # 블로그의 경우 MD 파일도 삭제
    if data_type == "blogs":
        primary_md = target_dir / f"{item_id}.md"
        if primary_md.exists():
            primary_md.unlink()
        for md in target_dir.glob(f"{item_id}_*.md"):
            md.unlink()
    
    return jsonify({"success": True, "message": "삭제되었습니다."})


# ============================================================
# Google Docs Export
# ============================================================

def _get_valid_google_token():
    """유효한 access_token 반환. 만료 5분 전이면 refresh_token으로 자동 갱신."""
    import time as _time
    access_token = session.get('google_token', '')
    expires_at = session.get('google_token_expires_at', 0)
    refresh_token = session.get('google_refresh_token', '')

    # 만료까지 5분 미만이면 갱신 시도
    if expires_at and _time.time() > expires_at - 300:
        if not refresh_token:
            return None  # refresh_token 없으면 재로그인 필요
        resp = http_requests.post('https://oauth2.googleapis.com/token', data={
            'client_id': GOOGLE_CLIENT_ID,
            'client_secret': GOOGLE_CLIENT_SECRET,
            'refresh_token': refresh_token,
            'grant_type': 'refresh_token',
        })
        if resp.status_code == 200:
            new_token = resp.json()
            session['google_token'] = new_token.get('access_token', '')
            session['google_token_expires_at'] = _time.time() + new_token.get('expires_in', 3600)
            return session['google_token']
        return None  # 갱신 실패

    return access_token or None


@app.route('/api/export/google-docs', methods=['POST'])
@login_required
def export_to_google_docs():
    """블로그 글을 Google Docs로 내보내기"""
    access_token = _get_valid_google_token()
    if not access_token:
        return jsonify({"error": "Google 로그인이 필요합니다. 다시 로그인해주세요.", "login_required": True}), 401

    data = request.get_json()
    data_type = data.get('type', 'blogs')
    item_id = data.get('id', '')

    if not item_id:
        return jsonify({"error": "항목 ID가 필요합니다."}), 400

    # 데이터 로드
    if data_type == 'blogs':
        fp = OUTPUT_DIR / f"{item_id}.json"
    elif data_type == 'dna':
        fp = DNA_DIR / f"{item_id}.json"
    elif data_type == 'business':
        fp = BUSINESS_DIR / f"{item_id}.json"
    else:
        return jsonify({"error": "지원되지 않는 데이터 유형입니다."}), 400

    if not fp.exists():
        return jsonify({"error": "파일을 찾을 수 없습니다."}), 404

    with open(fp, 'r', encoding='utf-8') as f:
        content_data = json.load(f)

    # Google Docs 문서 내용 구성
    doc_title, doc_body = _build_doc_content(data_type, content_data, item_id)

    # 1) Google Docs 문서 생성
    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/json'
    }

    create_resp = http_requests.post(
        'https://docs.googleapis.com/v1/documents',
        headers=headers,
        json={'title': doc_title}
    )

    if create_resp.status_code == 401:
        session.pop('google_token', None)
        session.pop('google_refresh_token', None)
        return jsonify({"error": "Google 인증이 만료되었습니다. 다시 로그인해주세요.", "login_required": True}), 401

    if create_resp.status_code != 200:
        return jsonify({"error": f"Google Docs 생성 실패: {create_resp.text}"}), 500

    doc = create_resp.json()
    doc_id = doc['documentId']
    doc_url = f"https://docs.google.com/document/d/{doc_id}/edit"

    # 2) 문서에 내용 삽입
    update_resp = http_requests.post(
        f'https://docs.googleapis.com/v1/documents/{doc_id}:batchUpdate',
        headers=headers,
        json={'requests': doc_body}
    )

    if update_resp.status_code != 200:
        print(f"[WARN] Google Docs 내용 입력 실패: {update_resp.text}")

    return jsonify({
        "success": True,
        "doc_url": doc_url,
        "doc_id": doc_id,
        "message": "Google Docs에 내보내기 완료!"
    })


def _build_doc_content(data_type, data, item_id):
    """데이터 유형별 Google Docs 내용 생성"""
    requests_list = []
    idx = 1  # 커서 위치 (1-indexed)

    if data_type == 'blogs':
        data = ensure_blog_package_shape(data)
        title = data.get('title', item_id)
        doc_title = f"[블로그] {title}"

        versions = data.get('versions', [])
        for i, v in enumerate(versions):
            # 두 번째 버전부터 페이지 나누기 삽입
            if i > 0:
                requests_list.append({
                    'insertPageBreak': {'location': {'index': idx}}
                })
                idx += 1

            ver_label = v.get('version_label', v.get('version_type', ''))
            ver_title = v.get('title', '')
            ver_content = v.get('content', '')
            tags = v.get('tags', [])

            # 버전 라벨
            header = f"{ver_label} 버전\n{'='*40}\n\n"
            requests_list.append({
                'insertText': {'location': {'index': idx}, 'text': header}
            })
            idx += len(header)

            # 제목
            title_text = f"제목: {ver_title}\n\n"
            requests_list.append({
                'insertText': {'location': {'index': idx}, 'text': title_text}
            })
            idx += len(title_text)

            # 태그
            if tags:
                tags_text = f"태그: {', '.join(tags)}\n\n"
                requests_list.append({
                    'insertText': {'location': {'index': idx}, 'text': tags_text}
                })
                idx += len(tags_text)

            # 본문
            body_text = f"{ver_content}\n\n"
            requests_list.append({
                'insertText': {'location': {'index': idx}, 'text': body_text}
            })
            idx += len(body_text)

    elif data_type == 'dna':
        doc_title = f"[DNA 분석] {data.get('blog_id', item_id)}"
        text = json.dumps(data, ensure_ascii=False, indent=2)
        requests_list.append({
            'insertText': {'location': {'index': idx}, 'text': text}
        })

    elif data_type == 'business':
        bp = data.get('business_personality', {})
        doc_title = f"[업무성격] {bp.get('type', item_id)}"
        text = json.dumps(data, ensure_ascii=False, indent=2)
        requests_list.append({
            'insertText': {'location': {'index': idx}, 'text': text}
        })


    else:
        doc_title = f"내보내기 - {item_id}"
        text = json.dumps(data, ensure_ascii=False, indent=2)
        requests_list.append({
            'insertText': {'location': {'index': idx}, 'text': text}
        })

    return doc_title, requests_list


# ============================================================
# Run Server
# ============================================================

if __name__ == '__main__':
    print("=" * 60)
    print("  페르소나 도구 웹 대시보드")
    print("=" * 60)
    print("  블로그 자동화 도구")
    port = int(os.getenv('PORT', 5050))
    print(f"  서버: http://localhost:{port}")
    print(f"  출력 폴더: {OUTPUT_DIR}")
    print(f"  스타일 템플릿: {len(STYLE_TEMPLATES)}개")
    print("=" * 60)

    from waitress import serve
    print("  WSGI: waitress (production)")
    serve(app, host='0.0.0.0', port=port, threads=4)

