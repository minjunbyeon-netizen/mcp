#!/usr/bin/env python3
"""
auto-blog 블로그 생성기
사용법: python run_blog_generator.py
"""

import sys
import os
import json
import io
import re
import shutil
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Dict

# Windows 터미널 UTF-8 출력 설정 (중복 래핑 방지)
if sys.platform == 'win32' and not isinstance(sys.stdout, io.TextIOWrapper):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from utils import LoadingSpinner, parse_json_response, load_api_key, extract_text_from_file
from blog_storage import build_blog_package, save_blog_package, sanitize_filename_component
from material_pipeline import build_material_bundle
from offline_engines import generate_single_blog_offline

# API 키 로드
load_api_key("GEMINI_API_KEY")

from google import genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 페르소나 버전 관리 시스템
from persona_version_manager import (
    load_latest_persona,
    create_upgraded_version,
    get_feedback_history,
    save_feedback_history,
    compare_versions,
    generate_default_blog_config
)

# [M-5] 파일 텍스트 추출은 utils.py의 extract_text_from_file로 통합
# pdfplumber/olefile 등은 utils.py 내부에서 lazy import 처리

# 지원 파일 확장자
SUPPORTED_EXTENSIONS = ['.txt', '.pdf', '.hwp', '.docx', '.jpg', '.jpeg', '.png']


# [M-5] extract_text_from_file은 utils.py에서 import하여 사용
# (위 import 라인: from utils import ..., extract_text_from_file)


def get_file_type_icon(ext: str) -> str:
    """파일 확장자에 따른 아이콘 반환"""
    icons = {
        '.txt': '📄 TXT',
        '.pdf': '📕 PDF',
        '.hwp': '📘 HWP',
        '.jpg': '🖼️ JPG',
        '.jpeg': '🖼️ JPEG',
        '.png': '🖼️ PNG',
    }
    return icons.get(ext.lower(), '📁 FILE')


def select_press_release():
    """폴더 및 파일 선택 (공용 함수)"""
    
    # 하위 폴더 목록 가져오기
    subfolders = [f for f in INPUT_DIR.iterdir() if f.is_dir()]
    
    # 현재 폴더의 파일도 가져오기
    root_files = [
        f for f in INPUT_DIR.iterdir() 
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS and f.name.lower() != "readme.txt"
    ]
    
    # 폴더가 있으면 먼저 폴더 선택
    if subfolders:
        print("\n📁 사용 가능한 폴더:")
        print("-" * 50)
        
        # 0번: 현재 폴더 (루트)
        if root_files:
            print(f"  0. [현재 폴더] 📂 ({len(root_files)}개 파일)")
        
        # 하위 폴더 목록
        for i, folder in enumerate(subfolders, 1):
            # 폴더 내 파일 수 계산
            folder_files = [
                f for f in folder.iterdir() 
                if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
            ]
            print(f"  {i}. {folder.name} 📂 ({len(folder_files)}개 파일)")
        
        print("\n🔢 폴더 번호를 입력하세요 (0: 현재 폴더):")
        try:
            folder_choice = int(input(">>> ").strip())
            
            if folder_choice == 0:
                if not root_files:
                    print("❌ 현재 폴더에 파일이 없습니다.")
                    return None
                target_dir = INPUT_DIR
            elif 1 <= folder_choice <= len(subfolders):
                target_dir = subfolders[folder_choice - 1]
                print(f"\n✅ 선택된 폴더: {target_dir.name}")
            else:
                print("❌ 잘못된 번호입니다.")
                return None
        except ValueError:
            print("❌ 숫자를 입력해주세요.")
            return None
    else:
        target_dir = INPUT_DIR
    
    # 선택된 폴더에서 파일 목록 가져오기
    press_files = [
        f for f in target_dir.iterdir() 
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS and f.name.lower() != "readme.txt"
    ]
    
    if not press_files:
        print("\n❌ 보도자료 파일이 없습니다.")
        print(f"   📂 이 폴더에 파일을 넣어주세요:")
        print(f"   {target_dir}")
        print(f"   지원 형식: {', '.join(SUPPORTED_EXTENSIONS)}")
        return None
    
    # 파일 목록 표시
    print("\n📂 사용 가능한 보도자료:")
    print("-" * 50)
    for i, f in enumerate(press_files, 1):
        size_kb = f.stat().st_size / 1024
        file_icon = get_file_type_icon(f.suffix)
        print(f"  {i}. {f.stem} {file_icon}")
        print(f"     ({f.name}, {size_kb:.1f}KB)")
    
    # 번호로 선택 (다중 선택 지원)
    print("\n🔢 사용할 보도자료 번호를 입력하세요:")
    print("   💡 여러 파일: 1,2,3 또는 1-3 또는 all")
    try:
        choice_input = input(">>> ").strip().lower()
        
        selected_indices = []
        
        if choice_input == "all":
            # 전체 선택
            selected_indices = list(range(1, len(press_files) + 1))
        elif "-" in choice_input and "," not in choice_input:
            # 범위 선택 (예: 1-3)
            parts = choice_input.split("-")
            start, end = int(parts[0]), int(parts[1])
            selected_indices = list(range(start, end + 1))
        elif "," in choice_input:
            # 개별 선택 (예: 1,3,5)
            selected_indices = [int(x.strip()) for x in choice_input.split(",")]
        else:
            # 단일 선택
            selected_indices = [int(choice_input)]
        
        # 유효성 검사
        for idx in selected_indices:
            if idx < 1 or idx > len(press_files):
                print(f"❌ 잘못된 번호입니다: {idx}")
                return None
        
        selected_files = [press_files[i - 1] for i in selected_indices]
        
    except ValueError:
        print("❌ 올바른 형식으로 입력해주세요. (예: 1 또는 1,2,3 또는 1-3)")
        return None
    
    # 선택된 파일 표시
    if len(selected_files) == 1:
        print(f"\n✅ 선택: {selected_files[0].name}")
    else:
        print(f"\n✅ 선택: {len(selected_files)}개 파일")
        for f in selected_files:
            print(f"   - {f.name}")
    
    # 모든 파일 읽기 및 합치기
    all_texts = []
    for selected_file in selected_files:
        try:
            text = extract_text_from_file(selected_file)
            if text.strip():
                if len(selected_files) > 1:
                    # 여러 파일인 경우 구분선 추가
                    all_texts.append(f"\n\n===== {selected_file.name} =====\n\n{text}")
                else:
                    all_texts.append(text)
        except Exception as e:
            print(f"⚠️ 파일 읽기 실패: {selected_file.name} - {e}")
    
    if not all_texts:
        print("❌ 파일에서 텍스트를 추출할 수 없습니다.")
        return None
    
    press_release = "\n".join(all_texts)
    print(f"📄 보도자료 길이: {len(press_release):,} 글자")
    return press_release


# 경로 설정
PERSONA_DIR = Path(__file__).parent / "output" / "personas"
PERSONA_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR = Path.home() / "mcp-data" / "outputs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Word 파일 전용 저장 위치
WORD_OUTPUT_DIR = Path(__file__).parent / "output" / "blog"
WORD_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Google Drive 자동 동기 폴더
GDRIVE_DIR = Path(r"G:\내 드라이브\01_auto_system\02_Archive_to_blog")

# 입력 폴더 (보도자료 텍스트 파일 넣는 곳)
INPUT_DIR = Path(__file__).parent / "input" / "2_blog_writing"
INPUT_DIR.mkdir(parents=True, exist_ok=True)


def list_personas():
    """저장된 페르소나 목록"""
    personas = []
    for file_path in PERSONA_DIR.glob("*.json"):
        # 피드백 파일 제외
        if file_path.name.endswith("_feedback.json"):
            continue
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if "client_id" in data and "persona_analysis" in data:
                    pa = data["persona_analysis"]
                    # [M-1] 스키마 정규화: formality_analysis.overall_score 사용
                    formality = (
                        pa.get("formality_analysis", {}).get("overall_score")
                        or pa.get("formality_level", {}).get("score", 5)
                    )
                    personas.append({
                        "client_id": data["client_id"],
                        "client_name": data["client_name"],
                        "organization": data["organization"],
                        "formality": formality
                    })
        except Exception:
            pass
    return personas


def _analyze_press_release(client, press_release: str) -> dict | None:
    """
    [V4-M2] 단계 1: 배포자료 분석 → press_analysis dict 반환.
    실패 시 None 반환 (폴백: 기존 단일호출 방식으로 전환).
    """
    prompt = f"""배포자료를 분석하여 아래 JSON 형식으로만 반환하세요. 다른 텍스트 없이 JSON만 출력하세요.

배포자료:
{press_release}

출력 JSON:
{{
  "key_messages": ["핵심 메시지1", "핵심 메시지2", "핵심 메시지3"],
  "target_audience": "주요 독자층 설명",
  "call_to_action": "독자에게 유도하고 싶은 행동",
  "emphasis_points": ["강조할 포인트1", "강조할 포인트2"]
}}"""
    try:
        resp = client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
        return parse_json_response(resp.text)
    except Exception:
        return None


def _design_blog_structure(client, press_analysis: dict, persona_data: dict) -> dict | None:
    """
    [V4-M2] 단계 2: 블로그 구조 설계 → blog_design dict 반환.
    실패 시 None 반환.
    """
    client_name = persona_data.get("client_name", "")
    prompt = f"""핵심 메시지와 페르소나를 고려해 블로그 구조를 JSON으로 설계하세요. 다른 텍스트 없이 JSON만 출력하세요.

페르소나: {client_name}
핵심 메시지: {json.dumps(press_analysis.get('key_messages', []), ensure_ascii=False)}
독자층: {press_analysis.get('target_audience', '')}
행동 유도: {press_analysis.get('call_to_action', '')}
강조 포인트: {json.dumps(press_analysis.get('emphasis_points', []), ensure_ascii=False)}

출력 JSON:
{{
  "intro_hook": "독자의 관심을 끄는 인트로 방향",
  "sections": [
    {{"title": "섹션 제목1", "content_points": ["다룰 내용1", "다룰 내용2"]}},
    {{"title": "섹션 제목2", "content_points": ["다룰 내용1"]}}
  ],
  "outro": "마무리 및 행동 유도 방향"
}}"""
    try:
        resp = client.models.generate_content(model='gemini-2.0-flash', contents=prompt)
        return parse_json_response(resp.text)
    except Exception:
        return None


def _build_persona_prompt_parts(persona_data: dict) -> dict:
    """페르소나 공통 파라미터를 한 번만 추출해 dict로 반환."""
    persona_analysis = persona_data.get("persona_analysis", {})
    formality_analysis = persona_analysis.get("formality_analysis", {})
    formality = (
        formality_analysis.get("overall_score")
        or persona_analysis.get("formality_level", {}).get("score", 5)
        or 5
    )
    writing_dna = persona_analysis.get("writing_dna", {})
    writing_chars = writing_dna if writing_dna else persona_analysis.get("writing_characteristics", {})
    comm_style = persona_analysis.get("communication_style", {})
    positive_triggers = persona_analysis.get("positive_triggers", {})
    green_flags = (
        persona_analysis.get("green_flags")
        or positive_triggers.get("favorite_expressions", [])
    )
    sensitive_areas = persona_analysis.get("sensitive_areas", {})
    red_flags = (
        persona_analysis.get("red_flags")
        or sensitive_areas.get("absolute_dont", {}).get("expressions", [])
    )

    if formality >= 8:
        sentence_ending = "합쇼체(~습니다, ~입니다) 위주로 작성"
        tone_desc = "매우 격식있고 공식적인 톤"
    elif formality >= 6:
        sentence_ending = "해요체(~해요, ~이에요) 70% + 합쇼체(~습니다) 30% 혼용"
        tone_desc = "정중하되 친근한 톤"
    elif formality >= 4:
        sentence_ending = "해요체(~해요) 위주, 가끔 반말 섞어도 무방"
        tone_desc = "편안하고 친근한 톤"
    else:
        sentence_ending = "반말(~해, ~야) 또는 해요체 자유롭게 사용"
        tone_desc = "매우 캐주얼하고 친구같은 톤"

    sentence_structure = writing_dna.get("sentence_structure", {})
    sentence_length = (
        sentence_structure.get("avg_length")
        or writing_chars.get("sentence_length", "medium")
    )
    if sentence_length == "short":
        length_guide = "짧고 간결한 문장 (15자 내외)"
    elif sentence_length == "long":
        length_guide = "상세하고 긴 문장 (30자 이상)"
    else:
        length_guide = "적당한 길이의 문장 (20자 내외)"

    emoji_usage_raw = (
        comm_style.get("emotional_expression", {}).get("emoji_usage")
        or writing_chars.get("emoji_usage", "moderate")
    )
    if isinstance(emoji_usage_raw, int):
        if emoji_usage_raw >= 7:
            emoji_usage = "frequent"
        elif emoji_usage_raw <= 3:
            emoji_usage = "none"
        else:
            emoji_usage = "moderate"
    else:
        emoji_usage = emoji_usage_raw

    if emoji_usage == "frequent":
        emoji_freq = "문단마다 1-2개 이상 이모티콘 필수"
    elif emoji_usage == "none":
        emoji_freq = "이모티콘 사용 금지"
    else:
        emoji_freq = "적절히 이모티콘 사용"

    writing_config = persona_data.get("blog_writing_config", {})
    human_config = writing_config.get("humanization", {})
    narrative_flow = human_config.get("narrative_flow", "flexible")
    insight_ratio = human_config.get("personal_insight_ratio", 0.3)
    catchphrases = human_config.get("human_catchphrases", [])
    avoid_cliches = human_config.get("avoid_cliches", [])

    unified_persona = persona_data.get("unified_persona")
    unified_guide_section = ""
    if unified_persona:
        unified_guide_section = f"""
  "unified_persona_guide": {{
    "writing_guide": "{unified_persona.get('writing_guide', '')}",
    "title_guide": "{unified_persona.get('title_guide', '')}",
    "structure_guide": "{unified_persona.get('structure_guide', '')}",
    "tone_guide": "{unified_persona.get('tone_guide', '')}",
    "hashtag_guide": "{unified_persona.get('hashtag_guide', '')}",
    "priority": "이 가이드는 블로그 DNA와 카톡 DNA를 병합한 최우선 가이드입니다. 위 strict_writing_rules보다 우선 적용하세요."
  }},"""

    return dict(
        formality=formality,
        tone_desc=tone_desc,
        sentence_ending=sentence_ending,
        length_guide=length_guide,
        emoji_freq=emoji_freq,
        narrative_flow=narrative_flow,
        insight_ratio=insight_ratio,
        catchphrases=catchphrases,
        avoid_cliches=avoid_cliches,
        green_flags=green_flags,
        red_flags=red_flags,
        comm_style=comm_style,
        unified_guide_section=unified_guide_section,
    )


def _generate_final_blog(
    client,
    blog_design: dict,
    persona_data: dict,
    keywords_str: str,
    few_shots: list,
    press_release_fallback: str = "",
) -> dict | None:
    """
    [V4-M2] 단계 3: 최종 블로그 생성.
    blog_design이 있으면 구조 기반, 없으면 press_release_fallback 직접 사용.
    few_shots: few_shot_examples 리스트 (없으면 []).
    title_variants 3개를 output_schema에 포함.
    """
    client_name = persona_data.get("client_name", "")
    custom_prompt = persona_data.get("custom_prompt", "")
    p = _build_persona_prompt_parts(persona_data)

    # 입력 컨텍스트: blog_design 우선, 없으면 press_release 원본
    if blog_design:
        input_context = f"""블로그 구조 설계:
인트로 방향: {blog_design.get('intro_hook', '')}
섹션 구성: {json.dumps(blog_design.get('sections', []), ensure_ascii=False)}
아웃트로 방향: {blog_design.get('outro', '')}"""
    else:
        input_context = f"배포자료:\n{press_release_fallback}"

    # Few-shot 섹션 (블로그 DNA — 글쓰기 스타일의 1순위 기준)
    few_shot_section = ""
    if few_shots:
        examples = []
        for idx, ex in enumerate(few_shots[:2], 1):
            examples.append(f"예시{idx} 제목: {ex.get('title', '')}\n예시{idx} 본문 발췌:\n{ex.get('excerpt', '')}")
        few_shot_section = f"""
  "blog_style_primary_reference": {{
    "priority": "HIGHEST — 아래 실제 블로그 글 예시가 글쓰기 스타일의 최우선 기준입니다. 문단 구조, 문장 길이, 어휘 수준, 소제목 방식, 마무리 패턴을 이 예시에서 학습하세요.",
    "CRITICAL_WARNING": "예시의 주제·소재·정보(정월대보름, 야외놀이장 등 예시 속 내용)를 절대 블로그에 옮기지 마세요. 오직 글쓰기 스타일(문체, 어휘 선택, 문단 구성)만 참고하고, 블로그의 실제 주제와 내용은 반드시 input_context의 배포자료에서만 가져오세요.",
    "examples": {json.dumps(examples, ensure_ascii=False)}
  }},"""

    blog_prompt = f"""
{{
  "system_settings": {{
    "role": "블로그 콘텐츠 전문 작가",
    "objective": "배포자료를 아래 [블로그 스타일 기준]에 맞는 완성도 높은 블로그 글로 변환. 광고주의 카카오톡 말투를 그대로 옮기는 것이 아니라, 블로그 글쓰기 DNA에서 학습한 스타일로 작성",
    "two_layer_rule": "Layer 1(블로그 DNA) = 실제 글쓰기 스타일 기준 (문장 구조·어휘·단락 구성). Layer 2(페르소나) = 톤·격식도·감성 조절용 참고. Layer 1이 항상 우선합니다."
  }},
  "input_context": {{
    "blog_design": "{input_context}",
    "target_keywords": ["{keywords_str}"],
    "persona_custom_request": "{custom_prompt}"
  }},
  {few_shot_section}
  "persona_tone_guide": {{
    "note": "아래는 글쓰기 스타일이 아닌 톤(격식도·감성) 조절용 참고 데이터입니다. 카카오톡 말버릇을 블로그에 그대로 쓰지 마세요.",
    "organization": "{persona_data.get('organization', '')}",
    "formality_level": "{p['formality']}/10",
    "tone_adjustment": {{
      "overall_tone": "{p['tone_desc']}",
      "sentence_ending_rule": "{p['sentence_ending']}",
      "sentence_length": "{p['length_guide']}",
      "emoji_usage": "{p['emoji_freq']}",
      "emotional_weight": "{p['comm_style'].get('emotional_expression', dict()).get('level', 'neutral') if isinstance(p['comm_style'].get('emotional_expression'), dict) else p['comm_style'].get('emotional_tone', 'neutral')}"
    }},
    "personality_traits_reference_only": {{
      "note": "아래 표현은 광고주의 성격을 이해하기 위한 참고용입니다. 블로그 본문에 직접 인용하지 마세요.",
      "kakao_communication_patterns": {json.dumps(p['catchphrases'], ensure_ascii=False)},
      "positive_signals": {json.dumps(p['green_flags'], ensure_ascii=False)},
      "avoid_signals": {json.dumps(p['red_flags'], ensure_ascii=False)}
    }}
  }},
    "blog_writing_rules": {{
      "style_source": "위 blog_style_primary_reference 예시의 글쓰기 패턴을 따르세요",
      "humanization": {{
        "narrative_flow_style": "{p['narrative_flow']}",
        "personal_insight": "본문의 {int(p['insight_ratio']*100)}% 이상은 배포자료에 없는 필자의 시각이나 해석을 포함할 것 (실명 노출 금지)",
        "burstiness_rule": "문장 길이를 의도적으로 다양하게 섞으세요 (아주 짧은 문장과 긴 문장의 조화)",
        "cliche_blacklist": {json.dumps(p['avoid_cliches'], ensure_ascii=False)}
      }},
      "banned_characters": [
        "마크다운 볼드 기호 완전 금지: ** 기호를 블로그 본문에 단 한 번도 사용하지 마세요. **굵게** 형식 절대 금지",
        "마크다운 헤더 완전 금지: ## 기호를 사용하지 마세요. 소제목은 「꺽쇠 괄호」로만 표현하세요",
        "마크다운 인용/목록 금지: >, - (목록 기호)로 시작하는 줄 금지",
        "스마트 따옴표 금지: \" \" ' ' 대신 일반 따옴표 사용",
        "말줄임표 금지: ... 또는 … 사용 금지",
        "중복/혼합 기호 금지: '.,', ';;', '!!', '??' 등 지저분한 기호 조합 절대 사용 금지"
      ]
    }},
  {p['unified_guide_section']}
  "content_structure": {{
    "intro": "독자의 관심을 끄는 시작 (블로그 DNA 스타일 기준)",
    "body": "배포자료 핵심 내용을 블로그 글쓰기 스타일로 풀어서 설명",
    "outro": "마무리 및 행동 유도 (블로그 DNA 마무리 패턴 따름)"
  }},
  "formatting_rules": {{
    "header_style": "소제목은 「꺽쇠 괄호」 사용. ## 마크다운 헤더 절대 금지",
    "emphasis_style": "강조 시 「」 괄호 활용. ** 볼드 기호 절대 금지",
    "layout_style": "가독성을 위해 적절한 줄바꿈 사용",
    "punctuation_cleanup": "문장 끝 외에 불필요한 마침표(.) 사용 자제",
    "emoji_minimalism": "이모지는 문단 끝에만 1개 이내로 아주 가끔씩만 사용",
    "anonymity_rule": "본문과 제목 어디에도 작성자 개인의 실명을 절대 노출하지 마세요. 소속 기관명({persona_data.get('organization', '')})은 사용 가능합니다"
  }},
  "task_requirements": {{
    "seo_optimization": {{
      "title": "60자 이내, 키워드 포함, 클릭 유도형 제목",
      "meta_description": "155자 이내 요약",
      "keyword_integration": "키워드를 본문에 자연스럽게 3회 이상 포함"
    }},
    "length": "1,500 ~ 2,000자 분량"
  }},
  "output_schema": {{
    "description": "반드시 아래 JSON 포맷으로만 출력",
    "format": {{
      "title_variants": ["클릭유도형 제목 (60자 이내, 작성자 실명 제외)", "정보형 제목 (60자 이내, 작성자 실명 제외)", "감성형 제목 (60자 이내, 작성자 실명 제외)"],
      "title": "title_variants[0] 과 동일한 값",
      "content": "Markdown 형식 본문 String (페르소나 톤 100% 반영)",
      "tags": ["태그1", "태그2", "태그3", "태그4", "태그5"],
      "meta_description": "155자 이내 String"
    }}
  }}
}}

**핵심 원칙**: 블로그 DNA 예시(blog_style_primary_reference)가 글쓰기 스타일의 최우선 기준입니다. 페르소나는 톤·격식도 조절용 참고이며, 카카오톡 말버릇을 블로그에 그대로 옮기지 마세요.
**절대 금지**: 본문과 제목에 작성자의 실명({client_name})을 절대 노출하지 마세요. 기관명은 사용 가능합니다.
(다른 텍스트 없이 오직 JSON만 출력해주세요)
"""
    try:
        resp = client.models.generate_content(model='gemini-2.0-flash', contents=blog_prompt)
        return parse_json_response(resp.text)
    except Exception:
        return None


def _self_review(client, blog_content: dict, persona_data: dict) -> tuple[int, dict | None]:
    """
    [V4-M3] 단계 4: 자기 검토 패스.
    블로그 글이 페르소나 기준에 맞는지 평가 후 total < 70이면 revised_blog 반환.
    실패 시 (0, None) 반환 → 호출자가 원본 사용.
    재생성/재귀/while 루프 절대 없음. AI가 직접 수정한 버전만 반환.
    """
    client_name = persona_data.get("client_name", "")
    p = _build_persona_prompt_parts(persona_data)

    review_prompt = f"""당신은 '{client_name}' 페르소나 블로그 품질 검토 전문가입니다.
아래 블로그 글이 페르소나 기준에 맞는지 10개 항목으로 평가하고, 필요 시 직접 수정된 버전을 제공하세요.

[페르소나 기준]
- 격식도: {p['formality']}/10
- 톤: {p['tone_desc']}
- 문장 끝: {p['sentence_ending']}
- 이모지 규칙: {p['emoji_freq']}
- 그린플래그(써야 할 표현): {json.dumps(p['green_flags'], ensure_ascii=False)}
- 레드플래그(쓰면 안 되는 표현): {json.dumps(p['red_flags'], ensure_ascii=False)}

[검토할 블로그 글]
제목: {blog_content.get('title', '')}
본문:
{blog_content.get('content', '')}

[평가 항목 10개] (각 10점 만점)
1. 페르소나 말투 일치 (문장 끝, 호칭)
2. 격식도 적정 여부
3. 그린플래그 표현 사용
4. 레드플래그 표현 미사용
5. 문장 길이 다양성 (burstiness)
6. 금지 기호 미사용 (스마트 따옴표, 말줄임표)
7. 이모지 규칙 준수
8. 본문 분량 (1500~2000자)
9. 개인 의견/통찰 포함 여부
10. 전체 가독성

[출력 JSON 형식 - JSON만 출력]
{{
  "scores": {{
    "persona_tone": 0,
    "formality": 0,
    "green_flags": 0,
    "red_flags": 0,
    "burstiness": 0,
    "banned_chars": 0,
    "emoji_rule": 0,
    "length": 0,
    "personal_insight": 0,
    "readability": 0
  }},
  "total": 0,
  "issues": ["발견된 문제점1", "발견된 문제점2"],
  "revised_blog": {{
    "title_variants": ["클릭유도형 제목", "정보형 제목", "감성형 제목"],
    "title": "수정된 제목",
    "content": "수정된 본문 (문제점을 직접 고친 버전)",
    "tags": ["태그1", "태그2"],
    "meta_description": "수정된 메타 설명"
  }}
}}"""
    try:
        resp = client.models.generate_content(model='gemini-2.0-flash', contents=review_prompt)
        review_data = parse_json_response(resp.text)
        if not isinstance(review_data, dict):
            return 0, None
        total = review_data.get("total", 0)
        revised = review_data.get("revised_blog")
        return total, revised
    except Exception:
        return 0, None


def generate_blog_post(client_id: str, press_release: str, target_keywords: list = None):
    """블로그 글 생성 (V4: 3단계 파이프라인 + 자기검토 + 제목 3변형)"""

    print(f"\n{'='*50}")
    print(f"  AI 블로그 글 생성 시작")
    print(f"{'='*50}")

    # 페르소나 최신 버전 로드
    result = load_latest_persona(client_id)
    if not result:
        print(f"페르소나를 찾을 수 없습니다: {client_id}")
        return None

    persona_data, version, persona_file = result
    client_name = persona_data.get("client_name", client_id)

    # blog_writing_config가 없으면 자동 생성
    if "blog_writing_config" not in persona_data:
        persona_data["blog_writing_config"] = generate_default_blog_config(persona_data)

    print(f"  페르소나: {client_name} (v{version})")
    print(f"  보도자료: {len(press_release):,} 글자")
    print(f"{'='*50}\n")

    # Step 1: API 연결
    print("[블로그 생성 1/4] AI 연결 준비")  # [티켓 #T002] 소비자팀 V4-MAJOR-2: 진행률 레이블 명확화
    spinner = LoadingSpinner("Gemini AI 연결 중")
    spinner.start()
    api_key = os.getenv("GEMINI_API_KEY")
    client = None
    if api_key:
        client = genai.Client(api_key=api_key)
        spinner.stop("API 연결 완료")
    else:
        spinner.stop("오프라인 모드")
        print("ℹ️ GEMINI_API_KEY가 없어 규칙 기반 초안으로 진행합니다.")

    # few_shot_examples 추출 (blog_dna에 있으면 사용, 없으면 빈 리스트)
    blog_dna = persona_data.get("blog_dna", {})
    few_shots = blog_dna.get("few_shot_examples", [])
    material_bundle = build_material_bundle(direct_text=press_release)
    prompt_source = material_bundle.get("briefing") or press_release

    # 키워드 문자열
    keywords_str = ", ".join(target_keywords) if target_keywords else ""

    if client:
        # Step 2: 배포자료 분석 (단계 1)
        print("\n[블로그 생성 2/4] 배포자료 분석")  # [티켓 #T002]
        spinner = LoadingSpinner("배포자료 핵심 메시지 추출 중")
        spinner.start()
        press_analysis = _analyze_press_release(client, prompt_source)
        if press_analysis:
            spinner.stop("배포자료 분석 완료")
        else:
            spinner.stop("분석 실패 — 기존 방식으로 진행")

        # Step 3: 블로그 구조 설계 (단계 2) — press_analysis 있을 때만
        blog_design = None
        if press_analysis:
            print("\n[블로그 생성 3/4] 블로그 구조 설계")  # [티켓 #T002]
            spinner = LoadingSpinner("페르소나 맞춤 블로그 구조 설계 중")
            spinner.start()
            blog_design = _design_blog_structure(client, press_analysis, persona_data)
            if blog_design:
                spinner.stop("블로그 구조 설계 완료")
            else:
                spinner.stop("구조 설계 실패 — 원본 배포자료 직접 사용")
        else:
            print("\n[블로그 생성 3/4] 블로그 구조 설계 건너뜀")  # [티켓 #T002]

        # Step 4: 최종 블로그 생성 (단계 3)
        print("\n[블로그 생성 4/4] 블로그 글 작성")  # [티켓 #T002]
        spinner = LoadingSpinner("AI가 페르소나 스타일로 글을 작성하고 있습니다")
        spinner.start()

        blog_content = _generate_final_blog(
            client,
            blog_design,
            persona_data,
            keywords_str,
            few_shots,
            press_release_fallback=prompt_source,
        )

        if not blog_content:
            spinner.stop("오류 발생")
            print("\n블로그 생성에 실패했습니다.")
            return None
        spinner.stop("블로그 글 생성 완료")

        # title_variants 하위 호환 보정: 없으면 title로 채움
        if "title_variants" not in blog_content or not isinstance(blog_content.get("title_variants"), list):
            blog_content["title_variants"] = [blog_content.get("title", ""), "", ""]

        # Step 5: 자기 검토 패스 (단계 4)
        print("\n  자기 검토 중...")
        spinner = LoadingSpinner("AI 품질 자기 검토 중")
        spinner.start()
        quality_score, revised = _self_review(client, blog_content, persona_data)
        if quality_score > 0:
            if quality_score < 70 and revised:
                spinner.stop(f"자기 검토 완료 (점수: {quality_score}/100) — 수정본 적용")
                # revised_blog의 title_variants 보정
                if "title_variants" not in revised or not isinstance(revised.get("title_variants"), list):
                    revised["title_variants"] = [revised.get("title", ""), "", ""]
                blog_content = revised
            else:
                spinner.stop(f"자기 검토 완료 (점수: {quality_score}/100) — 원본 유지")
        else:
            spinner.stop("자기 검토 실패 — 원본 사용")
    else:
        print("\n[블로그 생성 2/4] 오프라인 자료 요약")
        spinner = LoadingSpinner("자료에서 핵심 포인트를 정리하고 있습니다")
        spinner.start()
        blog_content = generate_single_blog_offline(persona_data, material_bundle, target_keywords)
        spinner.stop("오프라인 초안 생성 완료")

    # title 필드를 title_variants[0]과 일치시킴
    variants = blog_content.get("title_variants", [])
    if variants and variants[0]:
        blog_content["title"] = variants[0]

    # 파일 저장
    print("\n  파일 저장 중")
    spinner = LoadingSpinner("Word/Markdown 파일 생성 중")
    spinner.start()
    
    # 저장
    output_id = f"BLOG_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    blog_data = build_blog_package(
        output_id=output_id,
        client_id=client_id,
        client_name=client_name,
        versions=[{
            "version_type": "primary",
            "version_label": "메인",
            "title": blog_content.get("title", ""),
            "content": blog_content.get("content", ""),
            "tags": blog_content.get("tags", []),
            "meta_description": blog_content.get("meta_description", ""),
        }],
        source_bundle=material_bundle,
        extra={
            "content": blog_content,
            "title_variants": blog_content.get("title_variants", []),
        },
    )

    # JSON/Markdown 저장
    json_path, markdown_paths = save_blog_package(blog_data, OUTPUT_DIR)
    md_path = markdown_paths.get("primary", OUTPUT_DIR / f"{output_id}.md")
    
    # Word 파일 생성 (별도 위치에 페르소나명_제목_날짜 형식으로)
    # 파일명에 사용할 수 없는 문자 제거
    safe_title = sanitize_filename_component(blog_content['title'], limit=30)
    safe_client_name = sanitize_filename_component(client_name, limit=30).replace(' ', '_')
    date_str = datetime.now().strftime('%Y%m%d')
    docx_filename = f"{safe_client_name}_{safe_title}_{date_str}.docx"
    docx_path = WORD_OUTPUT_DIR / docx_filename
    doc = Document()
    
    # 기본 스타일에 한글 폰트 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    
    # 제목 추가
    title = doc.add_heading(blog_content['title'], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 제목에도 한글 폰트 적용
    for run in title.runs:
        run.font.name = '맑은 고딕'
    
    # 본문 추가 (마크다운 파싱 간소화)
    content = blog_content['content']
    paragraphs = content.split('\n\n')
    
    for para in paragraphs:
        if para.strip():
            # 소제목 처리 (「 」)
            if para.strip().startswith('「') and para.strip().endswith('」'):
                p = doc.add_heading(para.strip()[1:-1].strip(), level=2)
                for run in p.runs:
                    run.font.name = '맑은 고딕'
            # 구분선 처리
            elif para.strip() == '• • • • •':
                p = doc.add_paragraph('─' * 30)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = '맑은 고딕'
            # 이미지 자리 표시
            elif para.strip().startswith('[이미지'):
                p = doc.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.italic = True
                run.font.name = '맑은 고딕'
            else:
                # 일반 본문 - **굵게** 처리
                p = doc.add_paragraph()
                parts = para.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(11)
                    if i % 2 == 1:  # 홀수 인덱스는 굵게
                        run.bold = True
    
    # 태그 추가
    doc.add_paragraph()
    tags_para = doc.add_paragraph()
    tags_run = tags_para.add_run(f"태그: {', '.join(blog_content['tags'])}")
    tags_run.italic = True
    tags_run.font.name = '맑은 고딕'
    
    # 메타 설명 추가
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"메타 설명: {blog_content['meta_description']}")
    meta_run.italic = True
    meta_run.font.name = '맑은 고딕'
    
    doc.save(str(docx_path))
    
    # Google Drive에도 복사 (폴더가 있으면)
    gdrive_docx_path = None
    if GDRIVE_DIR.exists():
        try:
            gdrive_docx_path = GDRIVE_DIR / docx_filename
            shutil.copy2(docx_path, gdrive_docx_path)
            print(f"\r  [☁️] Google Drive 업로드 완료" + " " * 20)
        except Exception as e:
            print(f"\r  [⚠️] Google Drive 복사 실패: {e}")
    
    spinner.stop("파일 저장 완료")
    
    # 버전 정보 추가
    blog_data["version"] = version
    blog_data["client_id"] = client_id
    
    return blog_data, md_path, docx_path, gdrive_docx_path


def _parse_free_text_feedback_with_ai(free_text: str) -> Dict:
    """
    [M-4] 자유텍스트 피드백 → Gemini API → 조정 파라미터 변환.

    입력 예시: "말투가 너무 딱딱해요. 이모지도 좀 더 써주세요."
    출력 예시: {"tone_details.sentence_ending": "casual", "formatting.emoji_positions": ["intro","body","outro"]}

    Returns:
        dict: adjustments (persona_version_manager.create_upgraded_version에 전달 가능한 형태)
    """
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        print("❌ GEMINI_API_KEY가 없어 AI 분석을 건너뜁니다.")
        return {}

    client = genai.Client(api_key=api_key)

    prompt = f"""
당신은 블로그 스타일 조정 파라미터 변환 전문가입니다.
아래 사용자의 자유 피드백을 분석하여 블로그 설정 조정 파라미터(JSON)로 변환해주세요.

【사용자 피드백】
{free_text}

【출력 파라미터 규칙】
아래 키를 사용하여 변경이 필요한 항목만 포함하세요 (변경 불필요한 항목은 제외):

- content_rules.max_length: 정수 (1000~3000, 글 길이 조절)
- content_rules.paragraph_length: "short" | "medium" | "long"
- content_rules.technical_terms: "simplify" | "avoid" | "maintain"
- tone_details.sentence_length: "short" | "medium" | "long"
- tone_details.punctuation_style: "formal" | "friendly" | "casual"
- formatting.emoji_positions: 배열 ([], ["intro"], ["intro","outro"], ["intro","body","outro"], ["all"])
- humanization.personal_insight_ratio: 0.0~1.0 (개인 의견 비율)
- humanization.narrative_flow: "structured" | "flexible" | "storytelling" | "chatty"
- structure.body_sections: 정수 (2~5, 본문 섹션 수)
- formality_adjustment: "+1" | "-1" | "+2" | "-2" (격식도 조절)

【출력 형식】
반드시 JSON만 출력하세요 (설명 텍스트 없이).
예시: {{"content_rules.max_length": 1500, "formatting.emoji_positions": ["intro","outro"]}}
"""

    try:
        spinner = LoadingSpinner("자유 피드백 AI 분석 중")
        spinner.start()
        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=prompt
        )
        spinner.stop("분석 완료")

        adjustments = parse_json_response(response.text)
        return adjustments if isinstance(adjustments, dict) else {}

    except Exception as e:
        print(f"\n피드백 분석 중 오류가 발생했습니다. 잠시 후 다시 시도해주세요.")
        return {}


def _handle_free_text_feedback(client_id: str, version: int, output_id: str) -> bool:
    """
    [M-4] 자유텍스트 피드백 처리 흐름.
    입력 → AI 파라미터 변환 → 사용자 확인 → create_upgraded_version 호출
    """
    print("\n📝 자유 피드백을 입력해주세요.")
    print("   (예: '말투가 너무 딱딱해요', '이모지 좀 더 써주세요', '글이 너무 길어요')")
    free_text = input(">>> ").strip()

    if not free_text:
        print("❌ 피드백을 입력해주세요.")
        return False

    # AI로 파라미터 변환
    adjustments = _parse_free_text_feedback_with_ai(free_text)

    if not adjustments:
        print("\n피드백을 분석하지 못했습니다. 내용을 확인해주세요.")
        print("   (Tip: '말투가 딱딱해요', '이모지를 더 써주세요' 처럼 구체적으로 적어주시면 더 잘 반영됩니다.)")
        print("   번호 선택 방식으로 다시 피드백하시겠습니까? (Y/n): ", end="")
        retry = input().strip().lower()
        if retry != "n":
            return False  # 호출자(_collect_feedback_loop)가 번호 선택 메뉴로 안내
        return False

    # 변환 결과 표시
    print("\n🔄 AI가 다음과 같이 해석했습니다:")
    for key, value in adjustments.items():
        print(f"   - {key}: {value}")

    print("\n이 설정으로 새 버전을 생성할까요? (Y/n): ", end="")
    confirm = input().strip().lower()

    if confirm == 'n':
        print("취소되었습니다.")
        return False

    upgrade_reason = f"자유 피드백 (AI 해석): {free_text}"
    new_persona = create_upgraded_version(client_id, adjustments, upgrade_reason)

    if new_persona:
        print("\n✅ 다음 블로그부터 개선된 스타일로 작성됩니다!")
        # 피드백 히스토리에도 기록
        feedback_data = get_feedback_history(client_id)
        feedback_entry = {
            "timestamp": datetime.now().isoformat(),
            "blog_id": output_id,
            "version": version,
            "rating": None,
            "issues": [free_text],
            "adjustments_made": adjustments,
            "feedback_type": "free_text"
        }
        feedback_data["feedback_history"].append(feedback_entry)
        feedback_data["learning_stats"]["total_blogs"] = len(feedback_data["feedback_history"])
        save_feedback_history(client_id, feedback_data)
        return True

    return False


def collect_feedback_and_upgrade(blog_data: Dict) -> bool:
    """블로그 생성 후 피드백 수집 및 자동 업그레이드"""
    
    client_id = blog_data.get("client_id")
    version = blog_data.get("version", 1)
    output_id = blog_data.get("output_id")
    
    print("\n" + "=" * 60)
    print("📊 이번 블로그 글은 어떠셨나요?")
    print("=" * 60)
    print("1. 완벽해요! ⭐⭐⭐⭐⭐")
    print("2. 좋아요 ⭐⭐⭐⭐")
    print("3. 괜찮아요 ⭐⭐⭐")
    print("4. 아쉬워요 ⭐⭐")
    print("5. 다시 써주세요 ⭐")
    print("f. 자유 피드백 입력 (AI가 분석해서 자동 반영)")
    print("0. 피드백 건너뛰기")

    try:
        rating_input = input("\n>>> ").strip().lower()

        if rating_input == "0":
            return False

        # [M-4] 자유텍스트 피드백 분기
        if rating_input == "f":
            return _handle_free_text_feedback(client_id, version, output_id)

        rating = int(rating_input)
        if rating < 1 or rating > 5:
            print("❌ 잘못된 입력입니다.")
            return False

    except ValueError:
        print("❌ 숫자를 입력해주세요.")
        return False
    
    # 피드백 히스토리 로드
    feedback_data = get_feedback_history(client_id)
    
    # 새 피드백 추가
    feedback_entry = {
        "timestamp": datetime.now().isoformat(),
        "blog_id": output_id,
        "version": version,
        "rating": rating,
        "issues": [],
        "adjustments_made": {}
    }
    
    # 평점이 4점 이하면 문제점 파악
    if rating <= 4:
        print("\n🔧 어떤 부분이 아쉬웠나요? (번호로 선택, 여러 개 가능: 1,2,3)")
        print("1. 말투/어미가 안 맞아요")
        print("2. 너무 길어요")
        print("3. 너무 짧아요")
        print("4. 이모티콘이 너무 많아요")
        print("5. 이모티콘이 너무 적어요")
        print("6. 문장이 너무 길어요")
        print("7. 문장이 너무 짧아요")
        print("8. 전문용어가 어려워요")
        print("9. 구조가 이상해요")
        print("0. 기타 (직접 입력)")
        
        issues_input = input("\n>>> ").strip()
        
        if not issues_input:
            print("❌ 문제점을 선택해주세요.")
            return False
        
        issue_map = {
            "1": "말투/어미",
            "2": "글이 너무 깁",
            "3": "글이 너무 짧음",
            "4": "이모티콘 과다",
            "5": "이모티콘 부족",
            "6": "문장 너무 깁",
            "7": "문장 너무 짧음",
            "8": "전문용어 어려움",
            "9": "구조 문제"
        }
        
        selected_issues = []
        for num in issues_input.split(","):
            num = num.strip()
            if num == "0":
                custom_issue = input("기타 문제점을 입력하세요: ").strip()
                if custom_issue:
                    selected_issues.append(custom_issue)
            elif num in issue_map:
                selected_issues.append(issue_map[num])
        
        feedback_entry["issues"] = selected_issues
        
        # 자동 조정 로직
        adjustments = {}
        
        for issue in selected_issues:
            if "말투" in issue or "어미" in issue:
                # 격식도 1단계 상향 (최대 10)
                adjustments["formality_adjustment"] = "+1"
                
            elif "너무 깁" in issue:
                adjustments["content_rules.max_length"] = 1500
                adjustments["content_rules.paragraph_length"] = "short"
                
            elif "너무 짧" in issue:
                adjustments["content_rules.max_length"] = 2500
                adjustments["content_rules.paragraph_length"] = "medium"
                
            elif "이모티콘 과다" in issue:
                adjustments["formatting.emoji_positions"] = ["intro", "outro"]
                
            elif "이모티콘 부족" in issue:
                adjustments["formatting.emoji_positions"] = ["intro", "body", "outro"]
                
            elif "문장 너무 깁" in issue:
                adjustments["tone_details.sentence_length"] = "short"
                
            elif "문장 너무 짧" in issue:
                adjustments["tone_details.sentence_length"] = "long"
                
            elif "전문용어" in issue:
                adjustments["content_rules.technical_terms"] = "avoid"
                
            elif "구조" in issue:
                adjustments["structure.body_sections"] = 3
        
        feedback_entry["adjustments_made"] = adjustments
        
        # 조정 사항이 있으면 새 버전 생성
        if adjustments:
            print("\n🔄 피드백을 반영하여 새 버전을 생성할까요? (Y/n): ", end="")
            create_new = input().strip().lower()
            
            if create_new != 'n':
                upgrade_reason = f"사용자 피드백 (평점: {rating}/5): " + ", ".join(selected_issues)
                new_persona = create_upgraded_version(client_id, adjustments, upgrade_reason)
                
                if new_persona:
                    print("\n✅ 다음 블로그부터 개선된 스타일로 작성됩니다! 🎉")
    
    else:
        print(f"\n✨ 감사합니다! 현재 설정(v{version})을 유지합니다.")
    
    # 피드백 히스토리 업데이트
    feedback_data["feedback_history"].append(feedback_entry)
    
    # 통계 업데이트
    all_ratings = [f["rating"] for f in feedback_data["feedback_history"] if f.get("rating") is not None]
    feedback_data["learning_stats"]["total_blogs"] = len(feedback_data["feedback_history"])
    if all_ratings:
        feedback_data["learning_stats"]["average_rating"] = round(sum(all_ratings) / len(all_ratings), 1)
    
    # 최근 5개 vs 전체 평균 비교
    if len(all_ratings) >= 5:
        recent_avg = sum(all_ratings[-5:]) / 5
        overall_avg = sum(all_ratings) / len(all_ratings)
        feedback_data["learning_stats"]["improvement_trend"] = round(recent_avg - overall_avg, 1)
    
    # 공통 이슈 집계
    common_issues = {}
    for f in feedback_data["feedback_history"]:
        for issue in f.get("issues", []):
            common_issues[issue] = common_issues.get(issue, 0) + 1
    feedback_data["learning_stats"]["common_issues"] = common_issues
    
    # 저장
    save_feedback_history(client_id, feedback_data)
    
    return True


def generate_blog_with_persona(client_id: str):
    """페르소나 추출 후 바로 블로그 생성 (연계 호출용)"""
    print("\n" + "=" * 60)
    print("auto-blog 블로그 생성기")
    print("=" * 60)
    
    # 보도자료 선택 (폴더 및 파일)
    press_release = select_press_release()
    if not press_release:
        return
    
    # SEO 키워드 (선택)
    print("\n🔑 SEO 키워드를 입력하세요 (쉼표로 구분, 없으면 엔터):")
    keywords_input = input(">>> ").strip()
    keywords = [k.strip() for k in keywords_input.split(",")] if keywords_input else None
    
    # 블로그 생성
    result = generate_blog_post(client_id, press_release, keywords)
    
    if result:
        blog_data, md_path, docx_path, gdrive_path = result
        blog = blog_data["content"]
        
        print("\n" + "=" * 60)
        print("✅ 블로그 글 생성 완료!")
        print("=" * 60)
        
        print(f"\n📌 제목: {blog['title']}")
        print(f"🏷️ 태그: {', '.join(blog['tags'])}")
        print(f"\n💾 저장 위치:")
        print(f"   - Word: {docx_path}")
        if gdrive_path:
            print(f"   - ☁️ Google Drive: {gdrive_path}")
        
        # 피드백 수집
        collect_feedback_and_upgrade(blog_data)
        
        # 폴더 열기 옵션
        print("\n" + "=" * 60)
        print("📂 블로그 폴더를 여시겠습니까? (Y/n): ", end="")
        open_folder = input().strip().lower()
        if open_folder != 'n':
            subprocess.run(['explorer', str(WORD_OUTPUT_DIR)])
            print("   폴더를 열었습니다.")
    else:
        print("\n❌ 블로그 생성에 실패했습니다.")


def main():
    print("=" * 60)
    print("auto-blog 블로그 생성기")
    print("=" * 60)
    
    # 페르소나 목록 표시
    personas = list_personas()
    if not personas:
        print("\n페르소나가 없습니다. 페르소나를 먼저 추출해야 블로그를 생성할 수 있습니다.")
        print("   지금 페르소나를 만들어볼까요? (Y/n): ", end="")
        ans = input().strip().lower()
        if ans != "n":
            import subprocess as _sp
            _sp.run(["python", str(Path(__file__).parent / "run_persona_test.py")])
        return

    print("\n📋 사용 가능한 페르소나:")
    print("-" * 50)
    for i, p in enumerate(personas, 1):
        print(f"  {i}. {p['client_name']}")
        print(f"     ({p['organization']}) - 격식도: {p['formality']}/10")
    
    # 페르소나 선택
    print("\n🔢 사용할 페르소나 번호를 입력하세요:")
    try:
        choice = int(input(">>> ").strip())
        selected = personas[choice - 1]
        client_id = selected["client_id"]
    except (ValueError, IndexError):
        print("❌ 잘못된 선택입니다.")
        return
    
    print(f"\n✅ 선택된 페르소나: {selected['client_name']}")
    
    # 보도자료 선택 (폴더 및 파일)
    press_release = select_press_release()
    if not press_release:
        return
    
    # SEO 키워드 (선택)
    print("\n🔑 SEO 키워드를 입력하세요 (쉼표로 구분, 없으면 엔터):")
    keywords_input = input(">>> ").strip()
    keywords = [k.strip() for k in keywords_input.split(",")] if keywords_input else None
    
    # 블로그 생성
    result = generate_blog_post(client_id, press_release, keywords)
    
    if result:
        blog_data, md_path, docx_path, gdrive_path = result
        blog = blog_data["content"]
        
        print("\n" + "=" * 60)
        print("✅ 블로그 글 생성 완료!")
        print("=" * 60)
        
        print(f"\n📌 제목: {blog['title']}")
        print(f"🏷️ 태그: {', '.join(blog['tags'])}")
        print(f"\n💾 저장 위치:")
        print(f"   - Word: {docx_path}")
        if gdrive_path:
            print(f"   - ☁️ Google Drive: {gdrive_path}")
        
        # 피드백 수집
        collect_feedback_and_upgrade(blog_data)
        
        # 폴더 열기 옵션
        print("\n" + "=" * 60)
        print("📂 블로그 폴더를 여시겠습니까? (Y/n): ", end="")
        open_folder = input().strip().lower()
        if open_folder != 'n':
            subprocess.run(['explorer', str(WORD_OUTPUT_DIR)])
            print("   폴더를 열었습니다.")
    else:
        print("\n❌ 블로그 생성에 실패했습니다.")


def batch_blog_generation():
    """
    [M-3] 배치 블로그 생성.
    다수 페르소나 선택 → 1개 보도자료 → 순차 생성 (2초 딜레이)
    """
    import time

    print("=" * 60)
    print("📦 배치 블로그 생성 (다수 페르소나 x 1개 보도자료)")
    print("=" * 60)

    # 페르소나 목록 표시
    personas = list_personas()
    if not personas:
        print("\n페르소나가 없습니다. 페르소나를 먼저 추출해야 블로그를 생성할 수 있습니다.")
        print("   지금 페르소나를 만들어볼까요? (Y/n): ", end="")
        ans = input().strip().lower()
        if ans != "n":
            import subprocess as _sp
            _sp.run(["python", str(Path(__file__).parent / "run_persona_test.py")])
        return

    print("\n📋 사용 가능한 페르소나:")
    print("-" * 50)
    for i, p in enumerate(personas, 1):
        print(f"  {i}. {p['client_name']} ({p['organization']}) - 격식도: {p['formality']}/10")

    # 다중 페르소나 선택
    print("\n🔢 사용할 페르소나 번호를 입력하세요.")
    print("   💡 여러 페르소나: 1,2,3 또는 1-3 또는 all")
    try:
        choice_input = input(">>> ").strip().lower()

        selected_indices = []
        if choice_input == "all":
            selected_indices = list(range(1, len(personas) + 1))
        elif "-" in choice_input and "," not in choice_input:
            parts = choice_input.split("-")
            start, end = int(parts[0]), int(parts[1])
            selected_indices = list(range(start, end + 1))
        elif "," in choice_input:
            selected_indices = [int(x.strip()) for x in choice_input.split(",")]
        else:
            selected_indices = [int(choice_input)]

        for idx in selected_indices:
            if idx < 1 or idx > len(personas):
                print(f"❌ 잘못된 번호입니다: {idx}")
                return

        selected_personas = [personas[i - 1] for i in selected_indices]

    except (ValueError, IndexError):
        print("❌ 올바른 형식으로 입력해주세요. (예: 1 또는 1,2,3 또는 1-3)")
        return

    print(f"\n✅ 선택된 페르소나 {len(selected_personas)}개:")
    for p in selected_personas:
        print(f"   - {p['client_name']} ({p['organization']})")

    # 보도자료 선택 (1개)
    print("\n📄 보도자료를 선택하세요 (모든 페르소나에 동일하게 적용)")
    press_release = select_press_release()
    if not press_release:
        return

    # SEO 키워드 (공통)
    print("\n🔑 SEO 키워드를 입력하세요 (쉼표로 구분, 없으면 엔터):")
    keywords_input = input(">>> ").strip()
    keywords = [k.strip() for k in keywords_input.split(",")] if keywords_input else None

    # 배치 실행
    print("\n" + "=" * 60)
    print(f"🚀 배치 실행 시작: {len(selected_personas)}개 블로그")
    print("=" * 60)

    success_count = 0
    fail_count = 0
    results_summary = []

    for idx, persona in enumerate(selected_personas, 1):
        client_id = persona["client_id"]
        client_name = persona["client_name"]

        print(f"\n[{idx}/{len(selected_personas)}] {client_name} 처리 중...")

        try:
            result = generate_blog_post(client_id, press_release, keywords)

            if result:
                blog_data, md_path, docx_path, gdrive_path = result
                blog = blog_data["content"]
                success_count += 1
                results_summary.append({
                    "persona": client_name,
                    "title": blog.get("title", ""),
                    "docx": str(docx_path),
                    "status": "성공"
                })
                print(f"  [OK] {client_name}: {blog.get('title', '')}")
                print(f"       저장: {docx_path}")
            else:
                fail_count += 1
                results_summary.append({
                    "persona": client_name,
                    "title": "",
                    "docx": "",
                    "status": "실패"
                })
                print(f"  [FAIL] {client_name}: 블로그 생성 실패")

        except Exception as e:
            fail_count += 1
            results_summary.append({
                "persona": client_name,
                "title": "",
                "docx": "",
                "status": f"오류: {e}"
            })
            print(f"  [ERROR] {client_name}: {e}")

        # rate limit 대비 딜레이 (마지막 항목 제외)
        if idx < len(selected_personas):
            print("  다음 블로그 생성을 준비하고 있습니다...")
            time.sleep(2)

    # 결과 요약
    print("\n" + "=" * 60)
    print(f"📊 배치 완료: 성공 {success_count}개 / 실패 {fail_count}개")
    print("=" * 60)
    for r in results_summary:
        status_icon = "[OK]" if r["status"] == "성공" else "[FAIL]"
        print(f"  {status_icon} {r['persona']}: {r['title'] or r['status']}")

    if success_count > 0:
        print(f"\n💾 저장 위치: {WORD_OUTPUT_DIR}")
        print("\n📂 블로그 폴더를 여시겠습니까? (Y/n): ", end="")
        open_folder = input().strip().lower()
        if open_folder != 'n':
            subprocess.run(['explorer', str(WORD_OUTPUT_DIR)])
            print("   폴더를 열었습니다.")


if __name__ == "__main__":
    main()
